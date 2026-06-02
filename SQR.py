import os
import re
import json
import datetime
import time
import urllib.request
import urllib.error
import urllib.parse
from io import BytesIO
from functools import wraps

import jwt
import mysql.connector
from mysql.connector import pooling
from flask import Flask, request, jsonify, send_from_directory, send_file, render_template
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None
    Pt = None
    Inches = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
except Exception:
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    HRFlowable = None
    A4 = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    TA_CENTER = None
    TA_LEFT = None
    colors = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")
UPLOAD_DIR = os.path.join(BASE_DIR, os.getenv("UPLOAD_FOLDER", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = Flask(__name__, template_folder=TEMPLATES_DIR, static_folder=STATIC_DIR)
CORS(app, resources={r"/api/*": {"origins": os.getenv("CORS_ORIGINS", "*").split(",")}})

app.config["SECRET_KEY"] = os.getenv("SQR_SECRET_KEY") or os.getenv("SECRET_KEY") or "CHANGE_THIS_SECRET_KEY_BEFORE_DEPLOYMENT"
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_CONTENT_LENGTH", str(50 * 1024 * 1024)))
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR

DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost").strip(),
    "port": int(os.getenv("DB_PORT", "3306").strip() or "3306"),
    "user": os.getenv("DB_USER", "root").strip(),
    "password": os.getenv("DB_PASSWORD", "").strip(),
    "database": os.getenv("DB_NAME", "railway").strip(),
    "connection_timeout": int(os.getenv("DB_TIMEOUT", "10")),
    "autocommit": True,
}

pool = None
GEMINI_API_KEY = (os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or "").strip()
GEMINI_MODEL = (os.getenv("GEMINI_MODEL") or "gemini-2.5-flash").strip()

XAI_API_KEY = (os.getenv("XAI_API_KEY") or os.getenv("GROK_API_KEY") or "").strip()
XAI_MODEL = (os.getenv("XAI_MODEL") or os.getenv("GROK_MODEL") or "grok-4.3").strip()

OPENAI_API_KEY = (os.getenv("OPENAI_API_KEY") or "").strip()
OPENAI_MODEL = (os.getenv("OPENAI_MODEL") or "gpt-4.1-mini").strip()

AI_PROVIDER = (os.getenv("AI_PROVIDER") or "gemini").strip().lower()
AI_TIMEOUT = int(os.getenv("AI_TIMEOUT", "45"))
AI_MAX_RETRIES = max(1, int(os.getenv("AI_MAX_RETRIES", "3")))

# Gemini is called through the official REST endpoint. xAI/OpenAI use the OpenAI SDK.
# Use AI_PROVIDER=gemini, xai, openai, or auto.
gemini_client = bool(GEMINI_API_KEY)

TECH_SKILLS = [
    "python", "java", "javascript", "typescript", "html", "css", "sql", "mysql", "postgresql",
    "react", "node", "flask", "django", "api", "rest", "git", "github", "docker", "aws", "azure",
    "linux", "security", "cybersecurity", "networking", "forensics", "wireshark", "burp suite", "database",
    "machine learning", "ai", "data analysis", "data engineering", "etl", "cloud", "mongodb", "communication",
    "teamwork", "problem solving", "devops", "kubernetes", "php", "c++", "go", "rust", "swift",
    "figma", "ui", "ux", "testing", "automation", "incident response", "siem"
]

COURSE_LEVEL_META = {
    "beginner": {"label": "Beginner", "class": "level-beginner", "hex": "#22c55e"},
    "intermediate": {"label": "Intermediate", "class": "level-intermediate", "hex": "#eab308"},
    "advanced": {"label": "Advanced", "class": "level-advanced", "hex": "#ef4444"},
}

SPECIALIZATION_HINTS = {
    "cybersecurity": ["security", "cyber", "network", "linux", "forensics", "burp", "wireshark", "soc", "vulnerability"],
    "digital forensics": ["forensics", "evidence", "investigation", "incident", "malware", "security"],
    "software engineering": ["software", "java", "python", "problem", "api", "backend", "testing"],
    "web development": ["html", "css", "javascript", "react", "frontend", "backend", "node", "flask"],
    "data science": ["data", "python", "sql", "analysis", "machine learning", "statistics", "visualization"],
    "artificial intelligence": ["ai", "machine learning", "automation", "model", "python", "nlp", "vision"],
    "cloud computing": ["cloud", "aws", "azure", "deployment", "server", "docker", "devops"],
    "database administration": ["database", "sql", "mysql", "postgresql", "queries", "schema", "admin"],
    "computer networks": ["network", "tcp", "ip", "routing", "switching", "security", "linux"],
    "ui/ux engineering": ["ui", "ux", "design", "interface", "figma", "frontend", "user"],
}


def get_db():
    global pool
    if pool is None:
        pool = pooling.MySQLConnectionPool(
            pool_name="sqr_pool",
            pool_size=int(os.getenv("DB_POOL_SIZE", "5")),
            **DB_CONFIG
        )
    return pool.get_connection()


def query_db(sql, params=None, fetchone=False, fetchall=False, commit=False):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    try:
        cursor.execute(sql, params or ())
        result = None
        if fetchone:
            result = cursor.fetchone()
        elif fetchall:
            result = cursor.fetchall()
        if commit:
            db.commit()
            result = cursor.lastrowid
        return result
    except Exception:
        if commit:
            db.rollback()
        raise
    finally:
        cursor.close()
        db.close()


def exec_db(sql, params=None):
    return query_db(sql, params=params, commit=True)


def get_json():
    return request.get_json(silent=True) or {}


def safe_text(value):
    return str(value or "").strip()


def safe_int(value, default=0):
    try:
        if value is None or value == "":
            return default
        return int(float(value))
    except Exception:
        return default


def table_exists(table_name):
    try:
        row = query_db(
            """
            SELECT COUNT(*) AS total
            FROM INFORMATION_SCHEMA.TABLES
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
            """,
            (DB_CONFIG["database"], table_name),
            fetchone=True
        )
        return bool(row and row.get("total"))
    except Exception:
        return False


def column_exists(table_name, column_name):
    try:
        row = query_db(
            """
            SELECT COUNT(*) AS total
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s AND COLUMN_NAME=%s
            """,
            (DB_CONFIG["database"], table_name, column_name),
            fetchone=True
        )
        return bool(row and row.get("total"))
    except Exception:
        return False


def first_existing_column(table_name, names):
    for name in names:
        if column_exists(table_name, name):
            return name
    return names[0]


def row_value(row, *names):
    for name in names:
        if isinstance(row, dict) and row.get(name) not in [None, ""]:
            return row.get(name)
    return None


def upload_url(filename):
    value = safe_text(filename)
    if not value:
        return ""
    if value.startswith("http://") or value.startswith("https://") or value.startswith("/uploads/"):
        return value
    return f"/uploads/{value}"


def normalize_level(level):
    value = safe_text(level).lower() or "beginner"
    aliases = {
        "begginer": "beginner",
        "beginner": "beginner",
        "medium": "intermediate",
        "intermidiete": "intermediate",
        "intermediate": "intermediate",
        "advance": "advanced",
        "advanced": "advanced",
    }
    return aliases.get(value, "beginner")


def add_level_meta(course):
    level = normalize_level(course.get("level"))
    course["level"] = level
    course["level_badge"] = COURSE_LEVEL_META[level]
    return course


def normalize_specialization(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "specialization_id")
    row["specialization_id"] = row["id"]
    row["image_url"] = upload_url(row_value(row, "image", "image_url"))
    row["skills"] = row_value(row, "skills", "required_skills") or ""
    row["description"] = row.get("description") or ""
    return row


def normalize_course(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "course_id")
    row["course_id"] = row["id"]
    row["specialization_id"] = row_value(row, "specialization_id", "spec_id")
    row["spec_id"] = row["specialization_id"]
    row["link"] = row_value(row, "link", "course_link") or ""
    row["course_link"] = row["link"]
    row["image_url"] = upload_url(row_value(row, "image", "image_url", "thumbnail"))
    row["video_url"] = upload_url(row_value(row, "video", "video_url", "media_url"))
    row["specialization_name"] = row_value(row, "specialization_name", "specialization") or ""
    return add_level_meta(row)


def normalize_quiz(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "quiz_id")
    row["quiz_id"] = row["id"]
    row["course_id"] = row_value(row, "course_id", "course")
    return row


def normalize_question(row, include_answer=False):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "question_id")
    row["question_id"] = row["id"]
    row["question"] = row_value(row, "question", "question_text") or ""
    row["options"] = [
        row_value(row, "option1", "option_a") or "",
        row_value(row, "option2", "option_b") or "",
        row_value(row, "option3", "option_c") or "",
        row_value(row, "option4", "option_d") or "",
    ]
    if not include_answer:
        row.pop("answer", None)
        row.pop("correct_answer", None)
    return row


def normalize_job(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "job_id")
    row["job_id"] = row["id"]
    row["skills"] = row_value(row, "skills", "required_skills") or ""
    row["required_skills"] = row["skills"]
    row["salary"] = row_value(row, "salary", "average_salary") or ""
    row["average_salary"] = row["salary"]
    row["link"] = row_value(row, "link", "job_link") or ""
    row["job_link"] = row["link"]
    row["specialization"] = row_value(row, "specialization_name", "specialization") or ""
    return row


def normalize_certificate(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "certificate_id", "certification_id")
    row["certificate_id"] = row["id"]
    row["certification_id"] = row["id"]
    row["title"] = row_value(row, "title", "name") or "Certificate"
    row["name"] = row["title"]
    row["link"] = row_value(row, "link", "official_link", "certificate_url") or ""
    row["official_link"] = row["link"]
    row["specialization_id"] = row_value(row, "specialization_id", "spec_id")
    return row


def table_column_names(table_name):
    try:
        rows = query_db(
            """
            SELECT COLUMN_NAME
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
            """,
            (DB_CONFIG["database"], table_name),
            fetchall=True
        ) or []
        return {row.get("COLUMN_NAME") for row in rows}
    except Exception:
        return set()


def _condition_for_existing_columns(table_name, alias, candidates):
    existing = table_column_names(table_name)
    pieces = []
    for column in candidates:
        if column in existing:
            pieces.append(f"{alias}.`{column}`=%s")
    if not pieces:
        pieces.append(f"{alias}.`id`=%s")
    return "(" + " OR ".join(pieces) + ")", len(pieces)


def fetch_specialization_by_any_id(spec_id):
    spec_id = safe_int(spec_id, None)
    if not spec_id:
        return None
    table = admin_spec_table()
    condition, cols = where_existing_id(table, "s", ["id", "specialization_id", "spec_id"])
    return query_db(f"SELECT s.* FROM `{table}` s WHERE {condition} LIMIT 1", tuple([spec_id] * len(cols)), fetchone=True)


def fetch_course_by_any_id(course_id):
    course_id = safe_int(course_id, None)
    if not course_id:
        return None
    condition, count = _condition_for_existing_columns("courses", "c", ["id", "course_id"])
    return query_db(f"SELECT c.* FROM courses c WHERE {condition} LIMIT 1", tuple([course_id] * count), fetchone=True)


def fetch_quiz_by_any_id(quiz_id):
    quiz_id = safe_int(quiz_id, None)
    if not quiz_id:
        return None
    condition, count = _condition_for_existing_columns("quizzes", "q", ["id", "quiz_id"])
    return query_db(f"SELECT q.* FROM quizzes q WHERE {condition} LIMIT 1", tuple([quiz_id] * count), fetchone=True)


def real_id(row):
    return safe_int(row_value(row or {}, "id", "specialization_id", "course_id", "quiz_id", "job_id"), None)


def existing_table(*names):
    """Return the first existing table from a compatibility list."""
    for name in names:
        try:
            if table_exists(name):
                return name
        except Exception:
            continue
    return names[0]


def existing_pk(table_name, candidates):
    existing = table_column_names(table_name)
    for column in candidates:
        if column in existing:
            return column
    return candidates[0]


def existing_columns(table_name, candidates):
    existing = table_column_names(table_name)
    return [column for column in candidates if column in existing]


def coalesce_existing(alias, table_name, candidates, fallback="NULL"):
    cols = existing_columns(table_name, candidates)
    refs = [f"{alias}.`{column}`" for column in cols]
    if not refs:
        return fallback
    if len(refs) == 1:
        return refs[0]
    return "COALESCE(" + ", ".join(refs) + ")"


def where_existing_id(table_name, alias, candidates):
    cols = existing_columns(table_name, candidates)
    if not cols:
        cols = [candidates[0]]
    prefix = f"{alias}." if alias else ""
    condition = " OR ".join([f"{prefix}`{column}`=%s" for column in cols])
    return "(" + condition + ")", cols


def admin_insert_dynamic(table_name, payload):
    clean = {}
    existing = table_column_names(table_name)
    for key, value in (payload or {}).items():
        if key in existing:
            clean[key] = value
    if not clean:
        raise ValueError(f"No matching columns found for {table_name}")
    columns = list(clean.keys())
    placeholders = ",".join(["%s"] * len(columns))
    column_sql = ",".join([f"`{column}`" for column in columns])
    new_id = query_db(
        f"INSERT INTO `{table_name}` ({column_sql}) VALUES ({placeholders})",
        tuple(clean[column] for column in columns),
        commit=True
    )
    try:
        sync_compatibility_alias_columns()
    except Exception as exc:
        print("alias sync after insert skipped:", exc)
    return new_id


def admin_update_dynamic(table_name, item_id, id_candidates, payload):
    existing = table_column_names(table_name)
    clean = {}
    for key, value in (payload or {}).items():
        if key in existing:
            clean[key] = value
    if not clean:
        raise ValueError(f"No matching columns found for {table_name}")
    set_sql = ", ".join([f"`{column}`=%s" for column in clean.keys()])
    where_sql, cols = where_existing_id(table_name, "", id_candidates)
    params = list(clean.values()) + [item_id] * len(cols)
    query_db(f"UPDATE `{table_name}` SET {set_sql} WHERE {where_sql}", tuple(params), commit=True)


def admin_delete_dynamic(table_name, item_id, id_candidates):
    where_sql, cols = where_existing_id(table_name, "", id_candidates)
    query_db(f"DELETE FROM `{table_name}` WHERE {where_sql}", tuple([item_id] * len(cols)), commit=True)


def admin_order_sql(table_name, candidates):
    return f"`{existing_pk(table_name, candidates)}` DESC"


def admin_spec_table():
    return existing_table("specializations", "specialization")


def normalize_inserted_image(data):
    return save_file("image") or safe_text(data.get("image") or data.get("image_url"))


def normalize_inserted_video(data):
    return save_file("video") or safe_text(data.get("video") or data.get("video_url"))


def specialization_filter_sql(alias="c"):
    existing = table_column_names("courses")
    pieces = []
    if "spec_id" in existing:
        pieces.append(f"{alias}.spec_id=%s")
    if "specialization_id" in existing:
        pieces.append(f"{alias}.specialization_id=%s")
    if not pieces:
        pieces.append(f"{alias}.spec_id=%s")
    return "(" + " OR ".join(pieces) + ")", len(pieces)


def set_course_progress(user_id, course_id, progress_value, completed=False):
    progress_value = max(0, min(100, safe_int(progress_value, 0)))
    status = "Completed" if completed or progress_value >= 100 else ("In Progress" if progress_value > 0 else "Not Started")
    query_db(
        """
        UPDATE course_enrollments
        SET progress=GREATEST(COALESCE(progress,0),%s),
            progress_percentage=GREATEST(COALESCE(progress_percentage,0),%s),
            status=%s,
            completed_at=CASE WHEN %s >= 100 THEN CURRENT_TIMESTAMP ELSE completed_at END
        WHERE user_id=%s AND course_id=%s
        """,
        (progress_value, progress_value, status, progress_value, user_id, course_id),
        commit=True
    )


def ensure_course_enrollment(user_id, course_id, start_progress=0):
    start_progress = max(0, min(100, safe_int(start_progress, 0)))
    status = "In Progress" if start_progress > 0 else "Not Started"
    query_db(
        """
        INSERT INTO course_enrollments (user_id, course_id, progress, progress_percentage, status, enrolled_at)
        VALUES (%s,%s,%s,%s,%s,CURRENT_TIMESTAMP)
        ON DUPLICATE KEY UPDATE
            progress=GREATEST(COALESCE(progress,0), VALUES(progress)),
            progress_percentage=GREATEST(COALESCE(progress_percentage,0), VALUES(progress_percentage)),
            status=CASE
                WHEN LOWER(status) IN ('Completed') THEN status
                WHEN VALUES(progress) > 0 THEN 'In Progress'
                ELSE status
            END
        """,
        (user_id, course_id, start_progress, start_progress, status),
        commit=True
    )

def clean_user(user):
    if not user:
        return None
    user = dict(user)
    user.pop("password", None)
    user["id"] = row_value(user, "id", "user_id")
    user["user_id"] = user["id"]
    user["current_mode"] = user.get("current_mode") or user.get("role") or "student"
    user["banned"] = safe_int(row_value(user, "banned", "is_banned"), 0)
    return user


def strong_password(password):
    password = str(password or "")
    return (
        len(password) >= 8
        and re.search(r"[A-Z]", password)
        and re.search(r"[a-z]", password)
        and re.search(r"[0-9]", password)
        and re.search(r"[^A-Za-z0-9]", password)
        and not re.search(r"\s", password)
    )



def generate_username(name, email):
    base = (email.split("@")[0] if email and "@" in email else name).strip().lower()
    base = re.sub(r"[^a-z0-9_]+", "_", base).strip("_") or "student"
    if not column_exists("users", "username"):
        return base
    candidate = base
    counter = 1
    while query_db("SELECT id FROM users WHERE username=%s", (candidate,), fetchone=True):
        counter += 1
        candidate = f"{base}{counter}"
    return candidate



def generate_token(user):
    uid = row_value(user, "id", "user_id")
    payload = {
        "id": uid,
        "user_id": uid,
        "name": user.get("name"),
        "email": user.get("email"),
        "role": user.get("role", "student"),
        "current_mode": user.get("current_mode", user.get("role", "student")),
        "exp": datetime.datetime.utcnow() + datetime.timedelta(hours=int(os.getenv("JWT_HOURS", "24"))),
    }
    return jwt.encode(payload, app.config["SECRET_KEY"], algorithm="HS256")



def get_current_user():
    token = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
    if not token:
        return None
    try:
        data = jwt.decode(token, app.config["SECRET_KEY"], algorithms=["HS256"])
        uid = data.get("user_id") or data.get("id")
        user = query_db("SELECT * FROM users WHERE id=%s", (uid,), fetchone=True)
        if not user:
            return None
        if safe_int(row_value(user, "banned", "is_banned"), 0) == 1:
            return None
        return user
    except Exception:
        return None


def get_logged_user_id():
    """Compatibility helper used by older enrollment routes."""
    user = getattr(request, "current_user", None) or get_current_user()
    return row_value(user or {}, "id", "user_id")



def login_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def admin_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        if safe_text(user.get("role")).lower() != "admin":
            return jsonify({"error": "Admin only"}), 403
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def student_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        if safe_text(user.get("role")).lower() == "admin" and safe_text(user.get("current_mode") or "admin").lower() != "student":
            return jsonify({"error": "Admins can only access the admin page unless switched to student mode."}), 403
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def allowed_file(filename):
    allowed = {"png", "jpg", "jpeg", "gif", "webp", "mp4", "mov", "webm", "ogg", "pdf", "docx", "txt"}
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed


def save_file(field_name):
    file = request.files.get(field_name)
    if not file or not file.filename:
        return ""
    if not allowed_file(file.filename):
        return ""
    original = secure_filename(file.filename)
    stamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
    filename = f"{stamp}_{original}"
    file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
    return filename


def request_data():
    """Read JSON, multipart FormData, and normal HTML forms safely."""
    if request.is_json:
        return request.get_json(silent=True) or {}
    data = {}
    try:
        data.update(request.form.to_dict(flat=True))
    except Exception:
        pass
    if data:
        return data
    return request.get_json(silent=True) or {}


def calculate_match_percentage(profile_text, target_text):
    profile = safe_text(profile_text).lower()
    target = safe_text(target_text).lower()
    words = []
    for skill in TECH_SKILLS:
        if skill in profile or skill in target:
            words.append(skill)
    matched = [skill for skill in words if skill in profile and skill in target]
    unique_target = sorted(set([skill for skill in words if skill in target]))
    if not unique_target:
        tokens = set(re.findall(r"[a-zA-Z][a-zA-Z+#.]{2,}", target))
        user_tokens = set(re.findall(r"[a-zA-Z][a-zA-Z+#.]{2,}", profile))
        if not tokens:
            return 0, []
        score = round((len(tokens & user_tokens) / max(len(tokens), 1)) * 100)
        return max(0, min(100, score)), sorted(tokens & user_tokens)[:12]
    score = round((len(set(matched)) / max(len(unique_target), 1)) * 100)
    return max(0, min(100, score)), sorted(set(matched))[:12]


def extract_resume_text(file):
    if not file or not file.filename:
        return ""
    name = file.filename.lower()
    raw = file.read()
    file.seek(0)
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf") and PdfReader:
        reader = PdfReader(BytesIO(raw))
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    if name.endswith(".docx") and Document:
        doc = Document(BytesIO(raw))
        return "\n".join([p.text for p in doc.paragraphs])
    return ""


def extract_json_object(text_value, fallback=None):
    fallback = fallback if fallback is not None else None
    value = safe_text(text_value)
    if not value:
        return fallback

    value = re.sub(r"^```(?:json)?\s*", "", value, flags=re.I).strip()
    value = re.sub(r"\s*```$", "", value).strip()

    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, dict) else fallback
    except Exception:
        pass

    match = re.search(r"\{.*\}", value, re.S)
    if not match:
        return fallback

    try:
        parsed = json.loads(match.group(0))
        return parsed if isinstance(parsed, dict) else fallback
    except Exception:
        return fallback





def _resume_line_kind(line):
    value = safe_text(line)
    upper = value.upper().strip()
    heading_words = {
        "OBJECTIVE", "PROFESSIONAL SUMMARY", "SUMMARY", "ENHANCED SUMMARY",
        "TECHNICAL SKILLS", "SOFT SKILLS", "SKILLS", "LINKEDIN", "GITHUB", "LINKS",
        "PROJECTS", "EXPERIENCE", "WORK EXPERIENCE", "INTERNSHIP", "INTERNSHIPS",
        "EDUCATION", "CERTIFICATION", "CERTIFICATIONS", "CERTIFICATES", "CONTACT"
    }
    if upper in heading_words:
        return "heading"
    if value.startswith(("- ", "• ", "* ", "•")):
        return "bullet"
    return "normal"


def _resume_export_heading(line):
    """Normalize common resume headings while keeping the user's sections ATS-readable."""
    upper = safe_text(line).upper().strip().rstrip(":")
    heading_map = {
        "SUMMARY": "OBJECTIVE",
        "PROFESSIONAL SUMMARY": "OBJECTIVE",
        "ENHANCED SUMMARY": "OBJECTIVE",
        "CERTIFICATIONS": "CERTIFICATION",
        "CERTIFICATES": "CERTIFICATION",
        "WORK EXPERIENCE": "INTERNSHIP",
        "INTERNSHIPS": "INTERNSHIP",
        "EXPERIENCE": "INTERNSHIP",
    }
    return heading_map.get(upper, upper)


def _looks_like_contact_line(line):
    value = safe_text(line)
    lower = value.lower()
    if "|" in value and any(token in lower for token in ["@", "phone", "linkedin", "github", "portfolio", "http", "+"]):
        return True
    if re.search(r"\b(phone|email|location|address|linkedin|github|portfolio)\b", lower):
        return True
    return False


def _compact_resume_export_lines(text):
    lines = [safe_text(x).strip() for x in safe_text(text).replace("\r\n", "\n").split("\n")]
    compact = []
    last_blank = False
    for line in lines:
        if not line:
            if compact and not last_blank:
                compact.append("")
            last_blank = True
            continue
        compact.append(line)
        last_blank = False
    while compact and not compact[0]:
        compact.pop(0)
    while compact and not compact[-1]:
        compact.pop()
    return compact


def _split_resume_header(lines):
    name = lines[0] if lines else "FULL NAME"
    role = ""
    contact = ""
    cursor = 1
    while cursor < len(lines) and not lines[cursor]:
        cursor += 1
    if cursor < len(lines) and _resume_line_kind(lines[cursor]) != "heading" and not _looks_like_contact_line(lines[cursor]):
        role = lines[cursor]
        cursor += 1
    while cursor < len(lines) and not lines[cursor]:
        cursor += 1
    if cursor < len(lines) and _looks_like_contact_line(lines[cursor]):
        contact = lines[cursor]
        cursor += 1
    while cursor < len(lines) and not lines[cursor]:
        cursor += 1
    return name, role, contact, lines[cursor:]


def _build_contact_line_from_data(data):
    data = data or {}
    pieces = [
        safe_text(data.get("phone") or data.get("phone_number") or data.get("mobile")),
        safe_text(data.get("location") or data.get("address") or data.get("city")),
        safe_text(data.get("email") or data.get("email_address")),
        sqr_clean_link(data.get("linkedin") or data.get("linkedin_url")),
        sqr_clean_link(data.get("portfolio") or data.get("portfolio_url") or data.get("github") or data.get("github_url")),
    ]
    pieces = [p for p in pieces if p]
    return " | ".join(pieces)


def build_resume_pdf(text):
    """Build a one-page style resume PDF matching the clean SQR ATS template."""
    if not SimpleDocTemplate or not Paragraph or not Spacer or not A4 or not getSampleStyleSheet:
        return None

    def esc(value):
        return (safe_text(value)
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=58,
        leftMargin=58,
        topMargin=56,
        bottomMargin=46,
        title="SQR Resume",
    )
    styles = getSampleStyleSheet()
    dark = colors.HexColor("#27303d") if colors else None
    line_color = colors.HexColor("#6b7280") if colors else None

    normal = ParagraphStyle(
        "SQRResumeNormal",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.3,
        leading=13.8,
        textColor=dark,
        spaceAfter=3,
    )
    name_style = ParagraphStyle(
        "SQRResumeName",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        textColor=dark,
        spaceAfter=5,
    )
    role_style = ParagraphStyle(
        "SQRResumeRole",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=11.2,
        leading=14,
        alignment=TA_CENTER,
        textColor=dark,
        spaceAfter=20,
    )
    contact_style = ParagraphStyle(
        "SQRResumeContact",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.6,
        leading=13,
        alignment=TA_CENTER,
        textColor=dark,
        spaceBefore=6,
        spaceAfter=6,
    )
    heading = ParagraphStyle(
        "SQRResumeHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.2,
        leading=13.5,
        textColor=dark,
        spaceBefore=12,
        spaceAfter=3,
    )
    bullet = ParagraphStyle(
        "SQRResumeBullet",
        parent=normal,
        leftIndent=9,
        firstLineIndent=-7,
        spaceAfter=2,
    )

    lines = _compact_resume_export_lines(text)
    if not lines:
        return None

    name, role, contact, body_lines = _split_resume_header(lines)
    story = [Paragraph(esc(name.upper()), name_style)]
    if role:
        story.append(Paragraph(esc(role.upper()), role_style))
    if contact:
        if HRFlowable:
            story.append(HRFlowable(width="100%", thickness=0.7, color=line_color, spaceBefore=0, spaceAfter=6))
        story.append(Paragraph(esc(contact), contact_style))
        if HRFlowable:
            story.append(HRFlowable(width="100%", thickness=0.7, color=line_color, spaceBefore=0, spaceAfter=11))

    for raw in body_lines:
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 3))
            continue
        kind = _resume_line_kind(line)
        if kind == "heading":
            story.append(Paragraph(esc(_resume_export_heading(line)), heading))
            if HRFlowable:
                story.append(HRFlowable(width="100%", thickness=0.65, color=line_color, spaceBefore=0, spaceAfter=6))
        elif kind == "bullet":
            cleaned = re.sub(r"^[•\-*]\s*", "", line)
            story.append(Paragraph("• " + esc(cleaned), bullet))
        else:
            story.append(Paragraph(esc(line), normal))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _docx_paragraph_border(paragraph, top=False, bottom=False, color="6b7280", size="8"):
    if not OxmlElement or not qn:
        return
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for edge, enabled in (("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        tag = "w:" + edge
        element = pBdr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            pBdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), color)


def _set_docx_run(run, size=10.5, bold=False, color="27303d", all_caps=False):
    run.bold = bool(bold)
    if Pt:
        run.font.size = Pt(size)
    if RGBColor:
        run.font.color.rgb = RGBColor.from_string(color)
    if all_caps:
        run.text = safe_text(run.text).upper()


def build_resume_docx(text):
    """Build a DOCX resume using the same clean visual system as the PDF export."""
    if not Document:
        return None

    doc = Document()
    section = doc.sections[0]
    if Inches:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    if Pt:
        styles["Normal"].font.size = Pt(10.5)
    if RGBColor:
        styles["Normal"].font.color.rgb = RGBColor.from_string("27303d")

    lines = _compact_resume_export_lines(text)
    if not lines:
        return None
    name, role, contact, body_lines = _split_resume_header(lines)

    p = doc.add_paragraph()
    if WD_ALIGN_PARAGRAPH:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2) if Pt else None
    run = p.add_run(name.upper())
    _set_docx_run(run, size=24, bold=True, all_caps=True)

    if role:
        p = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14) if Pt else None
        run = p.add_run(role.upper())
        _set_docx_run(run, size=11.5, bold=True, all_caps=True)

    if contact:
        p = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_paragraph_border(p, top=True, bottom=True)
        p.paragraph_format.space_before = Pt(1) if Pt else None
        p.paragraph_format.space_after = Pt(14) if Pt else None
        run = p.add_run(contact)
        _set_docx_run(run, size=9.5, bold=False)

    for raw in body_lines:
        line = raw.strip()
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1) if Pt else None
            continue
        kind = _resume_line_kind(line)
        if kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7) if Pt else None
            p.paragraph_format.space_after = Pt(2) if Pt else None
            run = p.add_run(_resume_export_heading(line))
            _set_docx_run(run, size=11.2, bold=True, all_caps=True)
            _docx_paragraph_border(p, bottom=True)
        elif kind == "bullet":
            cleaned = re.sub(r"^[•\-*]\s*", "", line)
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.12) if Inches else None
            p.paragraph_format.first_line_indent = Inches(-0.1) if Inches else None
            p.paragraph_format.space_after = Pt(1) if Pt else None
            run = p.add_run("• " + cleaned)
            _set_docx_run(run, size=10.2)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2) if Pt else None
            run = p.add_run(line)
            _set_docx_run(run, size=10.3)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gemini_json(prompt, fallback=None):
    return ai_json(prompt, fallback)


def mask_gemini_error(value):
    text_value = safe_text(value)
    for secret in [GEMINI_API_KEY, XAI_API_KEY, OPENAI_API_KEY]:
        if secret:
            text_value = text_value.replace(secret, "****")
    return text_value


def extract_gemini_text(response_json):
    candidates = response_json.get("candidates") or []
    for candidate in candidates:
        content = candidate.get("content") or {}
        parts = content.get("parts") or []
        text_parts = []
        for part in parts:
            if isinstance(part, dict) and safe_text(part.get("text")):
                text_parts.append(safe_text(part.get("text")))
        if text_parts:
            return "\n".join(text_parts)
    return ""


def _ai_fallback(fallback, provider_name, error_message):
    fallback = dict(fallback or {})
    fallback["ai_powered"] = False
    fallback["ai_provider"] = "local_dynamic_fallback"
    fallback["ai_failed_provider"] = provider_name
    fallback["ai_error"] = mask_gemini_error(error_message or "AI provider returned no usable JSON.")
    return fallback


def _finalize_ai_json(raw_text, fallback, provider_name, model_name):
    parsed = extract_json_object(raw_text, None)
    if not isinstance(parsed, dict):
        raise RuntimeError(f"Could not parse {provider_name} JSON: {safe_text(raw_text)[:250]}")
    parsed["ai_powered"] = True
    parsed["ai_provider"] = provider_name
    parsed["ai_model"] = model_name
    return parsed


def _gemini_json_once(prompt, fallback=None):
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not configured. Add it in Render Environment variables.")

    system_instruction = (
        "You are the SQR AI engine. Return valid JSON only. "
        "Do not use markdown. Do not invent facts that the user did not provide. "
        "Separate technical skills from soft skills."
    )
    endpoint = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        + urllib.parse.quote(GEMINI_MODEL, safe="")
        + ":generateContent?key="
        + urllib.parse.quote(GEMINI_API_KEY, safe="")
    )
    request_payload = {
        "systemInstruction": {"parts": [{"text": system_instruction}]},
        "contents": [{"role": "user", "parts": [{"text": safe_text(prompt)}]}],
        "generationConfig": {
            "temperature": 0.25,
            "maxOutputTokens": 3200,
            "responseMimeType": "application/json"
        }
    }

    req = urllib.request.Request(
        endpoint,
        data=json.dumps(request_payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=AI_TIMEOUT) as response:
        response_text = response.read().decode("utf-8", errors="ignore")

    response_json = json.loads(response_text or "{}")
    if response_json.get("error"):
        raise RuntimeError(json.dumps(response_json.get("error"), ensure_ascii=False))

    return _finalize_ai_json(extract_gemini_text(response_json), fallback, "gemini", GEMINI_MODEL)


def _openai_compatible_json_once(prompt, fallback=None, provider_name="openai", api_key="", base_url=None, model_name=""):
    if not api_key:
        raise RuntimeError(f"{provider_name.upper()} API key is not configured in Render Environment variables.")
    if OpenAI is None:
        raise RuntimeError("The openai package is not installed. Add openai to requirements.txt and redeploy.")

    system_instruction = (
        "You are the SQR AI engine. Return valid JSON only. "
        "Do not use markdown. Do not invent facts that the user did not provide. "
        "Separate technical skills from soft skills."
    )
    client_kwargs = {"api_key": api_key}
    if base_url:
        client_kwargs["base_url"] = base_url
    client = OpenAI(**client_kwargs)

    kwargs = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": safe_text(prompt)},
        ],
        "temperature": 0.25,
        "max_tokens": 3200,
        "response_format": {"type": "json_object"},
    }
    try:
        response = client.chat.completions.create(**kwargs)
    except Exception:
        kwargs.pop("response_format", None)
        response = client.chat.completions.create(**kwargs)

    raw_text = response.choices[0].message.content if response and response.choices else ""
    return _finalize_ai_json(raw_text, fallback, provider_name, model_name)


def _call_provider_with_retries(provider_name, prompt, fallback):
    last_error = ""
    for attempt in range(AI_MAX_RETRIES):
        try:
            if provider_name == "gemini":
                return _gemini_json_once(prompt, fallback)
            if provider_name in {"xai", "grok"}:
                return _openai_compatible_json_once(
                    prompt, fallback, provider_name="xai", api_key=XAI_API_KEY,
                    base_url="https://api.x.ai/v1", model_name=XAI_MODEL
                )
            if provider_name == "openai":
                return _openai_compatible_json_once(
                    prompt, fallback, provider_name="openai", api_key=OPENAI_API_KEY,
                    base_url=None, model_name=OPENAI_MODEL
                )
            raise RuntimeError(f"Unsupported AI_PROVIDER: {provider_name}. Use gemini, xai, openai, or auto.")
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="ignore") if hasattr(exc, "read") else ""
            last_error = mask_gemini_error(body or str(exc))
            temporary = exc.code in [429, 500, 502, 503, 504]
        except Exception as exc:
            last_error = mask_gemini_error(str(exc))
            temporary = any(word in last_error.lower() for word in ["429", "500", "502", "503", "504", "rate", "timeout"])

        if temporary and attempt < AI_MAX_RETRIES - 1:
            time.sleep(min(2 + attempt, 5))
            continue
        break

    print(f"{provider_name.upper()} AI ERROR:", last_error)
    return _ai_fallback(fallback, provider_name, last_error)


def ai_json(prompt, fallback=None):
    fallback = dict(fallback or {})
    provider = safe_text(AI_PROVIDER).lower().replace(" ", "_").replace("-", "_") or "gemini"

    if provider in {"xai", "grok", "grok_ai"}:
        return _call_provider_with_retries("xai", prompt, fallback)

    if provider == "openai":
        return _call_provider_with_retries("openai", prompt, fallback)

    if provider == "auto":
        # Try Gemini first, then xAI/OpenAI only if you configured those keys.
        errors = []
        for candidate in ["gemini", "xai", "openai"]:
            if candidate == "gemini" and not GEMINI_API_KEY:
                continue
            if candidate == "xai" and not XAI_API_KEY:
                continue
            if candidate == "openai" and not OPENAI_API_KEY:
                continue
            result = _call_provider_with_retries(candidate, prompt, fallback)
            if result.get("ai_powered"):
                return result
            errors.append(f"{candidate}: {result.get('ai_error', '')}")
        return _ai_fallback(fallback, "auto", " | ".join(errors) or "No AI provider key is configured.")

    # Default: Gemini. If Gemini says location is unsupported, it is a Google/hosting-region block, not a bad key.
    result = _call_provider_with_retries("gemini", prompt, fallback)
    if not result.get("ai_powered") and "location is not supported" in safe_text(result.get("ai_error")).lower():
        result["ai_error"] = (
            "Google Gemini blocked the Render server location/IP. "
            "The API key is being read, but this host is not allowed by Google AI Studio. "
            "Use AI_PROVIDER=xai/openai with a matching key, or redeploy the backend in a Gemini-supported host/region."
        )
    return result


def init_db():
    """Create/patch the database without deleting project features.
    This version matches the Railway dump tables: specialization_id, course_id,
    quiz_id, job_id, specialization_enrollments, course_enrollments, and ats_id.
    """
    # Compatibility for old local databases that used the singular table name.
    # The Railway dump and this backend use `specializations` plural.
    try:
        if table_exists("specialization") and not table_exists("specializations"):
            exec_db("RENAME TABLE specialization TO specializations")
            print("Renamed specialization table to specializations")
    except Exception as exc:
        print("specialization table compatibility rename skipped:", exc)

    statements = [
        """
        CREATE TABLE IF NOT EXISTS users (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            email VARCHAR(150) NOT NULL UNIQUE,
            password VARCHAR(255) NOT NULL,
            role ENUM('student','admin') NOT NULL DEFAULT 'student',
            is_banned TINYINT(1) DEFAULT 0,
            current_mode ENUM('student','admin') DEFAULT 'student',
            banned TINYINT DEFAULT 0,
            skills TEXT,
            interests TEXT,
            goal TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS admins (
            admin_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL UNIQUE,
            admin_level ENUM('owner','manager') DEFAULT 'manager',
            can_manage_users TINYINT(1) DEFAULT 1,
            can_manage_specializations TINYINT(1) DEFAULT 1,
            can_manage_courses TINYINT(1) DEFAULT 1,
            can_manage_quizzes TINYINT(1) DEFAULT 1,
            can_view_reports TINYINT(1) DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS specializations (
            specialization_id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(150) NOT NULL,
            description TEXT,
            roadmap TEXT,
            job_titles TEXT,
            career_paths TEXT,
            image_url VARCHAR(255),
            skills TEXT,
            image VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS courses (
            course_id INT AUTO_INCREMENT PRIMARY KEY,
            specialization_id INT NOT NULL,
            title VARCHAR(150) NOT NULL,
            description TEXT,
            level ENUM('Beginner','Intermediate','Advanced') DEFAULT 'Beginner',
            course_link VARCHAR(255),
            video_url VARCHAR(255),
            image_url VARCHAR(255),
            spec_id INT,
            link VARCHAR(255),
            image VARCHAR(255),
            video VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS certificates (
            id INT AUTO_INCREMENT PRIMARY KEY,
            spec_id INT NOT NULL,
            name VARCHAR(150) NOT NULL,
            description TEXT,
            link VARCHAR(255),
            price VARCHAR(100),
            type ENUM('practical','theoretical','both') DEFAULT 'both',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS certifications (
            certification_id INT AUTO_INCREMENT PRIMARY KEY,
            specialization_id INT NOT NULL,
            name VARCHAR(150) NOT NULL,
            description TEXT,
            official_link VARCHAR(255),
            price VARCHAR(50),
            type ENUM('Practical','Theoretical','Both') DEFAULT 'Both',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS jobs (
            job_id INT AUTO_INCREMENT PRIMARY KEY,
            specialization_id INT NOT NULL,
            title VARCHAR(150) NOT NULL,
            description TEXT,
            required_skills TEXT,
            average_salary VARCHAR(100),
            job_link VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quizzes (
            quiz_id INT AUTO_INCREMENT PRIMARY KEY,
            course_id INT NOT NULL,
            title VARCHAR(150) NOT NULL,
            description TEXT,
            total_questions INT DEFAULT 0,
            spec_id INT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quiz_questions (
            question_id INT AUTO_INCREMENT PRIMARY KEY,
            quiz_id INT NOT NULL,
            question_text TEXT NOT NULL,
            option_a VARCHAR(255),
            option_b VARCHAR(255),
            option_c VARCHAR(255),
            option_d VARCHAR(255),
            correct_answer ENUM('A','B','C','D') NOT NULL,
            score DECIMAL(5,2) DEFAULT 1.00,
            question TEXT,
            option1 VARCHAR(255),
            option2 VARCHAR(255),
            option3 VARCHAR(255),
            option4 VARCHAR(255),
            answer VARCHAR(255)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS specialization_enrollments (
            enrollment_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            specialization_id INT NOT NULL,
            progress_percentage DECIMAL(5,2) DEFAULT 0.00,
            status ENUM('Not Started','In Progress','Completed') DEFAULT 'Not Started',
            enrolled_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            completed_at TIMESTAMP NULL DEFAULT NULL,
            UNIQUE KEY unique_user_specialization (user_id, specialization_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS course_enrollments (
            enrollment_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            course_id INT NOT NULL,
            progress_percentage DECIMAL(5,2) DEFAULT 0.00,
            status ENUM('Not Started','In Progress','Completed') DEFAULT 'Not Started',
            enrolled_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            completed_at TIMESTAMP NULL DEFAULT NULL,
            UNIQUE KEY unique_user_course (user_id, course_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quiz_attempts (
            attempt_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            quiz_id INT NOT NULL,
            score DECIMAL(5,2) DEFAULT 0.00,
            passed TINYINT(1) DEFAULT 0,
            attempted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS ats_results (
            ats_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            resume_text LONGTEXT,
            target_job VARCHAR(150),
            ats_score DECIMAL(5,2) DEFAULT 0.00,
            missing_keywords TEXT,
            matched_keywords TEXT,
            suggestions TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS recommendations (
            recommendation_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            specialization_id INT NOT NULL,
            assessment_id INT DEFAULT NULL,
            match_score DECIMAL(5,2) DEFAULT 0.00,
            explanation TEXT,
            generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS recommendation_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            recommendation_json LONGTEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS user_completed_courses (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            course_id INT NOT NULL,
            completed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY unique_completed_course (user_id, course_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS user_completed_quizzes (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            quiz_id INT NOT NULL,
            score INT DEFAULT 0,
            completed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY unique_completed_quiz (user_id, quiz_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS progress (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            spec_id INT NOT NULL,
            progress INT DEFAULT 0,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS specialization_progress (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            specialization_id INT NOT NULL,
            progress INT DEFAULT 0,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_spec_progress (user_id, specialization_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS students (
            student_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL UNIQUE,
            university VARCHAR(150),
            major VARCHAR(100),
            gpa DECIMAL(3,2),
            skills TEXT,
            interests TEXT,
            graduation_year INT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS assessments (
            assessment_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            title VARCHAR(150) DEFAULT 'Career Assessment',
            description TEXT,
            total_score DECIMAL(6,2) DEFAULT 0.00,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS assessment_answers (
            answer_id INT AUTO_INCREMENT PRIMARY KEY,
            assessment_id INT NOT NULL,
            question_text TEXT NOT NULL,
            selected_option VARCHAR(255),
            score DECIMAL(6,2) DEFAULT 0.00
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS cvs (
            cv_id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            file_url VARCHAR(255),
            generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
    ]
    for statement in statements:
        try:
            exec_db(statement)
        except Exception as exc:
            print("init_db statement skipped:", exc)

    try:
        if table_exists("specialization") and table_exists("specializations"):
            source_count = query_db("SELECT COUNT(*) AS total FROM specialization", fetchone=True) or {"total": 0}
            target_count = query_db("SELECT COUNT(*) AS total FROM specializations", fetchone=True) or {"total": 0}
            if safe_int(source_count.get("total"), 0) > 0 and safe_int(target_count.get("total"), 0) == 0:
                exec_db("""
                    INSERT INTO specializations (specialization_id, name, description, roadmap, job_titles, career_paths, image_url, created_at, id, skills, image)
                    SELECT specialization_id, name, description, roadmap, job_titles, career_paths, image_url, created_at, id, skills, image
                    FROM specialization
                """)
                print("Copied old specialization data into specializations")
    except Exception as exc:
        print("specialization table compatibility copy skipped:", exc)

    compatibility_columns = {
        "users": [
            ("current_mode", "ENUM('student','admin') DEFAULT 'student'"),
            ("banned", "TINYINT DEFAULT 0"),
            ("skills", "TEXT"),
            ("interests", "TEXT"),
            ("goal", "TEXT"),
        ],
        "specializations": [
            ("id", "INT"),
            ("spec_id", "INT"),
            ("roadmap", "TEXT"),
            ("job_titles", "TEXT"),
            ("career_paths", "TEXT"),
            ("image_url", "VARCHAR(255)"),
            ("skills", "TEXT"),
            ("image", "VARCHAR(255)"),
        ],
        "courses": [
            ("id", "INT"),
            ("spec_id", "INT"),
            ("name", "VARCHAR(150)"),
            ("difficulty", "VARCHAR(50)"),
            ("link", "VARCHAR(255)"),
            ("image", "VARCHAR(255)"),
            ("video", "VARCHAR(255)"),
            ("course_link", "VARCHAR(255)"),
            ("video_url", "VARCHAR(255)"),
            ("image_url", "VARCHAR(255)"),
        ],
        "jobs": [
            ("id", "INT"),
            ("skills", "TEXT"),
            ("salary", "VARCHAR(100)"),
            ("link", "VARCHAR(255)"),
        ],
        "quizzes": [
            ("id", "INT"),
            ("name", "VARCHAR(150)"),
            ("spec_id", "INT"),
            ("total_questions", "INT DEFAULT 0"),
            ("description", "TEXT"),
        ],
        "quiz_attempts": [
            ("id", "INT"),
            ("course_id", "INT"),
            ("total", "INT DEFAULT 0"),
            ("percentage", "DECIMAL(5,2) DEFAULT 0.00"),
            ("answers_json", "TEXT"),
        ],
        "quiz_questions": [
            ("id", "INT"),
            ("question", "TEXT"),
            ("option1", "VARCHAR(255)"),
            ("option2", "VARCHAR(255)"),
            ("option3", "VARCHAR(255)"),
            ("option4", "VARCHAR(255)"),
            ("answer", "VARCHAR(255)"),
        ],
        "ats_results": [
            ("id", "INT"),
            ("score", "DECIMAL(5,2) DEFAULT 0.00"),
            ("summary", "TEXT"),
        ],
        "specialization_enrollments": [
            ("id", "INT"),
            ("spec_id", "INT"),
            ("progress", "INT DEFAULT 0"),
            ("specialization_id", "INT"),
            ("progress_percentage", "DECIMAL(5,2) DEFAULT 0.00"),
        ],
        "course_enrollments": [
            ("id", "INT"),
            ("progress", "INT DEFAULT 0"),
            ("progress_percentage", "DECIMAL(5,2) DEFAULT 0.00"),
        ],
    }
    for table, columns in compatibility_columns.items():
        if not table_exists(table):
            continue
        for column, definition in columns:
            try:
                if not column_exists(table, column):
                    exec_db(f"ALTER TABLE `{table}` ADD COLUMN `{column}` {definition}")
            except Exception as exc:
                print(f"init_db alter skipped for {table}.{column}:", exc)

    try:
        sync_compatibility_alias_columns()
    except Exception as exc:
        print("compatibility alias sync skipped:", exc)


def sync_compatibility_alias_columns():
    """Keep old frontend/backend aliases like id/spec_id/progress synced with Railway-style columns."""
    def has(table, column):
        return table_exists(table) and column_exists(table, column)

    updates = []

    if has("specializations", "id") and has("specializations", "specialization_id"):
        updates.append("UPDATE specializations SET id=specialization_id WHERE id IS NULL OR id=0")
    if has("specializations", "spec_id") and has("specializations", "specialization_id"):
        updates.append("UPDATE specializations SET spec_id=specialization_id WHERE spec_id IS NULL OR spec_id=0")

    if has("courses", "id") and has("courses", "course_id"):
        updates.append("UPDATE courses SET id=course_id WHERE id IS NULL OR id=0")
    if has("courses", "spec_id") and has("courses", "specialization_id"):
        updates.append("UPDATE courses SET spec_id=specialization_id WHERE spec_id IS NULL OR spec_id=0")
    if has("courses", "name") and has("courses", "title"):
        updates.append("UPDATE courses SET name=title WHERE name IS NULL OR name=''")
    if has("courses", "difficulty") and has("courses", "level"):
        updates.append("UPDATE courses SET difficulty=level WHERE difficulty IS NULL OR difficulty=''")
    if has("courses", "link") and has("courses", "course_link"):
        updates.append("UPDATE courses SET link=course_link WHERE link IS NULL OR link=''")
    if has("courses", "image") and has("courses", "image_url"):
        updates.append("UPDATE courses SET image=image_url WHERE image IS NULL OR image=''")
    if has("courses", "video") and has("courses", "video_url"):
        updates.append("UPDATE courses SET video=video_url WHERE video IS NULL OR video=''")

    if has("quizzes", "id") and has("quizzes", "quiz_id"):
        updates.append("UPDATE quizzes SET id=quiz_id WHERE id IS NULL OR id=0")
    if has("quizzes", "name") and has("quizzes", "title"):
        updates.append("UPDATE quizzes SET name=title WHERE name IS NULL OR name=''")

    if has("quiz_questions", "id") and has("quiz_questions", "question_id"):
        updates.append("UPDATE quiz_questions SET id=question_id WHERE id IS NULL OR id=0")
    if has("quiz_questions", "question") and has("quiz_questions", "question_text"):
        updates.append("UPDATE quiz_questions SET question=question_text WHERE question IS NULL OR question=''")
    if has("quiz_questions", "option1") and has("quiz_questions", "option_a"):
        updates.append("UPDATE quiz_questions SET option1=option_a WHERE option1 IS NULL OR option1=''")
    if has("quiz_questions", "option2") and has("quiz_questions", "option_b"):
        updates.append("UPDATE quiz_questions SET option2=option_b WHERE option2 IS NULL OR option2=''")
    if has("quiz_questions", "option3") and has("quiz_questions", "option_c"):
        updates.append("UPDATE quiz_questions SET option3=option_c WHERE option3 IS NULL OR option3=''")
    if has("quiz_questions", "option4") and has("quiz_questions", "option_d"):
        updates.append("UPDATE quiz_questions SET option4=option_d WHERE option4 IS NULL OR option4=''")
    if has("quiz_questions", "answer") and has("quiz_questions", "correct_answer"):
        updates.append("UPDATE quiz_questions SET answer=correct_answer WHERE answer IS NULL OR answer=''")

    if has("jobs", "id") and has("jobs", "job_id"):
        updates.append("UPDATE jobs SET id=job_id WHERE id IS NULL OR id=0")
    if has("jobs", "skills") and has("jobs", "required_skills"):
        updates.append("UPDATE jobs SET skills=required_skills WHERE skills IS NULL OR skills=''")
    if has("jobs", "salary") and has("jobs", "average_salary"):
        updates.append("UPDATE jobs SET salary=average_salary WHERE salary IS NULL OR salary=''")
    if has("jobs", "link") and has("jobs", "job_link"):
        updates.append("UPDATE jobs SET link=job_link WHERE link IS NULL OR link=''")

    if has("specialization_enrollments", "id") and has("specialization_enrollments", "enrollment_id"):
        updates.append("UPDATE specialization_enrollments SET id=enrollment_id WHERE id IS NULL OR id=0")
    if has("specialization_enrollments", "spec_id") and has("specialization_enrollments", "specialization_id"):
        updates.append("UPDATE specialization_enrollments SET spec_id=specialization_id WHERE spec_id IS NULL OR spec_id=0")
    if has("specialization_enrollments", "progress") and has("specialization_enrollments", "progress_percentage"):
        updates.append("UPDATE specialization_enrollments SET progress=progress_percentage WHERE progress IS NULL OR progress=0")
    if has("course_enrollments", "id") and has("course_enrollments", "enrollment_id"):
        updates.append("UPDATE course_enrollments SET id=enrollment_id WHERE id IS NULL OR id=0")
    if has("course_enrollments", "progress") and has("course_enrollments", "progress_percentage"):
        updates.append("UPDATE course_enrollments SET progress=progress_percentage WHERE progress IS NULL OR progress=0")

    if has("quiz_attempts", "id") and has("quiz_attempts", "attempt_id"):
        updates.append("UPDATE quiz_attempts SET id=attempt_id WHERE id IS NULL OR id=0")
    if has("quiz_attempts", "percentage") and has("quiz_attempts", "score"):
        updates.append("UPDATE quiz_attempts SET percentage=score WHERE percentage IS NULL OR percentage=0")

    if has("ats_results", "id") and has("ats_results", "ats_id"):
        updates.append("UPDATE ats_results SET id=ats_id WHERE id IS NULL OR id=0")
    if has("ats_results", "score") and has("ats_results", "ats_score"):
        updates.append("UPDATE ats_results SET score=ats_score WHERE score IS NULL OR score=0")
    if has("ats_results", "summary") and has("ats_results", "suggestions"):
        updates.append("UPDATE ats_results SET summary=suggestions WHERE summary IS NULL OR summary=''")

    for sql in updates:
        try:
            exec_db(sql)
        except Exception as exc:
            print("alias sync skipped:", exc)



def render_page(template_name):
    """Render templates safely even if links use Courses.html while the file is courses.html."""
    templates_dir = os.path.join(app.root_path, "templates")
    requested = safe_text(template_name)
    direct_path = os.path.join(templates_dir, requested)
    if os.path.exists(direct_path):
        return render_template(requested)

    lower_requested = requested.lower()
    try:
        for filename in os.listdir(templates_dir):
            if filename.lower() == lower_requested:
                return render_template(filename)
    except Exception:
        pass

    return render_template(requested)


@app.route("/")
def home():
    return render_page("gp.html")


@app.route("/home")
def page_home():
    return render_page("gp.html")


@app.route("/specializations")
def page_specializations():
    return render_page("Specialization.html")


@app.route("/courses")
def page_courses():
    return render_page("courses.html")


@app.route("/quizzes")
def page_quizzes():
    return render_page("Quiz.html")


@app.route("/ats")
def page_ats():
    return render_page("ATS.html")


@app.route("/jobs")
def page_jobs():
    return render_page("jobs.html")


@app.route("/recommendation")
def page_recommendation():
    return render_page("recommendation.html")


@app.route("/profile")
def page_profile():
    return render_page("profile.html")


@app.route("/admin")
def page_admin():
    return render_page("admin.html")


@app.route("/signin")
def page_signin():
    return render_page("signin.html")


@app.route("/signup")
def page_signup():
    return render_page("signup.html")


@app.route("/<path:page>.html")
def legacy_html_pages(page):
    aliases = {
        "gp": "gp.html",
        "Specialization": "Specialization.html",
        "Sepecialization": "Specialization.html",
        "Courses": "courses.html",
        "courses": "courses.html",
        "Quiz": "Quiz.html",
        "ATS": "ATS.html",
        "ats": "ATS.html",
        "jobs": "jobs.html",
        "JobDetails": "JobDetails.html",
        "recommendation": "recommendation.html",
        "profile": "profile.html",
        "admin": "admin.html",
        "signin": "signin.html",
        "signup": "signup.html",
    }
    return render_page(aliases.get(page, f"{page}.html"))


@app.route("/uploads/<path:filename>")
def uploads(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


@app.route("/api/health")
def health():
    return jsonify({
        "message": "SQR Backend is running",
        "features": ["auth", "admin", "specializations", "courses", "quizzes", "jobs", "profile", "progress", "ATS", "recommendation"],
        "database": DB_CONFIG.get("database")
    })


@app.route("/api/ai/status")
@login_required
def ai_status():
    return jsonify({
        "ai_provider_mode": AI_PROVIDER,
        "gemini_configured": bool(GEMINI_API_KEY),
        "xai_configured": bool(XAI_API_KEY),
        "openai_configured": bool(OPENAI_API_KEY),
        "gemini_client_ready": bool(gemini_client),
        "gemini_model": GEMINI_MODEL if GEMINI_API_KEY else "",
        "xai_model": XAI_MODEL if XAI_API_KEY else "",
        "openai_model": OPENAI_MODEL if OPENAI_API_KEY else "",
        "max_retries": AI_MAX_RETRIES,
        "timeout_seconds": AI_TIMEOUT,
        "message": "AI provider is configured" if (GEMINI_API_KEY or XAI_API_KEY or OPENAI_API_KEY) else "Set an AI API key in Render Environment variables."
    })


@app.route("/api/signup", methods=["POST"])

def signup():
    data = get_json()
    name = safe_text(data.get("name"))
    email = safe_text(data.get("email")).lower()
    password = safe_text(data.get("password"))
    if not name or not email or not password:
        return jsonify({"error": "Name, email, and password are required"}), 400
    if not strong_password(password):
        return jsonify({"error": "Password must be at least 8 characters and include uppercase, lowercase, number, and symbol"}), 400
    if query_db("SELECT id FROM users WHERE email=%s", (email,), fetchone=True):
        return jsonify({"error": "Email already exists"}), 409
    hashed_password = generate_password_hash(password, method="pbkdf2:sha256", salt_length=16)
    if column_exists("users", "username"):
        username = generate_username(name, email)
        user_id = query_db(
            """
            INSERT INTO users (username, name, email, password, role, current_mode, banned)
            VALUES (%s,%s,%s,%s,'student','student',0)
            """,
            (username, name, email, hashed_password),
            commit=True
        )
    else:
        user_id = query_db(
            """
            INSERT INTO users (name, email, password, role, current_mode, banned)
            VALUES (%s,%s,%s,'student','student',0)
            """,
            (name, email, hashed_password),
            commit=True
        )
    user = query_db("SELECT * FROM users WHERE id=%s", (user_id,), fetchone=True)
    return jsonify({"message": "Account created", "token": generate_token(user), "user": clean_user(user)}), 201



@app.route("/api/signin", methods=["POST"])
def signin():
    data = get_json()
    email = safe_text(data.get("email")).lower()
    password = safe_text(data.get("password"))
    user = query_db("SELECT * FROM users WHERE email=%s", (email,), fetchone=True)
    if not user or not check_password_hash(user.get("password") or "", password):
        return jsonify({"error": "Invalid email or password"}), 401
    if safe_int(user.get("banned"), 0) == 1:
        return jsonify({"error": "Your account is banned"}), 403
    return jsonify({"message": "Login successful", "token": generate_token(user), "user": clean_user(user)})


@app.route("/api/me")
@login_required
def me():
    return jsonify(clean_user(request.current_user))




@app.route("/api/profile", methods=["GET"])
@login_required
def get_profile():
    user = clean_user(request.current_user)
    user_id = user["id"]
    quiz_history = []
    try:
        quiz_history = query_db(
            """
            SELECT
                qa.id,
                qa.id AS attempt_id,
                qa.score,
                qa.passed,
                qa.attempted_at AS created_at,
                COALESCE(q.title, q.name) AS quiz_title,
                COALESCE(c.title, c.name) AS course_title,
                c.id AS course_id,
                COALESCE(qa.percentage, qa.score) AS score_percentage,
                COALESCE(qa.total, (SELECT COUNT(*) FROM quiz_questions qq WHERE qq.quiz_id=qa.quiz_id), 0) AS total
            FROM quiz_attempts qa
            LEFT JOIN quizzes q ON q.id=qa.quiz_id
            LEFT JOIN courses c ON c.id=qa.course_id OR c.id=q.course_id
            WHERE qa.user_id=%s
            ORDER BY qa.attempted_at DESC
            LIMIT 30
            """,
            (user_id,),
            fetchall=True
        ) or []
    except Exception as exc:
        print("PROFILE QUIZ HISTORY ERROR:", exc)

    ats_history = []
    try:
        ats_history = query_db(
            """
            SELECT
                id,
                id AS ats_id,
                target_job,
                COALESCE(ats_score, score, 0) AS score,
                COALESCE(ats_score, score, 0) AS ats_score,
                suggestions AS summary,
                matched_keywords,
                missing_keywords,
                created_at
            FROM ats_results
            WHERE user_id=%s
            ORDER BY created_at DESC
            LIMIT 20
            """,
            (user_id,),
            fetchall=True
        ) or []
    except Exception as exc:
        print("PROFILE ATS HISTORY ERROR:", exc)

    return jsonify({"user": user, "quiz_history": quiz_history, "ats_history": ats_history})


@app.route("/api/profile", methods=["PUT"])
@login_required
def update_profile():
    data = get_json()
    name = safe_text(data.get("name")) or request.current_user.get("name")
    skills = safe_text(data.get("skills"))
    interests = safe_text(data.get("interests"))
    goal = safe_text(data.get("goal"))
    exec_db(
        "UPDATE users SET name=%s, skills=%s, interests=%s, goal=%s WHERE id=%s",
        (name, skills, interests, goal, request.current_user["id"])
    )
    user = query_db("SELECT * FROM users WHERE id=%s", (request.current_user["id"],), fetchone=True)
    return jsonify({"message": "Profile updated", "user": clean_user(user)})




def compute_user_progress(user_id):
    if not table_exists("specialization_enrollments"):
        return []

    specs = query_db(
        """
        SELECT
            s.*,
            se.id AS enrollment_id,
            COALESCE(se.progress, 0) AS enrollment_progress,
            se.status AS enrollment_status,
            se.enrolled_at
        FROM specialization_enrollments se
        JOIN specializations s ON s.id=se.spec_id
        WHERE se.user_id=%s
        ORDER BY se.enrolled_at DESC, s.name
        """,
        (user_id,),
        fetchall=True
    ) or []

    progress_rows = []
    for spec in specs:
        spec_id = safe_int(spec.get("id"), None)
        if not spec_id:
            continue

        total_courses_row = query_db(
            "SELECT COUNT(*) AS total FROM courses WHERE spec_id=%s OR specialization_id=%s",
            (spec_id, spec_id),
            fetchone=True
        ) or {"total": 0}
        total_courses = safe_int(total_courses_row.get("total"), 0)

        enrolled_row = query_db(
            """
            SELECT
                COUNT(DISTINCT ce.course_id) AS enrolled_courses,
                SUM(CASE WHEN COALESCE(ce.progress_percentage, ce.progress, 0) > 0 OR LOWER(ce.status) IN ('in progress','completed','in_progress') THEN 1 ELSE 0 END) AS opened_courses,
                SUM(CASE WHEN COALESCE(ce.progress_percentage, ce.progress, 0) >= 100 OR LOWER(ce.status) IN ('completed') THEN 1 ELSE 0 END) AS completed_courses
            FROM course_enrollments ce
            JOIN courses c ON c.id=ce.course_id
            WHERE ce.user_id=%s AND (c.spec_id=%s OR c.specialization_id=%s)
            """,
            (user_id, spec_id, spec_id),
            fetchone=True
        ) or {"enrolled_courses": 0, "opened_courses": 0, "completed_courses": 0}

        quiz_row = query_db(
            """
            SELECT
                COUNT(DISTINCT qa.quiz_id) AS completed_quizzes,
                COALESCE(ROUND(AVG(COALESCE(qa.percentage, qa.score)),0),0) AS average_score
            FROM quiz_attempts qa
            JOIN quizzes q ON q.id=qa.quiz_id
            JOIN courses c ON c.id=COALESCE(qa.course_id, q.course_id)
            JOIN course_enrollments ce ON ce.course_id=c.id AND ce.user_id=qa.user_id
            WHERE qa.user_id=%s
              AND (c.spec_id=%s OR c.specialization_id=%s)
              AND (qa.passed=1 OR COALESCE(qa.percentage, qa.score) >= 60)
            """,
            (user_id, spec_id, spec_id),
            fetchone=True
        ) or {"completed_quizzes": 0, "average_score": 0}

        total_quizzes_row = query_db(
            """
            SELECT COUNT(DISTINCT q.id) AS total
            FROM quizzes q
            JOIN courses c ON c.id=q.course_id
            WHERE c.spec_id=%s OR c.specialization_id=%s
            """,
            (spec_id, spec_id),
            fetchone=True
        ) or {"total": 0}

        enrolled_courses = safe_int(enrolled_row.get("enrolled_courses"), 0)
        opened_courses = safe_int(enrolled_row.get("opened_courses"), 0)
        completed_courses = safe_int(enrolled_row.get("completed_courses"), 0)
        completed_quizzes = safe_int(quiz_row.get("completed_quizzes"), 0)
        total_quizzes = safe_int(total_quizzes_row.get("total"), 0)
        average_score = safe_int(quiz_row.get("average_score"), 0)

        if total_courses <= 0:
            percent_value = 0
        else:
            opened_part = (opened_courses / total_courses) * 35
            completed_part = (completed_courses / total_courses) * 35
            quiz_part = (completed_quizzes / max(total_quizzes, 1)) * 30 if total_quizzes else 0
            percent_value = max(0, min(100, round(opened_part + completed_part + quiz_part)))

        try:
            if table_exists("specialization_progress"):
                query_db(
                    """
                    INSERT INTO specialization_progress (user_id, specialization_id, progress)
                    VALUES (%s,%s,%s)
                    ON DUPLICATE KEY UPDATE progress=%s
                    """,
                    (user_id, spec_id, percent_value, percent_value),
                    commit=True
                )
        except Exception as exc:
            print("SPECIALIZATION_PROGRESS SAVE ERROR:", exc)

        try:
            status = "Completed" if percent_value >= 100 else ("In Progress" if percent_value > 0 else "Not Started")
            query_db(
                """
                UPDATE specialization_enrollments
                SET progress=%s, status=%s, completed_at=CASE WHEN %s >= 100 THEN CURRENT_TIMESTAMP ELSE completed_at END
                WHERE user_id=%s AND spec_id=%s
                """,
                (percent_value, status, percent_value, user_id, spec_id),
                commit=True
            )
        except Exception as exc:
            print("SPECIALIZATION_ENROLLMENTS PROGRESS ERROR:", exc)

        progress_rows.append({
            "specialization_id": spec_id,
            "id": spec_id,
            "specialization_name": spec.get("name"),
            "name": spec.get("name"),
            "enrolled_courses": enrolled_courses,
            "total_courses": total_courses,
            "opened_courses": opened_courses,
            "completed_courses": completed_courses,
            "total_quizzes": total_quizzes,
            "completed_quizzes": completed_quizzes,
            "average_quiz_score": average_score,
            "progress": percent_value,
            "progress_percentage": percent_value,
            "percentage": percent_value,
            "status": "Completed" if percent_value >= 100 else ("In Progress" if percent_value > 0 else "Not Started"),
        })

    return progress_rows



@app.route("/api/profile/progress")
@login_required
def profile_progress():
    return jsonify({"progress": compute_user_progress(request.current_user["id"])})




@app.route("/api/admin/specializations", methods=["GET"])
@app.route("/api/specialization", methods=["GET"])
@app.route("/api/specializations", methods=["GET"])
def get_specializations():
    try:
        table = admin_spec_table()
        pk = existing_pk(table, ["specialization_id", "id", "spec_id"])
        rows = query_db(f"SELECT * FROM `{table}` ORDER BY `{pk}` DESC", fetchall=True) or []
        rows = [normalize_specialization(row) for row in rows]
        return jsonify({
            "ok": True,
            "specializations": rows,
            "data": rows
        })
    except Exception as e:
        print("GET /api/specializations ERROR:", e)
        return jsonify({"ok": False, "error": "Failed to load specializations", "details": str(e)}), 500



@app.route("/api/specializations/<int:spec_id>", methods=["GET"])
def get_specialization_detail(spec_id):
    try:
        spec = fetch_specialization_by_any_id(spec_id)
        if not spec:
            return jsonify({"ok": False, "error": "Specialization not found"}), 404
        spec = normalize_specialization(spec)
        real_spec_id = safe_int(spec.get("id"), spec_id)

        course_pk = existing_pk("courses", ["id", "course_id"])
        course_spec_cols = existing_columns("courses", ["spec_id", "specialization_id"])
        courses = []
        if course_spec_cols:
            where = " OR ".join([f"`{col}`=%s" for col in course_spec_cols])
            courses = query_db(
                f"SELECT * FROM courses WHERE {where} ORDER BY `{course_pk}` DESC",
                tuple([real_spec_id] * len(course_spec_cols)),
                fetchall=True
            ) or []

        job_pk = existing_pk("jobs", ["id", "job_id"])
        job_spec_cols = existing_columns("jobs", ["spec_id", "specialization_id"])
        jobs = []
        if job_spec_cols:
            where = " OR ".join([f"`{col}`=%s" for col in job_spec_cols])
            jobs = query_db(
                f"SELECT * FROM jobs WHERE {where} ORDER BY `{job_pk}` DESC",
                tuple([real_spec_id] * len(job_spec_cols)),
                fetchall=True
            ) or []

        certificates = []
        cert_table = existing_table("certificates", "certifications")
        if table_exists(cert_table):
            cert_pk = existing_pk(cert_table, ["id", "certificate_id", "certification_id"])
            cert_spec_cols = existing_columns(cert_table, ["spec_id", "specialization_id"])
            if cert_spec_cols:
                where = " OR ".join([f"`{col}`=%s" for col in cert_spec_cols])
                certificates = query_db(
                    f"SELECT * FROM `{cert_table}` WHERE {where} ORDER BY `{cert_pk}` DESC",
                    tuple([real_spec_id] * len(cert_spec_cols)),
                    fetchall=True
                ) or []

        return jsonify({
            "ok": True,
            "specialization": spec,
            "courses": [normalize_course(row) for row in courses],
            "jobs": [normalize_job(row) for row in jobs],
            "certificates": [normalize_certificate(row) for row in certificates],
            "data": spec,
        }), 200
    except Exception as e:
        print("GET /api/specializations/<id> ERROR:", e)
        return jsonify({"ok": False, "error": "Failed to load specialization details", "details": str(e)}), 500


@app.route("/api/specializations/<int:spec_id>/enrollment-status", methods=["GET"])
@login_required
def specialization_enrollment_status(spec_id):
    try:
        user_id = get_logged_user_id()
        if not user_id:
            return jsonify({"error": "Unauthorized"}), 401
        spec = fetch_specialization_by_any_id(spec_id)
        if not spec:
            return jsonify({"error": "Specialization not found"}), 404
        real_spec_id = safe_int(spec.get("id"), spec_id)
        enrollment = query_db(
            """
            SELECT id AS enrollment_id, progress, progress AS progress_percentage, status, enrolled_at, completed_at
            FROM specialization_enrollments
            WHERE user_id=%s AND spec_id=%s
            LIMIT 1
            """,
            (user_id, real_spec_id),
            fetchone=True
        )
        progress = safe_int(row_value(enrollment or {}, "progress_percentage", "progress"), 0)
        raw_status = safe_text((enrollment or {}).get("status")) or "not_started"
        return jsonify({
            "success": True,
            "specialization_id": real_spec_id,
            "id": real_spec_id,
            "enrolled": bool(enrollment),
            "progress": progress,
            "progress_percentage": progress,
            "status": raw_status,
            "status_label": raw_status.replace("_", " ").title(),
        })
    except Exception as e:
        print("SPECIALIZATION STATUS ERROR:", e)
        return jsonify({"error": "Failed to load enrollment status", "details": str(e)}), 500


@app.route("/api/specializations/<int:spec_id>/enroll", methods=["POST"])
@login_required
def enroll_specialization(spec_id):
    try:
        user_id = get_logged_user_id()
        if not user_id:
            return jsonify({"error": "Unauthorized"}), 401
        spec = fetch_specialization_by_any_id(spec_id)
        if not spec:
            return jsonify({"error": "Specialization not found"}), 404
        real_spec_id = safe_int(spec.get("id"), spec_id)
        query_db(
            """
            INSERT INTO specialization_enrollments (user_id, spec_id, specialization_id, progress, progress_percentage, status, enrolled_at)
            VALUES (%s, %s, %s, 0, 0, 'In Progress', CURRENT_TIMESTAMP)
            ON DUPLICATE KEY UPDATE
                spec_id=VALUES(spec_id),
                progress=GREATEST(COALESCE(progress,0), VALUES(progress)),
                progress_percentage=GREATEST(COALESCE(progress_percentage,0), VALUES(progress_percentage)),
                status=CASE WHEN LOWER(status) IN ('completed') THEN status ELSE 'In Progress' END,
                enrolled_at=enrolled_at
            """,
            (user_id, real_spec_id, real_spec_id),
            commit=True
        )
        compute_user_progress(user_id)
        return jsonify({"success": True, "message": "Enrolled successfully", "specialization_id": real_spec_id, "id": real_spec_id})
    except Exception as e:
        print("SPECIALIZATION ENROLL ERROR:", e)
        return jsonify({"error": "Failed to enroll specialization", "details": str(e)}), 500


@app.route("/api/specializations/<int:spec_id>/unenroll", methods=["DELETE", "POST"])
@login_required
def unenroll_specialization(spec_id):
    try:
        user_id = get_logged_user_id()
        if not user_id:
            return jsonify({"error": "Unauthorized"}), 401
        spec = fetch_specialization_by_any_id(spec_id)
        real_spec_id = safe_int((spec or {}).get("id"), spec_id)

        query_db("DELETE FROM specialization_enrollments WHERE user_id=%s AND spec_id=%s", (user_id, real_spec_id), commit=True)
        if table_exists("course_enrollments"):
            query_db(
                """
                DELETE ce
                FROM course_enrollments ce
                JOIN courses c ON c.id=ce.course_id
                WHERE ce.user_id=%s AND (c.spec_id=%s OR c.specialization_id=%s)
                """,
                (user_id, real_spec_id, real_spec_id),
                commit=True
            )
        if table_exists("user_completed_courses"):
            query_db(
                """
                DELETE ucc
                FROM user_completed_courses ucc
                JOIN courses c ON c.id=ucc.course_id
                WHERE ucc.user_id=%s AND (c.spec_id=%s OR c.specialization_id=%s)
                """,
                (user_id, real_spec_id, real_spec_id),
                commit=True
            )
        if table_exists("specialization_progress"):
            query_db("DELETE FROM specialization_progress WHERE user_id=%s AND specialization_id=%s", (user_id, real_spec_id), commit=True)
        return jsonify({"success": True, "message": "Unenrolled successfully", "specialization_id": real_spec_id, "id": real_spec_id})
    except Exception as e:
        print("SPECIALIZATION UNENROLL ERROR:", e)
        return jsonify({"error": "Failed to unenroll specialization", "details": str(e)}), 500



@app.route("/api/admin/specializations", methods=["POST"])
@app.route("/api/specializations", methods=["POST"])
@admin_required
def add_specialization():
    data = request_data()
    image = normalize_inserted_image(data)
    name = safe_text(data.get("name"))
    if not name:
        return jsonify({"ok": False, "error": "Specialization name is required"}), 400
    table = admin_spec_table()
    spec_id = admin_insert_dynamic(table, {
        "name": name,
        "description": safe_text(data.get("description")),
        "roadmap": safe_text(data.get("roadmap")),
        "job_titles": safe_text(data.get("job_titles")),
        "career_paths": safe_text(data.get("career_paths")),
        "skills": safe_text(data.get("skills")),
        "image_url": image,
        "image": image,
    })
    return jsonify({"ok": True, "message": "Specialization added", "id": spec_id, "specialization_id": spec_id})


@app.route("/api/admin/specializations/<int:spec_id>", methods=["PUT"])
@app.route("/api/specializations/<int:spec_id>", methods=["PUT"])
@admin_required
def admin_update_specialization(spec_id):
    data = request_data()
    image = normalize_inserted_image(data)
    name = safe_text(data.get("name"))
    if not name:
        return jsonify({"ok": False, "error": "Specialization name is required"}), 400
    table = admin_spec_table()
    payload = {
        "name": name,
        "description": safe_text(data.get("description")),
        "roadmap": safe_text(data.get("roadmap")),
        "job_titles": safe_text(data.get("job_titles")),
        "career_paths": safe_text(data.get("career_paths")),
        "skills": safe_text(data.get("skills")),
    }
    if image:
        payload["image_url"] = image
        payload["image"] = image
    admin_update_dynamic(table, spec_id, ["id", "specialization_id", "spec_id"], payload)
    return jsonify({"ok": True, "success": True, "message": "Specialization updated successfully"})


@app.route("/api/admin/specializations/<int:spec_id>", methods=["DELETE"])
@app.route("/api/specializations/<int:spec_id>", methods=["DELETE"])
@admin_required
def delete_specialization(spec_id):
    table = admin_spec_table()
    admin_delete_dynamic(table, spec_id, ["id", "specialization_id", "spec_id"])
    return jsonify({"ok": True, "message": "Specialization deleted"})





@app.route("/api/admin/courses", methods=["GET"])
@app.route("/api/courses", methods=["GET"])
def get_courses():
    try:
        search = safe_text(request.args.get("search"))
        spec_id = request.args.get("specialization_id") or request.args.get("spec_id")
        spec_table = admin_spec_table()
        spec_pk = existing_pk(spec_table, ["id", "specialization_id", "spec_id"])
        course_pk = existing_pk("courses", ["id", "course_id"])
        course_spec = coalesce_existing("c", "courses", ["spec_id", "specialization_id"], "NULL")
        searchable = existing_columns("courses", ["title", "name", "description", "level", "difficulty"])
        sql = f"""
            SELECT c.*, s.name AS specialization_name
            FROM courses c
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={course_spec}
            WHERE 1=1
        """
        params = []
        if search and searchable:
            sql += " AND (" + " OR ".join([f"c.`{col}` LIKE %s" for col in searchable]) + ")"
            params.extend([f"%{search}%"] * len(searchable))
        if spec_id and course_spec != "NULL":
            refs = existing_columns("courses", ["spec_id", "specialization_id"])
            sql += " AND (" + " OR ".join([f"c.`{col}`=%s" for col in refs]) + ")"
            params.extend([spec_id] * len(refs))
        sql += f" ORDER BY c.`{course_pk}` DESC"
        rows = query_db(sql, tuple(params), fetchall=True) or []
        return jsonify({"ok": True, "courses": [normalize_course(row) for row in rows], "data": [normalize_course(row) for row in rows]}), 200
    except Exception as e:
        print("GET /api/courses ERROR:", e)
        return jsonify({"ok": False, "error": "Server error", "details": str(e)}), 500


@app.route("/api/courses/<int:course_id>", methods=["GET"])
def get_course(course_id):
    try:
        spec_table = admin_spec_table()
        spec_pk = existing_pk(spec_table, ["id", "specialization_id", "spec_id"])
        course_pk = existing_pk("courses", ["id", "course_id"])
        course_spec = coalesce_existing("c", "courses", ["spec_id", "specialization_id"], "NULL")
        where_sql, cols = where_existing_id("courses", "c", ["id", "course_id"])
        row = query_db(
            f"""
            SELECT c.*, s.name AS specialization_name
            FROM courses c
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={course_spec}
            WHERE {where_sql}
            LIMIT 1
            """,
            tuple([course_id] * len(cols)),
            fetchone=True
        )
        if not row:
            return jsonify({"error": "Course not found"}), 404
        real_course_id = safe_int(row_value(row, course_pk, "id", "course_id"), course_id)
        quiz_pk = existing_pk("quizzes", ["id", "quiz_id"])
        quizzes = query_db(f"SELECT * FROM quizzes WHERE course_id=%s ORDER BY `{quiz_pk}` DESC", (real_course_id,), fetchall=True) or []
        return jsonify({"course": normalize_course(row), "quizzes": [normalize_quiz(q) for q in quizzes]}), 200
    except Exception as e:
        print("GET /api/courses/<id> ERROR:", e)
        return jsonify({"error": "Server error", "details": str(e)}), 500


@app.route("/api/admin/courses", methods=["POST"])
@app.route("/api/courses", methods=["POST"])
@admin_required
def add_course():
    data = request_data()
    image = normalize_inserted_image(data)
    video = normalize_inserted_video(data)
    title = safe_text(data.get("title") or data.get("name"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not title:
        return jsonify({"ok": False, "error": "Course title is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    spec = fetch_specialization_by_any_id(spec_id)
    if not spec:
        return jsonify({"ok": False, "error": "Specialization not found"}), 404
    real_spec_id = safe_int(row_value(spec, "id", "specialization_id", "spec_id"), spec_id)
    level = normalize_level(data.get("level") or data.get("difficulty"))
    link = safe_text(data.get("course_link") or data.get("link"))
    course_id = admin_insert_dynamic("courses", {
        "spec_id": real_spec_id,
        "specialization_id": real_spec_id,
        "title": title,
        "name": title,
        "description": safe_text(data.get("description") or data.get("content")),
        "level": level,
        "difficulty": level,
        "course_link": link,
        "link": link,
        "image_url": image,
        "image": image,
        "video_url": video,
        "video": video,
    })
    return jsonify({"ok": True, "message": "Course added", "id": course_id, "course_id": course_id})


@app.route("/api/admin/courses/<int:course_id>", methods=["PUT"])
@app.route("/api/courses/<int:course_id>", methods=["PUT"])
@admin_required
def admin_update_course(course_id):
    data = request_data()
    title = safe_text(data.get("title") or data.get("name"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not title:
        return jsonify({"ok": False, "error": "Course title is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    spec = fetch_specialization_by_any_id(spec_id)
    if not spec:
        return jsonify({"ok": False, "error": "Specialization not found"}), 404
    real_spec_id = safe_int(row_value(spec, "id", "specialization_id", "spec_id"), spec_id)
    image = normalize_inserted_image(data)
    video = normalize_inserted_video(data)
    level = normalize_level(data.get("level") or data.get("difficulty"))
    link = safe_text(data.get("course_link") or data.get("link"))
    payload = {
        "title": title,
        "name": title,
        "description": safe_text(data.get("description") or data.get("content")),
        "specialization_id": real_spec_id,
        "spec_id": real_spec_id,
        "level": level,
        "difficulty": level,
        "course_link": link,
        "link": link,
    }
    if image:
        payload["image_url"] = image
        payload["image"] = image
    if video:
        payload["video_url"] = video
        payload["video"] = video
    admin_update_dynamic("courses", course_id, ["id", "course_id"], payload)
    return jsonify({"ok": True, "success": True, "message": "Course updated successfully"})


@app.route("/api/admin/courses/<int:course_id>", methods=["DELETE"])
@app.route("/api/courses/<int:course_id>", methods=["DELETE"])
@admin_required
def delete_course(course_id):
    admin_delete_dynamic("courses", course_id, ["id", "course_id"])
    return jsonify({"ok": True, "message": "Course deleted"})


@app.route("/api/courses/<int:course_id>/enrollment-status", methods=["GET"])
@login_required
def course_enrollment_status(course_id):
    try:
        user_id = request.current_user["id"]
        course = fetch_course_by_any_id(course_id)
        if not course:
            return jsonify({"error": "Course not found"}), 404
        real_course_id = safe_int(course.get("id"), course_id)
        enrollment = query_db(
            """
            SELECT id AS enrollment_id, progress, progress_percentage, status, enrolled_at, completed_at
            FROM course_enrollments
            WHERE user_id=%s AND course_id=%s
            LIMIT 1
            """,
            (user_id, real_course_id),
            fetchone=True
        )
        spec_id = row_value(course, "spec_id", "specialization_id")
        spec_enrollment = None
        if spec_id:
            spec_enrollment = query_db(
                """
                SELECT id AS enrollment_id, progress, status
                FROM specialization_enrollments
                WHERE user_id=%s AND spec_id=%s
                LIMIT 1
                """,
                (user_id, spec_id),
                fetchone=True
            )
        progress = safe_int(row_value(enrollment or {}, "progress_percentage", "progress"), 0)
        raw_status = safe_text((enrollment or {}).get("status")) or "not_started"
        return jsonify({
            "success": True,
            "course_id": real_course_id,
            "id": real_course_id,
            "specialization_id": spec_id,
            "enrolled": bool(enrollment),
            "specialization_enrolled": bool(spec_enrollment),
            "progress": progress,
            "progress_percentage": progress,
            "status": raw_status,
            "status_label": raw_status.replace("_", " ").title(),
        })
    except Exception as exc:
        print("COURSE STATUS ERROR:", exc)
        return jsonify({"error": "Failed to load course enrollment status", "details": str(exc)}), 500


@app.route("/api/courses/<int:course_id>/enroll", methods=["POST"])
@student_required
def enroll_course(course_id):
    try:
        user_id = request.current_user["id"]
        course = fetch_course_by_any_id(course_id)
        if not course:
            return jsonify({"error": "Course not found"}), 404
        real_course_id = safe_int(course.get("id"), course_id)
        spec_id = row_value(course, "spec_id", "specialization_id")
        if spec_id:
            spec_enrollment = query_db("SELECT id FROM specialization_enrollments WHERE user_id=%s AND spec_id=%s LIMIT 1", (user_id, spec_id), fetchone=True)
            if not spec_enrollment:
                return jsonify({"error": "Enroll in the specialization first.", "needs_specialization_enrollment": True, "specialization_id": spec_id}), 409
        ensure_course_enrollment(user_id, real_course_id, 0)
        compute_user_progress(user_id)
        return jsonify({"success": True, "message": "Course enrolled successfully", "course_id": real_course_id, "id": real_course_id})
    except Exception as exc:
        print("COURSE ENROLL ERROR:", exc)
        return jsonify({"error": "Failed to enroll course", "details": str(exc)}), 500


@app.route("/api/courses/<int:course_id>/unenroll", methods=["DELETE", "POST"])
@student_required
def unenroll_course(course_id):
    try:
        user_id = request.current_user["id"]
        course = fetch_course_by_any_id(course_id)
        real_course_id = safe_int((course or {}).get("id"), course_id)
        query_db("DELETE FROM course_enrollments WHERE user_id=%s AND course_id=%s", (user_id, real_course_id), commit=True)
        if table_exists("user_completed_courses"):
            query_db("DELETE FROM user_completed_courses WHERE user_id=%s AND course_id=%s", (user_id, real_course_id), commit=True)
        compute_user_progress(user_id)
        return jsonify({"success": True, "message": "Course unenrolled successfully", "course_id": real_course_id, "id": real_course_id})
    except Exception as exc:
        print("COURSE UNENROLL ERROR:", exc)
        return jsonify({"error": "Failed to unenroll course", "details": str(exc)}), 500


@app.route("/api/courses/<int:course_id>/open", methods=["POST"])
@student_required
def open_course(course_id):
    try:
        course = fetch_course_by_any_id(course_id)
        if not course:
            return jsonify({"error": "Course not found"}), 404
        real_course_id = safe_int(course.get("id"), course_id)
        user_id = request.current_user["id"]
        data = get_json()
        completed = bool(data.get("completed"))
        spec_id = row_value(course, "spec_id", "specialization_id")
        if spec_id:
            spec_enrollment = query_db("SELECT id FROM specialization_enrollments WHERE user_id=%s AND spec_id=%s LIMIT 1", (user_id, spec_id), fetchone=True)
            if not spec_enrollment:
                return jsonify({"success": True, "message": "Course opened, but progress was not changed because you are not enrolled in the specialization.", "tracked": False, "enrolled": False, "course_id": real_course_id, "needs_specialization_enrollment": True, "specialization_id": spec_id})
        enrollment = query_db("SELECT id, progress, progress_percentage, status FROM course_enrollments WHERE user_id=%s AND course_id=%s LIMIT 1", (user_id, real_course_id), fetchone=True)
        if not enrollment:
            ensure_course_enrollment(user_id, real_course_id, 25)
            progress_value = 25
        else:
            current_progress = safe_int(row_value(enrollment, "progress_percentage", "progress"), 0)
            progress_value = max(current_progress, 100 if completed else 25)
            set_course_progress(user_id, real_course_id, progress_value, completed)
        if table_exists("course_activity"):
            try:
                query_db("INSERT INTO course_activity (user_id, course_id, action) VALUES (%s,%s,%s)", (user_id, real_course_id, "completed" if completed else "opened"), commit=True)
            except Exception as exc:
                print("COURSE ACTIVITY SAVE ERROR:", exc)
        if completed and table_exists("user_completed_courses"):
            query_db("INSERT IGNORE INTO user_completed_courses (user_id, course_id) VALUES (%s,%s)", (user_id, real_course_id), commit=True)
        compute_user_progress(user_id)
        return jsonify({"success": True, "message": "Course progress tracked", "tracked": True, "enrolled": True, "course_id": real_course_id, "id": real_course_id, "progress": progress_value})
    except Exception as exc:
        print("COURSE OPEN ERROR:", exc)
        return jsonify({"error": "Failed to track course", "details": str(exc)}), 500




@app.route("/api/admin/quizzes", methods=["GET"])
@app.route("/api/quizzes", methods=["GET"])
def get_quizzes():
    course_id = request.args.get("course_id")
    course_pk = existing_pk("courses", ["id", "course_id"])
    quiz_pk = existing_pk("quizzes", ["id", "quiz_id"])
    sql = f"""
        SELECT q.*, COALESCE(c.title, c.name) AS course_title
        FROM quizzes q
        LEFT JOIN courses c ON c.`{course_pk}`=q.course_id
        WHERE 1=1
    """
    params = []
    if course_id:
        sql += " AND q.course_id=%s"
        params.append(course_id)
    sql += f" ORDER BY q.`{quiz_pk}` DESC"
    rows = query_db(sql, tuple(params), fetchall=True) or []
    return jsonify({"ok": True, "quizzes": [normalize_quiz(row) for row in rows], "data": [normalize_quiz(row) for row in rows]})


@app.route("/api/quizzes/<int:quiz_id>", methods=["GET"])
def get_quiz(quiz_id):
    course_pk = existing_pk("courses", ["id", "course_id"])
    where_sql, cols = where_existing_id("quizzes", "q", ["id", "quiz_id"])
    quiz = query_db(
        f"""
        SELECT q.*, COALESCE(c.title, c.name) AS course_title
        FROM quizzes q
        LEFT JOIN courses c ON c.`{course_pk}`=q.course_id
        WHERE {where_sql}
        LIMIT 1
        """,
        tuple([quiz_id] * len(cols)),
        fetchone=True
    )
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404
    real_quiz_id = safe_int(row_value(quiz, "id", "quiz_id"), quiz_id)
    question_pk = existing_pk("quiz_questions", ["id", "question_id"])
    questions = query_db(f"SELECT * FROM quiz_questions WHERE quiz_id=%s ORDER BY `{question_pk}`", (real_quiz_id,), fetchall=True) or []
    return jsonify({"quiz": normalize_quiz(quiz), "questions": [normalize_question(row) for row in questions]})


@app.route("/api/admin/quizzes", methods=["POST"])
@app.route("/api/quizzes", methods=["POST"])
@admin_required
def add_quiz():
    data = get_json() if request.is_json else request_data()
    title = safe_text(data.get("title") or data.get("name"))
    course_id = safe_int(data.get("course_id"), None)
    if not title or not course_id:
        return jsonify({"ok": False, "error": "Quiz title and course are required"}), 400
    course = fetch_course_by_any_id(course_id)
    if not course:
        return jsonify({"ok": False, "error": "Course not found"}), 404
    real_course_id = safe_int(row_value(course, "id", "course_id"), course_id)
    quiz_id = admin_insert_dynamic("quizzes", {
        "course_id": real_course_id,
        "title": title,
        "name": title,
        "description": safe_text(data.get("description")),
        "total_questions": 0,
        "spec_id": row_value(course, "spec_id", "specialization_id"),
    })
    questions = data.get("questions") or []
    if not questions and data.get("questions_json"):
        try:
            questions = json.loads(data.get("questions_json"))
        except Exception:
            questions = []
    count = 0
    for q in questions:
        add_question_to_quiz(quiz_id, q)
        count += 1
    if column_exists("quizzes", "total_questions"):
        pk = existing_pk("quizzes", ["id", "quiz_id"])
        exec_db(f"UPDATE quizzes SET total_questions=%s WHERE `{pk}`=%s", (count, quiz_id))
    return jsonify({"ok": True, "message": "Quiz added", "id": quiz_id, "quiz_id": quiz_id})


def add_question_to_quiz(quiz_id, data):
    question = safe_text(data.get("question") or data.get("question_text"))
    option_a = safe_text(data.get("option_a") or data.get("option1"))
    option_b = safe_text(data.get("option_b") or data.get("option2"))
    option_c = safe_text(data.get("option_c") or data.get("option3"))
    option_d = safe_text(data.get("option_d") or data.get("option4"))
    answer = safe_text(data.get("correct_answer") or data.get("answer")).upper()
    aliases = {"1": "A", "2": "B", "3": "C", "4": "D", option_a.upper(): "A", option_b.upper(): "B", option_c.upper(): "C", option_d.upper(): "D"}
    answer = aliases.get(answer, answer if answer in ["A", "B", "C", "D"] else "A")
    return query_db(
        """
        INSERT INTO quiz_questions
            (quiz_id,question_text,option_a,option_b,option_c,option_d,correct_answer,question,option1,option2,option3,option4,answer)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """,
        (quiz_id, question, option_a, option_b, option_c, option_d, answer, question, option_a, option_b, option_c, option_d, answer),
        commit=True
    )


@app.route("/api/quizzes/<int:quiz_id>/questions", methods=["POST"])
@admin_required
def add_quiz_question(quiz_id):
    quiz = fetch_quiz_by_any_id(quiz_id)
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404
    real_quiz_id = safe_int(row_value(quiz, "id", "quiz_id"), quiz_id)
    question_id = add_question_to_quiz(real_quiz_id, get_json())
    row = query_db("SELECT COUNT(*) AS total FROM quiz_questions WHERE quiz_id=%s", (real_quiz_id,), fetchone=True) or {"total": 0}
    if column_exists("quizzes", "total_questions"):
        pk = existing_pk("quizzes", ["id", "quiz_id"])
        exec_db(f"UPDATE quizzes SET total_questions=%s WHERE `{pk}`=%s", (safe_int(row.get("total"), 0), real_quiz_id))
    return jsonify({"ok": True, "message": "Question added", "id": question_id})


@app.route("/api/admin/quizzes/<int:quiz_id>", methods=["PUT"])
@app.route("/api/quizzes/<int:quiz_id>", methods=["PUT"])
@admin_required
def admin_update_quiz(quiz_id):
    data = get_json() if request.is_json else request_data()
    title = safe_text(data.get("title") or data.get("name"))
    course_id = safe_int(data.get("course_id"), None)
    if not title:
        return jsonify({"ok": False, "error": "Quiz title is required"}), 400
    payload = {
        "title": title,
        "name": title,
        "description": safe_text(data.get("description")),
    }
    if course_id:
        course = fetch_course_by_any_id(course_id)
        if not course:
            return jsonify({"ok": False, "error": "Course not found"}), 404
        payload["course_id"] = safe_int(row_value(course, "id", "course_id"), course_id)
        payload["spec_id"] = row_value(course, "spec_id", "specialization_id")
    admin_update_dynamic("quizzes", quiz_id, ["id", "quiz_id"], payload)
    questions = data.get("questions") or []
    if not questions and data.get("questions_json"):
        try:
            questions = json.loads(data.get("questions_json"))
        except Exception:
            questions = []
    if questions:
        real_quiz = fetch_quiz_by_any_id(quiz_id)
        real_quiz_id = safe_int(row_value(real_quiz, "id", "quiz_id"), quiz_id)
        query_db("DELETE FROM quiz_questions WHERE quiz_id=%s", (real_quiz_id,), commit=True)
        for q in questions:
            add_question_to_quiz(real_quiz_id, q)
        if column_exists("quizzes", "total_questions"):
            pk = existing_pk("quizzes", ["id", "quiz_id"])
            exec_db(f"UPDATE quizzes SET total_questions=%s WHERE `{pk}`=%s", (len(questions), real_quiz_id))
    return jsonify({"ok": True, "success": True, "message": "Quiz updated successfully"})


@app.route("/api/admin/quizzes/<int:quiz_id>", methods=["DELETE"])
@app.route("/api/quizzes/<int:quiz_id>", methods=["DELETE"])
@admin_required
def delete_quiz(quiz_id):
    quiz = fetch_quiz_by_any_id(quiz_id)
    real_quiz_id = safe_int(row_value(quiz or {}, "id", "quiz_id"), quiz_id)
    query_db("DELETE FROM quiz_questions WHERE quiz_id=%s", (real_quiz_id,), commit=True)
    admin_delete_dynamic("quizzes", quiz_id, ["id", "quiz_id"])
    return jsonify({"ok": True, "message": "Quiz deleted"})


@app.route("/api/quizzes/<int:quiz_id>/submit", methods=["POST"])
@student_required
def submit_quiz(quiz_id):
    data = get_json()
    answers = data.get("answers") or {}
    quiz = fetch_quiz_by_any_id(quiz_id)
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404
    real_quiz_id = safe_int(quiz.get("id"), quiz_id)
    questions = query_db("SELECT * FROM quiz_questions WHERE quiz_id=%s ORDER BY id", (real_quiz_id,), fetchall=True) or []
    score = 0
    details = []
    for q in questions:
        qid_value = row_value(q, "id", "question_id")
        qid = str(qid_value)
        given = safe_text(answers.get(qid) or answers.get(str(row_value(q, "question_id", "id")))).upper()
        correct = safe_text(row_value(q, "correct_answer", "answer")).upper()
        normalized = {"1": "A", "2": "B", "3": "C", "4": "D"}.get(given, given)
        ok = normalized == correct
        if ok:
            score += 1
        details.append({"question_id": qid, "given": normalized, "correct": correct, "correct_boolean": bool(ok)})
    total = len(questions)
    percentage = round((score / total) * 100) if total else 0
    passed = 1 if percentage >= safe_int(quiz.get("passing_score"), 60) else 0
    user_id = request.current_user["id"]
    course_id = quiz.get("course_id")
    attempt_id = query_db(
        """
        INSERT INTO quiz_attempts (user_id, quiz_id, course_id, score, passed, total, percentage, answers_json)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
        """,
        (user_id, real_quiz_id, course_id, percentage, passed, total, percentage, json.dumps(answers)),
        commit=True
    )
    if passed and table_exists("user_completed_quizzes"):
        query_db(
            """
            INSERT INTO user_completed_quizzes (user_id, quiz_id, score)
            VALUES (%s,%s,%s)
            ON DUPLICATE KEY UPDATE score=GREATEST(score,VALUES(score)), completed_at=CURRENT_TIMESTAMP
            """,
            (user_id, real_quiz_id, percentage),
            commit=True
        )
    course_progress_tracked = False
    if course_id:
        enrollment = query_db("SELECT id, progress, progress_percentage FROM course_enrollments WHERE user_id=%s AND course_id=%s LIMIT 1", (user_id, course_id), fetchone=True)
        if enrollment:
            progress_value = 100 if passed else max(safe_int(row_value(enrollment, "progress_percentage", "progress"), 0), 50)
            set_course_progress(user_id, course_id, progress_value, bool(passed))
            course_progress_tracked = True
            if passed and table_exists("user_completed_courses"):
                query_db("INSERT IGNORE INTO user_completed_courses (user_id, course_id) VALUES (%s,%s)", (user_id, course_id), commit=True)
    compute_user_progress(user_id)
    return jsonify({
        "message": "Quiz submitted",
        "attempt_id": attempt_id,
        "score": score,
        "total": total,
        "score_percentage": percentage,
        "passed": bool(passed),
        "course_progress_tracked": course_progress_tracked,
        "details": details
    })



@app.route("/api/admin/jobs", methods=["GET"])
@app.route("/api/jobs", methods=["GET"])
def get_jobs():
    try:
        search = safe_text(request.args.get("search"))
        spec_id = request.args.get("specialization_id") or request.args.get("spec_id")
        spec_table = admin_spec_table()
        spec_pk = existing_pk(spec_table, ["id", "specialization_id", "spec_id"])
        job_pk = existing_pk("jobs", ["id", "job_id"])
        job_spec = coalesce_existing("j", "jobs", ["specialization_id", "spec_id"], "NULL")
        searchable = existing_columns("jobs", ["title", "description", "required_skills", "skills"])
        sql = f"""
            SELECT j.*, s.name AS specialization_name
            FROM jobs j
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={job_spec}
            WHERE 1=1
        """
        params = []
        if search and searchable:
            sql += " AND (" + " OR ".join([f"j.`{col}` LIKE %s" for col in searchable]) + ")"
            params.extend([f"%{search}%"] * len(searchable))
        if spec_id and job_spec != "NULL":
            refs = existing_columns("jobs", ["specialization_id", "spec_id"])
            sql += " AND (" + " OR ".join([f"j.`{col}`=%s" for col in refs]) + ")"
            params.extend([spec_id] * len(refs))
        sql += f" ORDER BY j.`{job_pk}` DESC"
        rows = query_db(sql, tuple(params), fetchall=True) or []
        return jsonify({"ok": True, "jobs": [normalize_job(row) for row in rows], "data": [normalize_job(row) for row in rows]}), 200
    except Exception as e:
        print("GET /api/jobs ERROR:", e)
        return jsonify({"ok": False, "error": "Server error", "details": str(e)}), 500


@app.route("/api/jobs/<int:job_id>", methods=["GET"])
def get_job(job_id):
    try:
        spec_table = admin_spec_table()
        spec_pk = existing_pk(spec_table, ["id", "specialization_id", "spec_id"])
        job_spec = coalesce_existing("j", "jobs", ["specialization_id", "spec_id"], "NULL")
        where_sql, cols = where_existing_id("jobs", "j", ["id", "job_id"])
        row = query_db(
            f"""
            SELECT j.*, s.name AS specialization_name
            FROM jobs j
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={job_spec}
            WHERE {where_sql}
            LIMIT 1
            """,
            tuple([job_id] * len(cols)),
            fetchone=True
        )
        if not row:
            return jsonify({"error": "Job not found"}), 404
        return jsonify({"job": normalize_job(row)}), 200
    except Exception as e:
        print("GET /api/jobs/<id> ERROR:", e)
        return jsonify({"error": "Server error", "details": str(e)}), 500

@app.route("/api/admin/jobs", methods=["POST"])
@app.route("/api/jobs", methods=["POST"])
@admin_required
def add_job():
    data = get_json() if request.is_json else request_data()
    title = safe_text(data.get("title"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not title:
        return jsonify({"ok": False, "error": "Job title is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    job_id = admin_insert_dynamic("jobs", {
        "specialization_id": spec_id,
        "spec_id": spec_id,
        "title": title,
        "description": safe_text(data.get("description")),
        "required_skills": safe_text(data.get("required_skills") or data.get("skills")),
        "skills": safe_text(data.get("required_skills") or data.get("skills")),
        "average_salary": safe_text(data.get("average_salary") or data.get("salary")),
        "salary": safe_text(data.get("average_salary") or data.get("salary")),
        "job_link": safe_text(data.get("job_link") or data.get("link")),
        "link": safe_text(data.get("job_link") or data.get("link")),
    })
    return jsonify({"ok": True, "message": "Job added", "id": job_id, "job_id": job_id})


@app.route("/api/admin/jobs/<int:job_id>", methods=["PUT"])
@app.route("/api/jobs/<int:job_id>", methods=["PUT"])
@admin_required
def admin_update_job(job_id):
    data = get_json() if request.is_json else request_data()
    title = safe_text(data.get("title"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not title:
        return jsonify({"ok": False, "error": "Job title is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    admin_update_dynamic("jobs", job_id, ["id", "job_id"], {
        "title": title,
        "description": safe_text(data.get("description")),
        "required_skills": safe_text(data.get("required_skills") or data.get("skills")),
        "skills": safe_text(data.get("required_skills") or data.get("skills")),
        "specialization_id": spec_id,
        "spec_id": spec_id,
        "average_salary": safe_text(data.get("average_salary") or data.get("salary")),
        "salary": safe_text(data.get("average_salary") or data.get("salary")),
        "job_link": safe_text(data.get("job_link") or data.get("link")),
        "link": safe_text(data.get("job_link") or data.get("link")),
    })
    return jsonify({"ok": True, "success": True, "message": "Job updated successfully"})


@app.route("/api/admin/jobs/<int:job_id>", methods=["DELETE"])
@app.route("/api/jobs/<int:job_id>", methods=["DELETE"])
@admin_required
def delete_job(job_id):
    admin_delete_dynamic("jobs", job_id, ["id", "job_id"])
    return jsonify({"ok": True, "message": "Job deleted"})


@app.route("/api/admin/certificates", methods=["GET"])
@app.route("/api/certificates", methods=["GET"])
def get_certificates():
    rows = []
    spec_table = admin_spec_table()
    spec_pk = existing_pk(spec_table, ["id", "specialization_id", "spec_id"])
    if table_exists("certificates"):
        cert_pk = existing_pk("certificates", ["id", "certificate_id", "certification_id"])
        cert_spec = coalesce_existing("c", "certificates", ["spec_id", "specialization_id"], "NULL")
        rows += query_db(
            f"""
            SELECT c.*, c.`{cert_pk}` AS id, {cert_spec} AS specialization_id, s.name AS specialization_name
            FROM certificates c
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={cert_spec}
            ORDER BY c.`{cert_pk}` DESC
            """,
            fetchall=True
        ) or []
    if table_exists("certifications"):
        cert_pk = existing_pk("certifications", ["certification_id", "id", "certificate_id"])
        cert_spec = coalesce_existing("c", "certifications", ["specialization_id", "spec_id"], "NULL")
        link_expr = coalesce_existing("c", "certifications", ["official_link", "link"], "''")
        rows += query_db(
            f"""
            SELECT c.*, c.`{cert_pk}` AS id, {cert_spec} AS specialization_id, {link_expr} AS link, s.name AS specialization_name
            FROM certifications c
            LEFT JOIN `{spec_table}` s ON s.`{spec_pk}`={cert_spec}
            ORDER BY c.`{cert_pk}` DESC
            """,
            fetchall=True
        ) or []
    return jsonify({"ok": True, "certificates": rows, "data": rows})


@app.route("/api/admin/certificates", methods=["POST"])
@app.route("/api/certificates", methods=["POST"])
@admin_required
def add_certificate():
    data = get_json() if request.is_json else request_data()
    name = safe_text(data.get("name"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not name:
        return jsonify({"ok": False, "error": "Certificate name is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    table = existing_table("certificates", "certifications")
    cert_id = admin_insert_dynamic(table, {
        "spec_id": spec_id,
        "specialization_id": spec_id,
        "name": name,
        "description": safe_text(data.get("description")),
        "link": safe_text(data.get("link") or data.get("official_link") or data.get("certificate_url")),
        "official_link": safe_text(data.get("link") or data.get("official_link") or data.get("certificate_url")),
        "price": safe_text(data.get("price")),
        "type": safe_text(data.get("type")).lower() or "both",
    })
    return jsonify({"ok": True, "message": "Certificate added", "id": cert_id, "certificate_id": cert_id})


@app.route("/api/admin/certificates/<int:cert_id>", methods=["PUT"])
@app.route("/api/certificates/<int:cert_id>", methods=["PUT"])
@admin_required
def admin_update_certificate(cert_id):
    data = get_json() if request.is_json else request_data()
    name = safe_text(data.get("name"))
    spec_id = safe_int(data.get("specialization_id") or data.get("spec_id"), None)
    if not name:
        return jsonify({"ok": False, "error": "Certificate name is required"}), 400
    if not spec_id:
        return jsonify({"ok": False, "error": "Specialization is required"}), 400
    table = "certificates" if table_exists("certificates") else "certifications"
    admin_update_dynamic(table, cert_id, ["id", "certificate_id", "certification_id"], {
        "spec_id": spec_id,
        "specialization_id": spec_id,
        "name": name,
        "description": safe_text(data.get("description")),
        "link": safe_text(data.get("link") or data.get("official_link") or data.get("certificate_url")),
        "official_link": safe_text(data.get("link") or data.get("official_link") or data.get("certificate_url")),
        "price": safe_text(data.get("price")),
        "type": safe_text(data.get("type")).lower() or "both",
    })
    return jsonify({"ok": True, "success": True, "message": "Certificate updated successfully"})


@app.route("/api/admin/certificates/<int:cert_id>", methods=["DELETE"])
@app.route("/api/certificates/<int:cert_id>", methods=["DELETE"])
@admin_required
def delete_certificate(cert_id):
    deleted = False
    if table_exists("certificates"):
        admin_delete_dynamic("certificates", cert_id, ["id", "certificate_id", "certification_id"])
        deleted = True
    if table_exists("certifications"):
        try:
            admin_delete_dynamic("certifications", cert_id, ["certification_id", "id", "certificate_id"])
            deleted = True
        except Exception:
            pass
    return jsonify({"ok": True, "message": "Certificate deleted" if deleted else "Certificate delete checked"})




def sqr_slug(value):
    value = safe_text(value).lower()
    value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
    aliases = {
        "cyber_security": "cybersecurity",
        "cybersecurity": "cybersecurity",
        "digital_forensic": "digital_forensics",
        "digital_forensics": "digital_forensics",
        "software_engineer": "software_engineering",
        "software_engineering": "software_engineering",
        "web_development": "web_development",
        "web_developer": "web_development",
        "data_science": "data_science",
        "artificial_intelligence": "artificial_intelligence",
        "ai": "artificial_intelligence",
        "cloud": "cloud_computing",
        "cloud_computing": "cloud_computing",
        "database": "database_administration",
        "database_administration": "database_administration",
        "computer_networks": "computer_networks",
        "networking": "computer_networks",
        "ui_ux": "ui_ux_engineering",
        "ui_ux_engineering": "ui_ux_engineering",
    }
    return aliases.get(value, value)



SQR_AI_SPECIALIZATION_LIBRARY = {
    "data_engineering": {
        "name": "Data Engineering",
        "description": "Designing data pipelines, ETL workflows, databases, warehouses, and reliable data platforms for analytics and AI.",
        "keywords": ["data", "sql", "python", "etl", "pipeline", "warehouse", "database", "cloud", "spark", "airflow", "analytics"],
        "roadmap": ["SQL and Python", "Database design", "ETL and data pipelines", "Cloud data warehouses", "Spark or Airflow projects"],
    },
    "cybersecurity": {
        "name": "Cybersecurity",
        "description": "Protecting systems, monitoring threats, handling incidents, and improving security posture.",
        "keywords": ["security", "cyber", "network", "linux", "forensics", "soc", "siem", "incident", "vulnerability"],
        "roadmap": ["Networking and Linux", "Security fundamentals", "SOC tools and SIEM", "Incident response", "Security portfolio labs"],
    },
    "artificial_intelligence": {
        "name": "Artificial Intelligence",
        "description": "Building intelligent systems using machine learning, data, automation, NLP, and computer vision.",
        "keywords": ["ai", "machine learning", "python", "model", "automation", "nlp", "vision", "statistics", "data"],
        "roadmap": ["Python and math basics", "Machine learning", "Model evaluation", "NLP or computer vision", "AI projects"],
    },
    "software_engineering": {
        "name": "Software Engineering",
        "description": "Designing, building, testing, and maintaining reliable software systems and applications.",
        "keywords": ["software", "java", "python", "api", "backend", "testing", "problem solving", "oop", "database"],
        "roadmap": ["Programming fundamentals", "Data structures", "APIs and databases", "Testing", "Full-stack projects"],
    },
    "web_development": {
        "name": "Web Development",
        "description": "Creating frontend and backend web applications using modern web technologies.",
        "keywords": ["html", "css", "javascript", "frontend", "backend", "react", "flask", "node", "ui"],
        "roadmap": ["HTML CSS JavaScript", "Frontend framework", "Backend API", "Database", "Deployment"],
    },
    "cloud_computing": {
        "name": "Cloud Computing",
        "description": "Deploying, scaling, and managing applications and infrastructure on cloud platforms.",
        "keywords": ["cloud", "aws", "azure", "docker", "deployment", "server", "devops", "kubernetes"],
        "roadmap": ["Linux basics", "Cloud fundamentals", "Docker", "Networking", "Deployment projects"],
    },
    "devops_engineering": {
        "name": "DevOps Engineering",
        "description": "Automating deployment, monitoring, CI/CD, containers, and production operations.",
        "keywords": ["devops", "docker", "kubernetes", "ci", "cd", "deployment", "linux", "cloud", "monitoring"],
        "roadmap": ["Linux and Git", "CI/CD", "Docker", "Cloud deployment", "Monitoring"],
    },
    "database_administration": {
        "name": "Database Administration",
        "description": "Managing databases, schemas, queries, performance, backups, and data security.",
        "keywords": ["database", "sql", "mysql", "postgresql", "schema", "queries", "backup", "admin"],
        "roadmap": ["SQL", "Database design", "Performance tuning", "Backup and recovery", "Security"],
    },
    "computer_networks": {
        "name": "Computer Networks",
        "description": "Designing and troubleshooting networks, routing, switching, protocols, and infrastructure.",
        "keywords": ["network", "tcp", "ip", "routing", "switching", "linux", "security", "protocols"],
        "roadmap": ["Network fundamentals", "Routing and switching", "Linux networking", "Security basics", "Hands-on labs"],
    },
    "ui_ux_engineering": {
        "name": "UI/UX Engineering",
        "description": "Designing usable interfaces, prototypes, user flows, and frontend experiences.",
        "keywords": ["ui", "ux", "design", "interface", "figma", "frontend", "user", "prototype"],
        "roadmap": ["Design principles", "Figma", "User flows", "Frontend basics", "Portfolio case studies"],
    },
}

SPECIALIZATION_HINTS.update({info["name"].lower(): info["keywords"] for info in SQR_AI_SPECIALIZATION_LIBRARY.values()})
SPECIALIZATION_HINTS.update({key: info["keywords"] for key, info in SQR_AI_SPECIALIZATION_LIBRARY.items()})


def sqr_all_specialization_library():
    library = dict(SQR_AI_SPECIALIZATION_LIBRARY)
    for key, hints in SPECIALIZATION_HINTS.items():
        slug = sqr_slug(key)
        if slug not in library:
            library[slug] = {
                "name": safe_text(key).replace("_", " ").title(),
                "description": f"A computer science path focused on {safe_text(key).replace('_', ' ')} skills and related career roles.",
                "keywords": hints,
                "roadmap": ["Learn fundamentals", "Build practical projects", "Complete related courses", "Prepare your resume"],
            }
    return library


def sqr_recommendation_questions():
    base_questions = list(globals().get("SQR_RECOMMENDATION_QUESTION_BANK") or [])
    existing_ids = {safe_text(q.get("id")) for q in base_questions if isinstance(q, dict)}
    for key, info in sqr_all_specialization_library().items():
        name = info.get("name") or key.replace("_", " ").title()
        keywords = info.get("keywords") or []
        skill_text = ", ".join(keywords[:4])
        generated = [
            {
                "id": f"{key}_interest",
                "specialization_key": key,
                "specialization": name,
                "dimension": "interest",
                "question": f"How interested are you in {name} tasks such as {skill_text}?",
                "keywords": keywords,
                "weight": 5,
            },
            {
                "id": f"{key}_skill",
                "specialization_key": key,
                "specialization": name,
                "dimension": "skill",
                "question": f"How confident are you with skills related to {name}?",
                "keywords": keywords,
                "weight": 5,
            },
            {
                "id": f"{key}_career",
                "specialization_key": key,
                "specialization": name,
                "dimension": "career",
                "question": f"Would you enjoy a future job connected to {name}?",
                "keywords": keywords,
                "weight": 5,
            },
        ]
        for q in generated:
            if q["id"] not in existing_ids:
                base_questions.append(q)
                existing_ids.add(q["id"])
    return base_questions

def sqr_rating(value, default=3):
    text = safe_text(value).lower()
    aliases = {
        "very_low": 1, "low": 1, "no": 1, "never": 1,
        "medium_low": 2, "maybe_no": 2,
        "medium": 3, "neutral": 3, "maybe": 3,
        "medium_high": 4, "yes": 4,
        "very_high": 5, "high": 5, "strong": 5, "love": 5,
    }
    if text in aliases:
        return aliases[text]
    try:
        number = int(float(text))
        return max(1, min(5, number))
    except Exception:
        return default


def sqr_parse_recommendation_answers(data):
    answers = []
    raw = data.get("answers") or data.get("quiz_answers") or data.get("recommendation_answers")
    if isinstance(raw, str) and raw.strip():
        try:
            raw = json.loads(raw)
        except Exception:
            raw = None
    if isinstance(raw, dict):
        for key, value in raw.items():
            answers.append({"id": key, "value": value})
    elif isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict):
                answers.append(item)
    for key, value in (data or {}).items():
        if key.startswith("q_") or key.startswith("quiz_") or key.startswith("answer_"):
            clean_key = re.sub(r"^(q_|quiz_|answer_)", "", key)
            answers.append({"id": clean_key, "value": value})
    return answers


def sqr_score_recommendation_quiz(data):
    answers = sqr_parse_recommendation_answers(data)
    questions = sqr_recommendation_questions()
    by_id = {safe_text(q.get("id")): q for q in questions}
    scores = {}
    max_scores = {}
    chosen = []
    for answer in answers:
        qid = safe_text(answer.get("id") or answer.get("question_id") or answer.get("name"))
        value = sqr_rating(answer.get("value") or answer.get("answer") or answer.get("score"), 3)
        q = by_id.get(qid)
        if not q:
            continue
        spec_key = sqr_slug(q.get("specialization_key") or q.get("specialization"))
        weight = safe_int(q.get("weight"), 5) or 5
        scores[spec_key] = scores.get(spec_key, 0) + (value * weight)
        max_scores[spec_key] = max_scores.get(spec_key, 0) + (5 * weight)
        chosen.append({
            "question_id": qid,
            "specialization_key": spec_key,
            "question": q.get("question"),
            "value": value,
            "dimension": q.get("dimension"),
        })
    percentages = {}
    for key, score in scores.items():
        percentages[key] = round((score / max(max_scores.get(key, 1), 1)) * 100)
    return percentages, chosen



def pct_value(value):
    try:
        number = int(round(float(value or 0)))
    except Exception:
        number = 0
    return max(0, min(100, number))

def sqr_spec_key(row):
    return sqr_slug(row_value(row, "key", "slug", "name", "title") or "")


def sqr_recommendation_profile_text(data, user=None):
    parts = [
        safe_text(data.get("interests")),
        safe_text(data.get("skills")),
        safe_text(data.get("technical_skills")),
        safe_text(data.get("soft_skills")),
        safe_text(data.get("preferred_work")),
        safe_text(data.get("work_style")),
        safe_text(data.get("goal")),
        safe_text(data.get("answers")),
        safe_text(data.get("extra_answers")),
    ]
    if user:
        parts.extend([safe_text(user.get("interests")), safe_text(user.get("skills")), safe_text(user.get("goal"))])
    q_by_id = {safe_text(q.get("id")): q for q in sqr_recommendation_questions()}
    for answer in sqr_parse_recommendation_answers(data):
        q = q_by_id.get(safe_text(answer.get("id") or answer.get("question_id")))
        if q and sqr_rating(answer.get("value") or answer.get("answer"), 3) >= 3:
            parts.extend(q.get("keywords") or [])
            parts.append(q.get("specialization") or "")
    return " ".join([p for p in parts if p]).lower()

@app.route("/api/recommendations", methods=["POST"])
@student_required
def recommendations():
    data = get_json()
    user = request.current_user
    user_id = user.get("id") or user.get("user_id")
    quiz_scores, quiz_answers = sqr_score_recommendation_quiz(data)
    profile_text = sqr_recommendation_profile_text({"answers": quiz_answers}, user)

    specs = [normalize_specialization(row) for row in (query_db("SELECT * FROM specializations ORDER BY specialization_id DESC", fetchall=True) or [])]
    library = sqr_all_specialization_library()

    recommended_specs = []
    system_slugs = set()
    for spec in specs:
        spec_key = sqr_spec_key(spec)
        system_slugs.add(spec_key)
        target = f"{spec.get('name','')} {spec.get('description','')} {spec.get('skills','')} {spec.get('roadmap','')} {spec.get('career_paths','')}"
        text_score, matches = calculate_match_percentage(profile_text, target)
        quiz_score = safe_int(quiz_scores.get(spec_key), 0)
        lower_name = safe_text(spec.get("name")).lower()
        hint_score = 0
        for key, hints in SPECIALIZATION_HINTS.items():
            if sqr_slug(key) == spec_key or key in lower_name:
                hint_score = min(100, len([h for h in hints if h in profile_text]) * 18)
        final_score = round((quiz_score * 0.82) + (max(text_score, hint_score) * 0.18)) if quiz_answers else max(text_score, hint_score)
        final_score = max(0, min(100, final_score))
        recommended_specs.append({
            "id": spec.get("id"),
            "specialization_id": spec.get("specialization_id") or spec.get("id"),
            "name": spec.get("name"),
            "description": spec.get("description") or "",
            "match_percentage": final_score,
            "score": final_score,
            "quiz_score": quiz_score,
            "text_score": text_score,
            "matched_skills": matches,
            "reason": "Matched from your quiz answers and compared with specializations currently stored in SQR.",
            "in_system": True,
            "source": "database",
        })

    for key, info in library.items():
        if key in system_slugs:
            continue
        quiz_score = safe_int(quiz_scores.get(key), 0)
        target = " ".join([info.get("name", ""), info.get("description", ""), " ".join(info.get("keywords") or [])])
        text_score, matches = calculate_match_percentage(profile_text, target)
        final_score = round((quiz_score * 0.88) + (text_score * 0.12)) if quiz_answers else text_score
        if final_score <= 0 and quiz_answers:
            final_score = quiz_score
        final_score = max(0, min(100, final_score))
        recommended_specs.append({
            "id": None,
            "specialization_id": f"external-{key}",
            "name": info.get("name") or key.replace("_", " ").title(),
            "description": info.get("description") or "",
            "match_percentage": final_score,
            "score": final_score,
            "quiz_score": quiz_score,
            "text_score": text_score,
            "matched_skills": matches or (info.get("keywords") or [])[:5],
            "reason": "AI-ready suggestion from your quiz answers. It can appear even if this specialization is not stored in the database yet.",
            "roadmap": info.get("roadmap") or [],
            "in_system": False,
            "source": "ai_catalog",
        })

    recommended_specs.sort(key=lambda item: item.get("match_percentage", 0), reverse=True)
    deterministic_top = recommended_specs[:6]

    ai_fallback = {
        "summary": "Your specialization recommendation is based on the quiz answers only. Jobs are recommended separately on the Jobs page.",
        "recommended_specializations": deterministic_top,
        "roadmap": [
            "Start with the highest quiz-matched specialization.",
            "Review the suggested skills and roadmap.",
            "Open the Jobs page to get job recommendations separately.",
            "Use courses and quizzes to build measurable progress.",
        ],
    }

    ai_payload = ai_json(
        """
Return valid JSON only with these keys: summary, recommended_specializations, roadmap.
Recommend computer science specializations for an SQR student using the quiz answers only.
You may recommend a specialization even if it is not currently in the SQR database.
Do not recommend jobs. Jobs belong only on the Jobs page.
For each recommended_specializations item use keys: name, description, reason, match_percentage, roadmap, in_system.
Use the deterministic candidates and scores as guidance, but you may include one strong external specialization if the quiz indicates it.
"""
        + f"\nQuiz answers: {json.dumps(quiz_answers, ensure_ascii=False)}"
        + f"\nQuiz scores: {json.dumps(quiz_scores, ensure_ascii=False)}"
        + f"\nCurrent SQR database specializations: {json.dumps([{'name': s.get('name'), 'description': s.get('description'), 'specialization_id': s.get('specialization_id')} for s in specs], ensure_ascii=False)}"
        + f"\nAvailable external specialization library: {json.dumps(library, ensure_ascii=False)}"
        + f"\nDeterministic top candidates: {json.dumps(deterministic_top, ensure_ascii=False)}",
        ai_fallback
    )

    ai_specs_raw = ai_payload.get("recommended_specializations") if isinstance(ai_payload, dict) else None
    final_specs = []
    if isinstance(ai_specs_raw, list) and ai_specs_raw:
        db_by_name = {safe_text(s.get("name")).lower(): s for s in specs}
        for idx, item in enumerate(ai_specs_raw[:6]):
            if not isinstance(item, dict):
                continue
            name = safe_text(item.get("name")) or safe_text(deterministic_top[idx].get("name") if idx < len(deterministic_top) else "Specialization")
            db_match = db_by_name.get(name.lower())
            score = pct_value(item.get("match_percentage") or item.get("score") or (deterministic_top[idx].get("score") if idx < len(deterministic_top) else 0))
            final_specs.append({
                "id": db_match.get("id") if db_match else None,
                "specialization_id": (db_match.get("specialization_id") or db_match.get("id")) if db_match else f"external-{sqr_slug(name)}",
                "name": name,
                "description": safe_text(item.get("description")) or (db_match.get("description") if db_match else ""),
                "reason": safe_text(item.get("reason")) or "Recommended from your quiz answers.",
                "match_percentage": score,
                "score": score,
                "roadmap": item.get("roadmap") if isinstance(item.get("roadmap"), list) else [],
                "in_system": bool(db_match) if "in_system" not in item else bool(item.get("in_system")) and bool(db_match),
                "source": "ai" if ai_payload.get("ai_powered") else "dynamic_quiz",
            })
    if not final_specs:
        final_specs = deterministic_top

    result = {
        "summary": safe_text(ai_payload.get("summary")) or ai_fallback["summary"],
        "recommendation_basis": "quiz",
        "quiz_answers": quiz_answers,
        "quiz_scores": quiz_scores,
        "recommended_specializations": final_specs[:6],
        "roadmap": ai_payload.get("roadmap") if isinstance(ai_payload.get("roadmap"), list) else ai_fallback["roadmap"],
        "recommended_jobs": [],
        "jobs_location": "jobs.html",
        "ai_powered": bool(ai_payload.get("ai_powered")),
        "ai_provider": safe_text(ai_payload.get("ai_provider") or "local_dynamic_fallback"),
        "ai_error": safe_text(ai_payload.get("ai_error")),
    }

    try:
        if table_exists("recommendation_results"):
            query_db(
                "INSERT INTO recommendation_results (user_id,recommendation_json) VALUES (%s,%s)",
                (user_id, json.dumps(result)),
                commit=True
            )
        elif table_exists("recommendations") and result["recommended_specializations"]:
            top = result["recommended_specializations"][0]
            top_id = top.get("specialization_id") if safe_text(top.get("specialization_id")).isdigit() else None
            if top_id:
                query_db(
                    "INSERT INTO recommendations (user_id,specialization_id,match_score,explanation) VALUES (%s,%s,%s,%s)",
                    (user_id, top_id, top.get("match_percentage", 0), json.dumps(result)),
                    commit=True
                )
    except Exception as exc:
        print("RECOMMENDATION SAVE ERROR:", exc)

    return jsonify(result)


@app.route("/api/recommendations/analyze", methods=["POST"])
@student_required
def recommendations_analyze():
    return recommendations()


@app.route("/api/ats/check", methods=["POST"])
@student_required
def ats_check():
    # The checker now analyzes only the uploaded resume file. No pasted resume text and no job description are required.
    resume_file = request.files.get("resume_file") or request.files.get("resume") or request.files.get("file")
    if not resume_file:
        return jsonify({"error": "Please upload a PDF or DOCX resume"}), 400

    filename = secure_filename(resume_file.filename or "").lower()
    if not filename.endswith((".pdf", ".docx")):
        return jsonify({"error": "ATS checker accepts uploaded PDF or DOCX resumes only"}), 400

    resume_text = extract_resume_text(resume_file)
    if not resume_text.strip():
        return jsonify({"error": "Could not read resume text from this uploaded file"}), 400

    resume_lower = resume_text.lower()
    words = resume_text.split()
    detected_keywords = sorted({skill for skill in TECH_SKILLS if skill.lower() in resume_lower})[:25]
    required_sections = ["summary", "objective", "education", "skills", "experience", "project", "certification"]
    found_sections = [section for section in required_sections if re.search(r"\b" + re.escape(section) + r"s?\b", resume_lower)]
    missing_sections = [section.title() for section in ["summary/objective", "education", "skills", "projects or experience"] if not re.search(section.replace("/", "|"), resume_lower)]
    has_email = bool(re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", resume_text, re.I))
    has_phone = bool(re.search(r"(\+?\d[\d\s().-]{7,}\d)", resume_text))
    has_link = bool(re.search(r"linkedin|github|portfolio|https?://", resume_lower))

    structure_score = min(100, 45 + (len(found_sections) * 8) + (10 if has_email else 0) + (10 if has_phone else 0) + (5 if has_link else 0))
    keyword_score = min(100, 35 + len(detected_keywords) * 4)
    clarity_score = 88 if 180 <= len(words) <= 750 else (72 if len(words) >= 100 else 58)
    ats_score = round((structure_score * 0.4) + (keyword_score * 0.35) + (clarity_score * 0.25))

    fallback = {
        "ats_score": ats_score,
        "score": ats_score,
        "summary": "ATS resume check completed from the uploaded resume only. Configure GEMINI_API_KEY to receive a fully AI-written critique.",
        "detected_keywords": detected_keywords,
        "matched_keywords": detected_keywords,
        "missing_keywords": missing_sections,
        "strengths": [
            "Resume file was readable",
            "Contact information detected" if (has_email or has_phone) else "Resume content was extracted successfully",
            "Relevant technical keywords detected" if detected_keywords else "Ready for keyword improvement",
        ],
        "weaknesses": missing_sections or ["Add more measurable results and role-specific wording"],
        "improvements": [
            "Add phone number and email in the header if either is missing.",
            "Use clear sections such as Objective, Education, Internship/Experience, Skills, Certification, and Projects.",
            "Add measurable project or internship results such as accuracy, users, time saved, or tools used.",
            "Keep bullets concise and start them with action verbs.",
        ],
        "section_scores": {
            "structure": structure_score,
            "keywords": keyword_score,
            "clarity": clarity_score,
        },
        "source": "uploaded_resume_only",
        "ai_powered": False,
        "ai_provider": "local_dynamic_fallback",
    }

    prompt = f"""
You are an expert ATS resume reviewer inside the SQR website.
Analyze ONLY the uploaded resume text below. Do not use any job description, target role, or external user data.
Return valid JSON only with these exact keys:
ats_score, score, summary, detected_keywords, matched_keywords, missing_keywords, strengths, weaknesses, improvements, section_scores, source.
Rules:
- Base the review only on the uploaded resume text.
- Do not invent education, experience, projects, certificates, dates, grades, names, phone numbers, or emails.
- Score the resume for ATS readability, section structure, contact information, skill keywords, clarity, action verbs, and measurable achievements.
- missing_keywords should list missing resume elements or missing skill areas found from the resume itself, not from a job description.
- improvements must be specific, practical, and based on the uploaded resume.
- source must be "uploaded_resume_only".

Uploaded resume text:
{safe_text(resume_text)[:9000]}
"""
    result = ai_json(prompt, fallback)
    if not isinstance(result, dict):
        result = fallback

    result["source"] = "uploaded_resume_only"
    result["ats_score"] = safe_int(result.get("ats_score") or result.get("score"), ats_score)
    result["score"] = result["ats_score"]
    if "detected_keywords" not in result:
        result["detected_keywords"] = result.get("matched_keywords") or detected_keywords
    if "matched_keywords" not in result:
        result["matched_keywords"] = result.get("detected_keywords") or detected_keywords
    if "missing_keywords" not in result:
        result["missing_keywords"] = missing_sections

    suggestions = safe_text(result.get("summary")) or json.dumps(result.get("improvements", []))
    query_db(
        """
        INSERT INTO ats_results (user_id,resume_text,target_job,ats_score,matched_keywords,missing_keywords,suggestions)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (
            request.current_user["id"],
            resume_text[:60000],
            "Uploaded resume only",
            result["ats_score"],
            json.dumps(result.get("matched_keywords", []), ensure_ascii=False),
            json.dumps(result.get("missing_keywords", []), ensure_ascii=False),
            suggestions,
        ),
        commit=True
    )
    return jsonify(result)




def sqr_compact_list(text, limit=10):
    parts = re.split(r"[,\n;|]+", safe_text(text))
    clean_parts = []
    for part in parts:
        item = safe_text(part)
        if item and item.lower() not in [x.lower() for x in clean_parts]:
            clean_parts.append(item)
    return clean_parts[:limit]


def sqr_local_enhanced_summary(target_role, original_summary, technical_skills, soft_skills, education, experience, projects, certifications):
    role = safe_text(target_role) or "technology role"
    tech = sqr_compact_list(technical_skills, 6)
    soft = sqr_compact_list(soft_skills, 3)
    tech_text = ", ".join(tech) if tech else "relevant technical tools"
    soft_text = ", ".join(soft) if soft else "communication and problem solving"
    education_text = safe_text(education)
    base = f"{role} professional with skills in {tech_text} and strengths in {soft_text}."
    if projects:
        base += " Experienced in building practical projects and presenting technical work clearly."
    elif experience:
        base += " Experienced in applying technical knowledge in practical environments."
    else:
        base += " Prepared to apply academic knowledge, projects, and continuous learning to real technical work."
    if education_text:
        base += f" Education background includes {education_text[:140]}."
    if original_summary:
        base += " " + safe_text(original_summary)[:180]
    base = re.sub(r"\b(candidate|ATS-friendly career readiness)\b", "", base, flags=re.I)
    base = re.sub(r"\s+", " ", base).strip()
    return base[:650]


def extract_resume_text_from_upload(file):
    if not file or not file.filename:
        return ""
    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        raw = file.read()
        file.seek(0)
        if ext == "pdf" and PdfReader:
            reader = PdfReader(BytesIO(raw))
            return "\n".join([(page.extract_text() or "") for page in reader.pages]).strip()
        if ext == "docx" and Document:
            doc = Document(BytesIO(raw))
            return "\n".join([p.text for p in doc.paragraphs]).strip()
        if ext == "txt":
            return raw.decode("utf-8", errors="ignore").strip()
    except Exception as exc:
        print("ATS RESUME READ ERROR:", exc)
        return ""
    return ""


def local_dynamic_summary(name, target_role, technical_skills, soft_skills, resume_text):
    name = safe_text(name) or "The candidate"
    target_role = safe_text(target_role) or "a technology role"
    tech = sqr_compact_list(technical_skills, 10)
    soft = sqr_compact_list(soft_skills, 6)
    detected = [skill for skill in TECH_SKILLS if skill.lower() in safe_text(resume_text).lower()]
    tech_text = ", ".join(tech or detected[:8]) or "software, data, and problem-solving fundamentals"
    soft_text = ", ".join(soft) or "communication, teamwork, ownership, and continuous learning"
    resume_hint = ""
    if re.search(r"project|built|developed|implemented|designed", safe_text(resume_text), re.I):
        resume_hint = " The profile shows practical project work and the ability to turn technical knowledge into deliverables."
    elif safe_text(resume_text):
        resume_hint = " The submitted resume text was used to keep this summary specific to the candidate."
    return (
        f"{name} is preparing for {target_role} roles with a foundation in {tech_text}. "
        f"They bring strengths in {soft_text}, with focus on clear problem solving and reliable execution."
        f"{resume_hint} This summary can become stronger by adding measurable achievements, tools used, project outcomes, and role-specific keywords."
    ).strip()


SOFT_SKILL_KEYWORDS = [
    "communication", "teamwork", "problem solving", "problem-solving", "leadership",
    "time management", "adaptability", "critical thinking", "creativity", "collaboration",
    "presentation", "continuous learning", "attention to detail", "organization",
    "work ethic", "ownership", "reliability", "active listening", "decision making"
]


def sqr_clean_link(value):
    link = safe_text(value)
    if not link:
        return ""
    if link.lower().startswith(("http://", "https://")):
        return link
    if link.lower().startswith(("linkedin.com", "www.linkedin.com", "github.com", "www.github.com")):
        return "https://" + link
    return link


def sqr_is_soft_skill(skill):
    value = safe_text(skill).lower()
    value = re.sub(r"\s+", " ", value)
    return any(keyword in value for keyword in SOFT_SKILL_KEYWORDS)


def sqr_split_resume_skills(technical_text="", soft_text="", resume_text=""):
    entered_technical = sqr_compact_list(technical_text, 50)
    entered_soft = sqr_compact_list(soft_text, 50)

    technical = []
    soft = []
    for item in entered_technical:
        if sqr_is_soft_skill(item):
            soft.append(item)
        else:
            technical.append(item)
    for item in entered_soft:
        if item and item.lower() not in [x.lower() for x in soft]:
            soft.append(item)

    resume_lower = safe_text(resume_text).lower()
    if not technical and resume_lower:
        for skill in TECH_SKILLS:
            if skill.lower() in resume_lower and not sqr_is_soft_skill(skill):
                technical.append(skill)
    if not soft:
        for skill in SOFT_SKILL_KEYWORDS:
            if skill.lower() in resume_lower:
                soft.append(skill.title())

    if not soft:
        soft = ["Communication", "Teamwork", "Problem Solving", "Continuous Learning"]

    # Remove accidental duplicates and soft skills from technical skills.
    final_technical = []
    for item in technical:
        clean = safe_text(item)
        if clean and not sqr_is_soft_skill(clean) and clean.lower() not in [x.lower() for x in final_technical]:
            final_technical.append(clean)

    final_soft = []
    for item in soft:
        clean = safe_text(item)
        if clean and clean.lower() not in [x.lower() for x in final_soft]:
            final_soft.append(clean)

    return final_technical[:14], final_soft[:10]


def build_dynamic_resume_payload(data, resume_text):
    data = data or {}
    name = safe_text(data.get("name")) or "Candidate"
    target_role = safe_text(data.get("target_role") or data.get("role") or data.get("target_job")) or "Technology Role"
    technical_skills_text = safe_text(data.get("technical_skills") or data.get("skills"))
    soft_skills_text = safe_text(data.get("soft_skills"))
    education = safe_text(data.get("education"))
    experience = safe_text(data.get("experience"))
    projects = safe_text(data.get("projects"))
    certifications = safe_text(data.get("certifications") or data.get("certificates"))
    original_summary = safe_text(data.get("summary") or data.get("original_summary"))
    linkedin = sqr_clean_link(data.get("linkedin") or data.get("linkedin_url"))
    github = sqr_clean_link(data.get("github") or data.get("github_url"))
    phone = safe_text(data.get("phone") or data.get("phone_number") or data.get("mobile"))
    email = safe_text(data.get("email") or data.get("email_address"))
    location = safe_text(data.get("location") or data.get("address") or data.get("city"))
    portfolio = sqr_clean_link(data.get("portfolio") or data.get("portfolio_url"))
    contact_line = _build_contact_line_from_data({
        "phone": phone,
        "location": location,
        "email": email,
        "linkedin": linkedin,
        "portfolio": portfolio or github,
    })

    technical_list, soft_list = sqr_split_resume_skills(technical_skills_text, soft_skills_text, resume_text)

    summary = sqr_local_enhanced_summary(
        target_role,
        original_summary,
        ", ".join(technical_list),
        ", ".join(soft_list),
        education,
        experience,
        projects,
        certifications,
    )
    if not summary:
        summary = local_dynamic_summary(name, target_role, ", ".join(technical_list), ", ".join(soft_list), resume_text)

    sections = []
    sections.append(name.upper())
    sections.append(target_role)
    if contact_line:
        sections.append(contact_line)
    sections.extend(["", "PROFESSIONAL SUMMARY", summary])
    sections.extend(["", "TECHNICAL SKILLS", ", ".join(technical_list) if technical_list else "Add role-specific tools, languages, platforms, and frameworks."])
    sections.extend(["", "SOFT SKILLS", ", ".join(soft_list)])
    if projects:
        sections.extend(["", "PROJECTS", projects])
    if experience:
        sections.extend(["", "EXPERIENCE", experience])
    if education:
        sections.extend(["", "EDUCATION", education])
    if certifications:
        sections.extend(["", "CERTIFICATIONS", certifications])

    return {
        "headline": target_role,
        "summary": summary,
        "enhanced_summary": summary,
        "technical_skills": technical_list,
        "soft_skills": soft_list,
        "phone": phone,
        "email": email,
        "location": location,
        "address": location,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "contact_line": contact_line,
        "projects": projects,
        "experience": experience,
        "education": education,
        "certifications": certifications,
        "full_resume": "\n".join(sections).strip(),
        "improvements": [
            "Add numbers to achievements, such as percentage improvements, users served, or project size.",
            "Add exact tools and technologies beside each project.",
            "Mirror the most important keywords from the target job description.",
            "Keep sections clear: Summary, Technical Skills, Soft Skills, Projects, Experience, Education, Certifications."
        ],
        "missing_information": [
            item for item, value in {
                "phone": phone,
                "email": email,
                "linkedin": linkedin,
                "github": github,
                "projects": projects,
                "experience": experience,
                "education": education,
                "certifications": certifications,
            }.items() if not value
        ],
        "ai_powered": False,
        "ai_provider": "local_dynamic_fallback",
    }


def generate_ai_resume_payload(data, resume_text):
    data = data or {}
    fallback = build_dynamic_resume_payload(data, resume_text)

    target_role = safe_text(data.get("target_role") or data.get("role") or data.get("target_job") or fallback.get("headline"))
    original_summary = safe_text(data.get("summary") or data.get("original_summary"))
    linkedin = fallback.get("linkedin", "")
    github = fallback.get("github", "")
    phone = fallback.get("phone", "")
    email = fallback.get("email", "")
    location = fallback.get("location", "")
    portfolio = fallback.get("portfolio", "")

    prompt = f"""
You are an expert resume writer inside the SQR website. Rewrite and improve the user's resume details like a real AI assistant.

Return valid JSON only with these exact keys:
headline, summary, enhanced_summary, technical_skills, soft_skills, phone, email, location, linkedin, github, portfolio, projects, experience, education, certifications, full_resume, improvements, missing_information.

Rules:
- Use only the information provided in the form and uploaded/pasted resume text.
- Do not invent companies, dates, GPA, certificates, degrees, projects, jobs, links, or years.
- Do not use fixed generic text.
- Do not say "ATS-friendly career readiness".
- enhanced_summary must be 3 to 5 specific sentences based on the user's actual summary, skills, target job, projects, education, and resume text.
- Separate technical_skills from soft_skills.
- technical_skills must include only programming languages, tools, frameworks, databases, platforms, methods, and technologies.
- soft_skills must include only personal/workplace strengths such as communication, teamwork, problem solving, leadership, time management, and continuous learning.
- Put phone only in phone, email only in email, LinkedIn only in linkedin, GitHub only in github, and portfolio only in portfolio if provided.
- Do not put soft skills inside technical_skills.
- Improve wording, action verbs, clarity, and ATS keyword alignment for the target role.
- If information is missing, list it in missing_information instead of inventing it.
- full_resume must start with: NAME, target role, then one contact row formatted as Phone | Location | Email | LinkedIn | Portfolio/Github when those fields exist. After that, use clean ATS-readable headings when data exists: PROFESSIONAL SUMMARY, TECHNICAL SKILLS, SOFT SKILLS, PROJECTS, EXPERIENCE, EDUCATION, CERTIFICATIONS.
- If a section has no user data, do not invent bullet points for it.

Target role: {target_role}
Phone: {phone}
Email: {email}
Location: {location}
LinkedIn: {linkedin}
GitHub: {github}
Portfolio: {portfolio}
User's original summary: {original_summary}
User form data JSON:
{json.dumps(data, ensure_ascii=False)}

Uploaded or pasted resume text:
{safe_text(resume_text)[:9000]}
"""

    payload = ai_json(prompt, fallback)
    if not isinstance(payload, dict):
        payload = fallback

    for key, value in fallback.items():
        if key not in payload or payload.get(key) in [None, "", []]:
            payload[key] = value

    for key in ("technical_skills", "soft_skills", "improvements", "missing_information"):
        value = payload.get(key)
        if isinstance(value, str):
            payload[key] = sqr_compact_list(value, 20)
        elif isinstance(value, list):
            payload[key] = [safe_text(x) for x in value if safe_text(x)][:20]
        else:
            payload[key] = fallback.get(key, [])

    # Clean any skill mixing after AI returns.
    technical_clean, soft_clean = sqr_split_resume_skills(", ".join(payload.get("technical_skills") or []), ", ".join(payload.get("soft_skills") or []), resume_text)
    payload["technical_skills"] = technical_clean or fallback.get("technical_skills", [])
    payload["soft_skills"] = soft_clean or fallback.get("soft_skills", [])

    payload["phone"] = safe_text(payload.get("phone") or fallback.get("phone"))
    payload["email"] = safe_text(payload.get("email") or fallback.get("email"))
    payload["location"] = safe_text(payload.get("location") or payload.get("address") or fallback.get("location"))
    payload["address"] = payload["location"]
    payload["linkedin"] = sqr_clean_link(payload.get("linkedin") or fallback.get("linkedin"))
    payload["github"] = sqr_clean_link(payload.get("github") or fallback.get("github"))
    payload["portfolio"] = sqr_clean_link(payload.get("portfolio") or fallback.get("portfolio"))
    payload["contact_line"] = _build_contact_line_from_data({
        "phone": payload["phone"],
        "location": payload["location"],
        "email": payload["email"],
        "linkedin": payload["linkedin"],
        "portfolio": payload["portfolio"] or payload["github"],
    })
    payload["summary"] = safe_text(payload.get("summary") or payload.get("enhanced_summary") or fallback["summary"])
    payload["enhanced_summary"] = safe_text(payload.get("enhanced_summary") or payload.get("summary") or fallback["summary"])
    payload["headline"] = safe_text(payload.get("headline") or fallback.get("headline") or target_role)

    sections = [
        safe_text(data.get("name") or "Candidate").upper(),
        payload["headline"],
    ]
    if payload.get("contact_line"):
        sections.append(payload["contact_line"])
    sections.extend([
        "", "PROFESSIONAL SUMMARY", payload["enhanced_summary"],
        "", "TECHNICAL SKILLS", ", ".join(payload.get("technical_skills") or fallback.get("technical_skills") or []),
        "", "SOFT SKILLS", ", ".join(payload.get("soft_skills") or fallback.get("soft_skills") or []),
    ])
    for label, key in (("PROJECTS", "projects"), ("EXPERIENCE", "experience"), ("EDUCATION", "education"), ("CERTIFICATIONS", "certifications")):
        value = safe_text(payload.get(key) or fallback.get(key))
        if value:
            sections.extend(["", label, value])
    payload["full_resume"] = "\n".join([x for x in sections if x is not None]).strip()

    provider = safe_text(payload.get("ai_provider")).lower()
    payload["ai_powered"] = provider in {"gemini", "google_gemini", "google-genai", "xai", "grok", "openai"} or bool(payload.get("ai_powered"))
    payload["ai_provider"] = provider if payload["ai_powered"] else "local_dynamic_fallback"

    bad_phrases = ["ATS-friendly career readiness", "fixed text", "lorem ipsum"]
    if any(bad.lower() in json.dumps(payload, ensure_ascii=False).lower() for bad in bad_phrases):
        fallback["ai_provider"] = "local_dynamic_fallback"
        fallback["ai_powered"] = False
        return fallback

    return payload

def generate_ai_enhanced_summary(name, target_role, technical_skills, soft_skills, resume_text):
    payload = generate_ai_resume_payload({
        "name": name,
        "target_role": target_role,
        "technical_skills": technical_skills,
        "soft_skills": soft_skills,
    }, resume_text)
    return payload.get("enhanced_summary") or payload.get("summary") or local_dynamic_summary(name, target_role, technical_skills, soft_skills, resume_text)

@app.route("/api/debug/ai", methods=["GET"])
@app.route("/api/debug/gemini", methods=["GET"])
@app.route("/api/debug/xai", methods=["GET"])
def debug_gemini():
    try:
        payload = ai_json(
            'Return this exact JSON only: {"test":"AI works"}',
            {"test": "fallback"}
        )

        return jsonify({
            "ok": bool(payload.get("ai_powered")),
            "provider_mode": AI_PROVIDER,
            "provider_used": payload.get("ai_provider"),
            "model_used": payload.get("ai_model", ""),
            "gemini_configured": bool(GEMINI_API_KEY),
            "xai_configured": bool(XAI_API_KEY),
            "openai_configured": bool(OPENAI_API_KEY),
            "message": payload,
            "error": payload.get("ai_error", "")
        })

    except Exception as e:
        print("AI DEBUG ERROR:", repr(e))
        return jsonify({
            "ok": False,
            "error": mask_gemini_error(str(e))
        }), 500
        
@app.route("/api/ats/generate", methods=["POST"])

@login_required
def ats_generate():
    if request.content_type and "multipart/form-data" in request.content_type:
        data = dict(request.form)
        resume_file = request.files.get("resume") or request.files.get("resume_file") or request.files.get("file")
        resume_text = extract_resume_text_from_upload(resume_file)
    else:
        data = get_json()
        resume_text = safe_text(data.get("resume_text") or data.get("resume"))

    payload = generate_ai_resume_payload(data, resume_text)
    payload.update({
        "success": True,
        "resume_text_found": bool(safe_text(resume_text)),
        "target_role": safe_text(data.get("target_role") or data.get("role") or data.get("target_job")),
    })
    return jsonify(payload)


@app.route("/api/ats/export/pdf", methods=["POST"])
@login_required
def export_pdf():
    data = get_json()
    text = safe_text(data.get("resume") or data.get("text"))
    if not text:
        return jsonify({"error": "Resume text is required"}), 400
    pdf = build_resume_pdf(text)
    if not pdf:
        return jsonify({"error": "PDF export library is not installed"}), 500
    return send_file(pdf, mimetype="application/pdf", as_attachment=True, download_name="sqr_resume.pdf")


@app.route("/api/ats/export/docx", methods=["POST"])
@login_required
def export_docx():
    if not Document:
        return jsonify({"error": "DOCX export library is not installed"}), 500
    data = get_json()
    text = safe_text(data.get("resume") or data.get("text"))
    if not text:
        return jsonify({"error": "Resume text is required"}), 400
    docx_file = build_resume_docx(text)
    if not docx_file:
        return jsonify({"error": "DOCX export could not be generated"}), 500
    return send_file(docx_file, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="sqr_resume.docx")


@app.route("/api/admin/stats")
@admin_required
def admin_stats():
    def count(table):
        if not table_exists(table):
            return 0
        row = query_db(f"SELECT COUNT(*) AS total FROM `{table}`", fetchone=True) or {"total": 0}
        return safe_int(row.get("total"), 0)
    return jsonify({
        "users": count("users"),
        "specializations": count(admin_spec_table()),
        "courses": count("courses"),
        "quizzes": count("quizzes"),
        "jobs": count("jobs"),
        "certificates": count("certificates"),
        "ats_results": count("ats_results"),
    })



@app.route("/api/admin/users")
@admin_required
def admin_users():
    rows = query_db("SELECT * FROM users ORDER BY id DESC", fetchall=True) or []
    return jsonify({"users": [clean_user(row) for row in rows]})




@app.route("/api/admin/users/<int:user_id>/ban", methods=["PUT", "POST"])
@admin_required
def ban_user(user_id):
    if user_id == request.current_user["id"]:
        return jsonify({"error": "You cannot ban yourself"}), 400
    if column_exists("users", "banned") and column_exists("users", "is_banned"):
        exec_db("UPDATE users SET banned=1, is_banned=1 WHERE id=%s", (user_id,))
    elif column_exists("users", "banned"):
        exec_db("UPDATE users SET banned=1 WHERE id=%s", (user_id,))
    else:
        exec_db("UPDATE users SET is_banned=1 WHERE id=%s", (user_id,))
    return jsonify({"message": "User banned"})




@app.route("/api/admin/users/<int:user_id>/unban", methods=["PUT", "POST"])
@admin_required
def unban_user(user_id):
    if column_exists("users", "banned") and column_exists("users", "is_banned"):
        exec_db("UPDATE users SET banned=0, is_banned=0 WHERE id=%s", (user_id,))
    elif column_exists("users", "banned"):
        exec_db("UPDATE users SET banned=0 WHERE id=%s", (user_id,))
    else:
        exec_db("UPDATE users SET is_banned=0 WHERE id=%s", (user_id,))
    return jsonify({"message": "User unbanned"})




@app.route("/api/admin/users/<int:user_id>/role", methods=["PUT", "POST"])
@admin_required
def change_user_role(user_id):
    data = get_json()
    role = safe_text(data.get("role")).lower()
    if role not in ["student", "admin"]:
        return jsonify({"error": "Role must be student or admin"}), 400
    current_mode = "admin" if role == "admin" else "student"
    exec_db("UPDATE users SET role=%s,current_mode=%s WHERE id=%s", (role, current_mode, user_id))
    if role == "admin":
        try:
            exec_db("INSERT IGNORE INTO admins (user_id,admin_level) VALUES (%s,'manager')", (user_id,))
        except Exception as exc:
            print("ADMIN INSERT SKIPPED:", exc)
    else:
        try:
            exec_db("DELETE FROM admins WHERE user_id=%s", (user_id,))
        except Exception:
            pass
    return jsonify({"message": "User role updated"})



@app.route("/api/mode", methods=["PUT", "POST"])
@login_required
def switch_mode():
    data = get_json()
    mode = safe_text(data.get("mode")).lower()
    if request.current_user.get("role") != "admin":
        return jsonify({"error": "Only admins can switch mode"}), 403
    if mode not in ["student", "admin"]:
        return jsonify({"error": "Mode must be student or admin"}), 400
    exec_db("UPDATE users SET current_mode=%s WHERE id=%s", (mode, request.current_user["id"]))
    user = query_db("SELECT * FROM users WHERE id=%s", (request.current_user["id"],), fetchone=True)
    return jsonify({"message": "Mode updated", "token": generate_token(user), "user": clean_user(user)})



# =====================================================
# SQR LONG DYNAMIC PATCH LAYER
# Non-destructive extension: adds dynamic bootstrap, diagnostics,
# recommendation question bank, and rich profile/admin summary routes.
# =====================================================

SQR_RECOMMENDATION_QUESTION_BANK = [
    {
        "id": "cybersecurity_interest",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "interest",
        "question": "How interested are you in Cybersecurity tasks such as linux, networking, and security?",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 5
    },
    {
        "id": "cybersecurity_skill",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "skill",
        "question": "How confident are you with linux, networking, or security for a Cybersecurity path?",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 5
    },
    {
        "id": "cybersecurity_work_style",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "work_style",
        "question": "Do you prefer work that involves linux, networking, and practical problem solving for Cybersecurity?",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 5
    },
    {
        "id": "cybersecurity_project",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using linux, networking, and security?",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 5
    },
    {
        "id": "cybersecurity_career",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "career",
        "question": "Would you consider a future job connected to Cybersecurity and skills like linux and networking?",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 5
    },
    {
        "id": "digital_forensics_interest",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "interest",
        "question": "How interested are you in Digital Forensics tasks such as forensics, evidence, and malware?",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 5
    },
    {
        "id": "digital_forensics_skill",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "skill",
        "question": "How confident are you with forensics, evidence, or malware for a Digital Forensics path?",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 5
    },
    {
        "id": "digital_forensics_work_style",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "work_style",
        "question": "Do you prefer work that involves forensics, evidence, and practical problem solving for Digital Forensics?",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 5
    },
    {
        "id": "digital_forensics_project",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using forensics, evidence, and malware?",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 5
    },
    {
        "id": "digital_forensics_career",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "career",
        "question": "Would you consider a future job connected to Digital Forensics and skills like forensics and evidence?",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 5
    },
    {
        "id": "software_engineering_interest",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "interest",
        "question": "How interested are you in Software Engineering tasks such as java, python, and testing?",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 5
    },
    {
        "id": "software_engineering_skill",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "skill",
        "question": "How confident are you with java, python, or testing for a Software Engineering path?",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 5
    },
    {
        "id": "software_engineering_work_style",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "work_style",
        "question": "Do you prefer work that involves java, python, and practical problem solving for Software Engineering?",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 5
    },
    {
        "id": "software_engineering_project",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using java, python, and testing?",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 5
    },
    {
        "id": "software_engineering_career",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "career",
        "question": "Would you consider a future job connected to Software Engineering and skills like java and python?",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 5
    },
    {
        "id": "web_development_interest",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "interest",
        "question": "How interested are you in Web Development tasks such as html, css, and javascript?",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 5
    },
    {
        "id": "web_development_skill",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "skill",
        "question": "How confident are you with html, css, or javascript for a Web Development path?",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 5
    },
    {
        "id": "web_development_work_style",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "work_style",
        "question": "Do you prefer work that involves html, css, and practical problem solving for Web Development?",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 5
    },
    {
        "id": "web_development_project",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using html, css, and javascript?",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 5
    },
    {
        "id": "web_development_career",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "career",
        "question": "Would you consider a future job connected to Web Development and skills like html and css?",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 5
    },
    {
        "id": "data_science_interest",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "interest",
        "question": "How interested are you in Data Science tasks such as python, sql, and statistics?",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 5
    },
    {
        "id": "data_science_skill",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "skill",
        "question": "How confident are you with python, sql, or statistics for a Data Science path?",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 5
    },
    {
        "id": "data_science_work_style",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "work_style",
        "question": "Do you prefer work that involves python, sql, and practical problem solving for Data Science?",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 5
    },
    {
        "id": "data_science_project",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using python, sql, and statistics?",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 5
    },
    {
        "id": "data_science_career",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "career",
        "question": "Would you consider a future job connected to Data Science and skills like python and sql?",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 5
    },
    {
        "id": "ai_ml_interest",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "interest",
        "question": "How interested are you in AI and Machine Learning tasks such as python, machine learning, and deep learning?",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 5
    },
    {
        "id": "ai_ml_skill",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "skill",
        "question": "How confident are you with python, machine learning, or deep learning for a AI and Machine Learning path?",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 5
    },
    {
        "id": "ai_ml_work_style",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "work_style",
        "question": "Do you prefer work that involves python, machine learning, and practical problem solving for AI and Machine Learning?",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 5
    },
    {
        "id": "ai_ml_project",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using python, machine learning, and deep learning?",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 5
    },
    {
        "id": "ai_ml_career",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "career",
        "question": "Would you consider a future job connected to AI and Machine Learning and skills like python and machine learning?",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 5
    },
    {
        "id": "cloud_devops_interest",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "interest",
        "question": "How interested are you in Cloud and DevOps tasks such as aws, docker, and linux?",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 5
    },
    {
        "id": "cloud_devops_skill",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "skill",
        "question": "How confident are you with aws, docker, or linux for a Cloud and DevOps path?",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 5
    },
    {
        "id": "cloud_devops_work_style",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "work_style",
        "question": "Do you prefer work that involves aws, docker, and practical problem solving for Cloud and DevOps?",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 5
    },
    {
        "id": "cloud_devops_project",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using aws, docker, and linux?",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 5
    },
    {
        "id": "cloud_devops_career",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "career",
        "question": "Would you consider a future job connected to Cloud and DevOps and skills like aws and docker?",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 5
    },
    {
        "id": "database_interest",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "interest",
        "question": "How interested are you in Database Systems tasks such as sql, mysql, and postgresql?",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 5
    },
    {
        "id": "database_skill",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "skill",
        "question": "How confident are you with sql, mysql, or postgresql for a Database Systems path?",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 5
    },
    {
        "id": "database_work_style",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "work_style",
        "question": "Do you prefer work that involves sql, mysql, and practical problem solving for Database Systems?",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 5
    },
    {
        "id": "database_project",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using sql, mysql, and postgresql?",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 5
    },
    {
        "id": "database_career",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "career",
        "question": "Would you consider a future job connected to Database Systems and skills like sql and mysql?",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 5
    },
    {
        "id": "mobile_interest",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "interest",
        "question": "How interested are you in Mobile App Development tasks such as flutter, swift, and kotlin?",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 5
    },
    {
        "id": "mobile_skill",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "skill",
        "question": "How confident are you with flutter, swift, or kotlin for a Mobile App Development path?",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 5
    },
    {
        "id": "mobile_work_style",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "work_style",
        "question": "Do you prefer work that involves flutter, swift, and practical problem solving for Mobile App Development?",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 5
    },
    {
        "id": "mobile_project",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using flutter, swift, and kotlin?",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 5
    },
    {
        "id": "mobile_career",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "career",
        "question": "Would you consider a future job connected to Mobile App Development and skills like flutter and swift?",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 5
    },
    {
        "id": "networks_interest",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "interest",
        "question": "How interested are you in Computer Networks tasks such as routing, tcp/ip, and switching?",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 5
    },
    {
        "id": "networks_skill",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "skill",
        "question": "How confident are you with routing, tcp/ip, or switching for a Computer Networks path?",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 5
    },
    {
        "id": "networks_work_style",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "work_style",
        "question": "Do you prefer work that involves routing, tcp/ip, and practical problem solving for Computer Networks?",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 5
    },
    {
        "id": "networks_project",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using routing, tcp/ip, and switching?",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 5
    },
    {
        "id": "networks_career",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "career",
        "question": "Would you consider a future job connected to Computer Networks and skills like routing and tcp/ip?",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 5
    },
    {
        "id": "uiux_interest",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "interest",
        "question": "How interested are you in UI/UX Engineering tasks such as design, accessibility, and prototyping?",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 5
    },
    {
        "id": "uiux_skill",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "skill",
        "question": "How confident are you with design, accessibility, or prototyping for a UI/UX Engineering path?",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 5
    },
    {
        "id": "uiux_work_style",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "work_style",
        "question": "Do you prefer work that involves design, accessibility, and practical problem solving for UI/UX Engineering?",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 5
    },
    {
        "id": "uiux_project",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using design, accessibility, and prototyping?",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 5
    },
    {
        "id": "uiux_career",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "career",
        "question": "Would you consider a future job connected to UI/UX Engineering and skills like design and accessibility?",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 5
    },
    {
        "id": "game_interest",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "interest",
        "question": "How interested are you in Game Development tasks such as c++, unity, and graphics?",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 5
    },
    {
        "id": "game_skill",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "skill",
        "question": "How confident are you with c++, unity, or graphics for a Game Development path?",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 5
    },
    {
        "id": "game_work_style",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "work_style",
        "question": "Do you prefer work that involves c++, unity, and practical problem solving for Game Development?",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 5
    },
    {
        "id": "game_project",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "project",
        "question": "Would you enjoy building portfolio projects using c++, unity, and graphics?",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 5
    },
    {
        "id": "game_career",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "career",
        "question": "Would you consider a future job connected to Game Development and skills like c++ and unity?",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 5
    },
    {
        "id": "cybersecurity_scenario_1",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "scenario",
        "question": "For Cybersecurity, choose how much you like a beginner scenario that uses networking and security.",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 4
    },
    {
        "id": "cybersecurity_scenario_2",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "scenario",
        "question": "For Cybersecurity, choose how much you like a intermediate scenario that uses security and incident response.",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 4
    },
    {
        "id": "cybersecurity_scenario_3",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "scenario",
        "question": "For Cybersecurity, choose how much you like a advanced scenario that uses incident response and linux.",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 4
    },
    {
        "id": "cybersecurity_scenario_4",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "scenario",
        "question": "For Cybersecurity, choose how much you like a project scenario that uses linux and networking.",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 4
    },
    {
        "id": "cybersecurity_scenario_5",
        "specialization_key": "cybersecurity",
        "specialization": "Cybersecurity",
        "dimension": "scenario",
        "question": "For Cybersecurity, choose how much you like a career scenario that uses networking and security.",
        "keywords": [
            "linux",
            "networking",
            "security",
            "incident response"
        ],
        "weight": 4
    },
    {
        "id": "digital_forensics_scenario_1",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "scenario",
        "question": "For Digital Forensics, choose how much you like a beginner scenario that uses evidence and malware.",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 4
    },
    {
        "id": "digital_forensics_scenario_2",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "scenario",
        "question": "For Digital Forensics, choose how much you like a intermediate scenario that uses malware and investigation.",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 4
    },
    {
        "id": "digital_forensics_scenario_3",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "scenario",
        "question": "For Digital Forensics, choose how much you like a advanced scenario that uses investigation and forensics.",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 4
    },
    {
        "id": "digital_forensics_scenario_4",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "scenario",
        "question": "For Digital Forensics, choose how much you like a project scenario that uses forensics and evidence.",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 4
    },
    {
        "id": "digital_forensics_scenario_5",
        "specialization_key": "digital_forensics",
        "specialization": "Digital Forensics",
        "dimension": "scenario",
        "question": "For Digital Forensics, choose how much you like a career scenario that uses evidence and malware.",
        "keywords": [
            "forensics",
            "evidence",
            "malware",
            "investigation"
        ],
        "weight": 4
    },
    {
        "id": "software_engineering_scenario_1",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "scenario",
        "question": "For Software Engineering, choose how much you like a beginner scenario that uses python and testing.",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 4
    },
    {
        "id": "software_engineering_scenario_2",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "scenario",
        "question": "For Software Engineering, choose how much you like a intermediate scenario that uses testing and architecture.",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 4
    },
    {
        "id": "software_engineering_scenario_3",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "scenario",
        "question": "For Software Engineering, choose how much you like a advanced scenario that uses architecture and java.",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 4
    },
    {
        "id": "software_engineering_scenario_4",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "scenario",
        "question": "For Software Engineering, choose how much you like a project scenario that uses java and python.",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 4
    },
    {
        "id": "software_engineering_scenario_5",
        "specialization_key": "software_engineering",
        "specialization": "Software Engineering",
        "dimension": "scenario",
        "question": "For Software Engineering, choose how much you like a career scenario that uses python and testing.",
        "keywords": [
            "java",
            "python",
            "testing",
            "architecture"
        ],
        "weight": 4
    },
    {
        "id": "web_development_scenario_1",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "scenario",
        "question": "For Web Development, choose how much you like a beginner scenario that uses css and javascript.",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 4
    },
    {
        "id": "web_development_scenario_2",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "scenario",
        "question": "For Web Development, choose how much you like a intermediate scenario that uses javascript and react.",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 4
    },
    {
        "id": "web_development_scenario_3",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "scenario",
        "question": "For Web Development, choose how much you like a advanced scenario that uses react and html.",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 4
    },
    {
        "id": "web_development_scenario_4",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "scenario",
        "question": "For Web Development, choose how much you like a project scenario that uses html and css.",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 4
    },
    {
        "id": "web_development_scenario_5",
        "specialization_key": "web_development",
        "specialization": "Web Development",
        "dimension": "scenario",
        "question": "For Web Development, choose how much you like a career scenario that uses css and javascript.",
        "keywords": [
            "html",
            "css",
            "javascript",
            "react"
        ],
        "weight": 4
    },
    {
        "id": "data_science_scenario_1",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "scenario",
        "question": "For Data Science, choose how much you like a beginner scenario that uses sql and statistics.",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 4
    },
    {
        "id": "data_science_scenario_2",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "scenario",
        "question": "For Data Science, choose how much you like a intermediate scenario that uses statistics and machine learning.",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 4
    },
    {
        "id": "data_science_scenario_3",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "scenario",
        "question": "For Data Science, choose how much you like a advanced scenario that uses machine learning and python.",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 4
    },
    {
        "id": "data_science_scenario_4",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "scenario",
        "question": "For Data Science, choose how much you like a project scenario that uses python and sql.",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 4
    },
    {
        "id": "data_science_scenario_5",
        "specialization_key": "data_science",
        "specialization": "Data Science",
        "dimension": "scenario",
        "question": "For Data Science, choose how much you like a career scenario that uses sql and statistics.",
        "keywords": [
            "python",
            "sql",
            "statistics",
            "machine learning"
        ],
        "weight": 4
    },
    {
        "id": "ai_ml_scenario_1",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "scenario",
        "question": "For AI and Machine Learning, choose how much you like a beginner scenario that uses machine learning and deep learning.",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 4
    },
    {
        "id": "ai_ml_scenario_2",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "scenario",
        "question": "For AI and Machine Learning, choose how much you like a intermediate scenario that uses deep learning and models.",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 4
    },
    {
        "id": "ai_ml_scenario_3",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "scenario",
        "question": "For AI and Machine Learning, choose how much you like a advanced scenario that uses models and python.",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 4
    },
    {
        "id": "ai_ml_scenario_4",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "scenario",
        "question": "For AI and Machine Learning, choose how much you like a project scenario that uses python and machine learning.",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 4
    },
    {
        "id": "ai_ml_scenario_5",
        "specialization_key": "ai_ml",
        "specialization": "AI and Machine Learning",
        "dimension": "scenario",
        "question": "For AI and Machine Learning, choose how much you like a career scenario that uses machine learning and deep learning.",
        "keywords": [
            "python",
            "machine learning",
            "deep learning",
            "models"
        ],
        "weight": 4
    },
    {
        "id": "cloud_devops_scenario_1",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "scenario",
        "question": "For Cloud and DevOps, choose how much you like a beginner scenario that uses docker and linux.",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 4
    },
    {
        "id": "cloud_devops_scenario_2",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "scenario",
        "question": "For Cloud and DevOps, choose how much you like a intermediate scenario that uses linux and ci/cd.",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 4
    },
    {
        "id": "cloud_devops_scenario_3",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "scenario",
        "question": "For Cloud and DevOps, choose how much you like a advanced scenario that uses ci/cd and aws.",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 4
    },
    {
        "id": "cloud_devops_scenario_4",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "scenario",
        "question": "For Cloud and DevOps, choose how much you like a project scenario that uses aws and docker.",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 4
    },
    {
        "id": "cloud_devops_scenario_5",
        "specialization_key": "cloud_devops",
        "specialization": "Cloud and DevOps",
        "dimension": "scenario",
        "question": "For Cloud and DevOps, choose how much you like a career scenario that uses docker and linux.",
        "keywords": [
            "aws",
            "docker",
            "linux",
            "ci/cd"
        ],
        "weight": 4
    },
    {
        "id": "database_scenario_1",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "scenario",
        "question": "For Database Systems, choose how much you like a beginner scenario that uses mysql and postgresql.",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 4
    },
    {
        "id": "database_scenario_2",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "scenario",
        "question": "For Database Systems, choose how much you like a intermediate scenario that uses postgresql and data modeling.",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 4
    },
    {
        "id": "database_scenario_3",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "scenario",
        "question": "For Database Systems, choose how much you like a advanced scenario that uses data modeling and sql.",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 4
    },
    {
        "id": "database_scenario_4",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "scenario",
        "question": "For Database Systems, choose how much you like a project scenario that uses sql and mysql.",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 4
    },
    {
        "id": "database_scenario_5",
        "specialization_key": "database",
        "specialization": "Database Systems",
        "dimension": "scenario",
        "question": "For Database Systems, choose how much you like a career scenario that uses mysql and postgresql.",
        "keywords": [
            "sql",
            "mysql",
            "postgresql",
            "data modeling"
        ],
        "weight": 4
    },
    {
        "id": "mobile_scenario_1",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "scenario",
        "question": "For Mobile App Development, choose how much you like a beginner scenario that uses swift and kotlin.",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 4
    },
    {
        "id": "mobile_scenario_2",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "scenario",
        "question": "For Mobile App Development, choose how much you like a intermediate scenario that uses kotlin and ui.",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 4
    },
    {
        "id": "mobile_scenario_3",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "scenario",
        "question": "For Mobile App Development, choose how much you like a advanced scenario that uses ui and flutter.",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 4
    },
    {
        "id": "mobile_scenario_4",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "scenario",
        "question": "For Mobile App Development, choose how much you like a project scenario that uses flutter and swift.",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 4
    },
    {
        "id": "mobile_scenario_5",
        "specialization_key": "mobile",
        "specialization": "Mobile App Development",
        "dimension": "scenario",
        "question": "For Mobile App Development, choose how much you like a career scenario that uses swift and kotlin.",
        "keywords": [
            "flutter",
            "swift",
            "kotlin",
            "ui"
        ],
        "weight": 4
    },
    {
        "id": "networks_scenario_1",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "scenario",
        "question": "For Computer Networks, choose how much you like a beginner scenario that uses tcp/ip and switching.",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 4
    },
    {
        "id": "networks_scenario_2",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "scenario",
        "question": "For Computer Networks, choose how much you like a intermediate scenario that uses switching and security.",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 4
    },
    {
        "id": "networks_scenario_3",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "scenario",
        "question": "For Computer Networks, choose how much you like a advanced scenario that uses security and routing.",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 4
    },
    {
        "id": "networks_scenario_4",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "scenario",
        "question": "For Computer Networks, choose how much you like a project scenario that uses routing and tcp/ip.",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 4
    },
    {
        "id": "networks_scenario_5",
        "specialization_key": "networks",
        "specialization": "Computer Networks",
        "dimension": "scenario",
        "question": "For Computer Networks, choose how much you like a career scenario that uses tcp/ip and switching.",
        "keywords": [
            "routing",
            "tcp/ip",
            "switching",
            "security"
        ],
        "weight": 4
    },
    {
        "id": "uiux_scenario_1",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "scenario",
        "question": "For UI/UX Engineering, choose how much you like a beginner scenario that uses accessibility and prototyping.",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 4
    },
    {
        "id": "uiux_scenario_2",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "scenario",
        "question": "For UI/UX Engineering, choose how much you like a intermediate scenario that uses prototyping and frontend.",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 4
    },
    {
        "id": "uiux_scenario_3",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "scenario",
        "question": "For UI/UX Engineering, choose how much you like a advanced scenario that uses frontend and design.",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 4
    },
    {
        "id": "uiux_scenario_4",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "scenario",
        "question": "For UI/UX Engineering, choose how much you like a project scenario that uses design and accessibility.",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 4
    },
    {
        "id": "uiux_scenario_5",
        "specialization_key": "uiux",
        "specialization": "UI/UX Engineering",
        "dimension": "scenario",
        "question": "For UI/UX Engineering, choose how much you like a career scenario that uses accessibility and prototyping.",
        "keywords": [
            "design",
            "accessibility",
            "prototyping",
            "frontend"
        ],
        "weight": 4
    },
    {
        "id": "game_scenario_1",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "scenario",
        "question": "For Game Development, choose how much you like a beginner scenario that uses unity and graphics.",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 4
    },
    {
        "id": "game_scenario_2",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "scenario",
        "question": "For Game Development, choose how much you like a intermediate scenario that uses graphics and logic.",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 4
    },
    {
        "id": "game_scenario_3",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "scenario",
        "question": "For Game Development, choose how much you like a advanced scenario that uses logic and c++.",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 4
    },
    {
        "id": "game_scenario_4",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "scenario",
        "question": "For Game Development, choose how much you like a project scenario that uses c++ and unity.",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 4
    },
    {
        "id": "game_scenario_5",
        "specialization_key": "game",
        "specialization": "Game Development",
        "dimension": "scenario",
        "question": "For Game Development, choose how much you like a career scenario that uses unity and graphics.",
        "keywords": [
            "c++",
            "unity",
            "graphics",
            "logic"
        ],
        "weight": 4
    }
]

SQR_PAGE_BLUEPRINTS = {
    "home": [
        "homeSpecializations",
        "homeCourses",
        "homeJobs"
    ],
    "profile": [
        "profileSummary",
        "profileProgressBars",
        "profileQuizHistory",
        "profileAtsHistory"
    ],
    "specializations": [
        "specializationDetails",
        "specializationsBox"
    ],
    "courses": [
        "courseDetails",
        "coursesBox"
    ],
    "quiz": [
        "quizDetails",
        "quizResult",
        "quizzesBox"
    ],
    "ats": [
        "atsCheckForm",
        "atsGenerateForm",
        "atsResult",
        "generatedResume"
    ],
    "jobs": [
        "jobsBox",
        "jobDetails"
    ],
    "recommendation": [
        "recommendationForm",
        "recommendationResult",
        "recommendationQuestionBank"
    ],
    "admin": [
        "adminStatsBox",
        "adminSpecializationsList",
        "adminCoursesList",
        "adminJobsList",
        "adminQuizzesList",
        "adminCertificatesList",
        "adminUsersList"
    ]
}

SQR_COLOR_THEME_TOKENS = {
    "background": "#020617",
    "surface": "rgba(15,23,42,0.82)",
    "cyan": "#22d3ee",
    "blue": "#3b82f6",
    "purple": "#8b5cf6",
    "pink": "#ec4899",
    "green": "#22c55e",
    "orange": "#f97316",
    "red": "#ef4444"
}


def sqr_patch_safe_count(table_name):
    try:
        if not table_exists(table_name):
            return 0
        row = query_db(f"SELECT COUNT(*) AS total FROM `{table_name}`", fetchone=True)
        return int(row.get("total") or 0) if row else 0
    except Exception:
        return 0


def sqr_patch_table_columns(table_name):
    try:
        if not DB_CONFIG.get("database"):
            return []
        rows = query_db(
            """
            SELECT COLUMN_NAME AS name
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
            ORDER BY ORDINAL_POSITION
            """,
            (DB_CONFIG.get("database"), table_name),
            fetchall=True
        ) or []
        return [r.get("name") for r in rows if r.get("name")]
    except Exception:
        return []


def sqr_patch_select_rows(table_name, limit=8, order_col="created_at"):
    try:
        if not table_exists(table_name):
            return []
        columns = sqr_patch_table_columns(table_name)
        order_sql = f" ORDER BY `{order_col}` DESC" if order_col in columns else ""
        safe_limit = max(1, min(int(limit or 8), 50))
        return query_db(f"SELECT * FROM `{table_name}`{order_sql} LIMIT {safe_limit}", fetchall=True) or []
    except Exception:
        return []


def sqr_patch_normalize_rows(table_name, rows):
    normalized = []
    for row in rows or []:
        try:
            if table_name == "specializations":
                normalized.append(normalize_specialization(row))
            elif table_name == "courses":
                normalized.append(normalize_course(row))
            elif table_name == "jobs":
                normalized.append(normalize_job(row))
            elif table_name == "quizzes":
                normalized.append(normalize_quiz(row))
            else:
                normalized.append(dict(row))
        except Exception:
            normalized.append(dict(row))
    return normalized


def sqr_patch_public_stats():
    return {
        "users": sqr_patch_safe_count("users"),
        "specializations": sqr_patch_safe_count("specializations"),
        "courses": sqr_patch_safe_count("courses"),
        "quizzes": sqr_patch_safe_count("quizzes"),
        "jobs": sqr_patch_safe_count("jobs"),
        "certificates": sqr_patch_safe_count("certificates"),
        "ats_results": sqr_patch_safe_count("ats_results"),
        "quiz_attempts": sqr_patch_safe_count("quiz_attempts")
    }


def sqr_patch_keyword_list(text):
    source = safe_text(text).lower()
    found = []
    for skill in TECH_SKILLS:
        if skill.lower() in source and skill not in found:
            found.append(skill)
    return found


def sqr_patch_score_text_against_keywords(text, keywords):
    body = safe_text(text).lower()
    keys = [safe_text(k).lower() for k in (keywords or []) if safe_text(k)]
    if not keys:
        return 0, []
    matched = [k for k in keys if k in body]
    score = round((len(matched) / max(len(keys), 1)) * 100)
    return min(100, score), matched


def sqr_patch_recommend_from_text(text, limit=5):
    profile_text = safe_text(text)
    specs = sqr_patch_normalize_rows("specializations", sqr_patch_select_rows("specializations", 50))
    jobs = sqr_patch_normalize_rows("jobs", sqr_patch_select_rows("jobs", 50))
    spec_matches = []
    for spec in specs:
        keywords = sqr_patch_keyword_list(" ".join([safe_text(spec.get("name")), safe_text(spec.get("description")), safe_text(spec.get("skills"))]))
        score, matched = sqr_patch_score_text_against_keywords(profile_text, keywords)
        if score or safe_text(spec.get("name")).lower() in profile_text.lower():
            spec_matches.append({
                "id": spec.get("id"),
                "name": spec.get("name"),
                "match_percentage": max(score, 20 if safe_text(spec.get("name")).lower() in profile_text.lower() else 0),
                "matched_skills": matched,
                "reason": "Matched profile text with specialization keywords."
            })
    job_matches = []
    for job in jobs:
        keywords = sqr_patch_keyword_list(" ".join([safe_text(job.get("title")), safe_text(job.get("description")), safe_text(job.get("skills")), safe_text(job.get("required_skills"))]))
        score, matched = sqr_patch_score_text_against_keywords(profile_text, keywords)
        if score or safe_text(job.get("title")).lower() in profile_text.lower():
            job_matches.append({
                "id": job.get("id"),
                "title": job.get("title"),
                "match_percentage": max(score, 20 if safe_text(job.get("title")).lower() in profile_text.lower() else 0),
                "matched_skills": matched,
                "reason": "Matched profile text with job skills."
            })
    spec_matches.sort(key=lambda item: item.get("match_percentage", 0), reverse=True)
    job_matches.sort(key=lambda item: item.get("match_percentage", 0), reverse=True)
    return {
        "recommended_specializations": spec_matches[:limit],
        "recommended_jobs": job_matches[:limit],
        "detected_skills": sqr_patch_keyword_list(profile_text),
        "roadmap": [
            "Choose the highest matching specialization.",
            "Open linked courses to start progress tracking.",
            "Complete course quizzes to raise profile progress.",
            "Use ATS tools for the target job.",
            "Apply first to jobs with stronger skill matches."
        ]
    }


def sqr_patch_profile_text(user):
    if not user:
        return ""
    return " ".join([
        safe_text(user.get("name")),
        safe_text(user.get("skills")),
        safe_text(user.get("interests")),
        safe_text(user.get("goal")),
        safe_text(user.get("work_style")) if isinstance(user, dict) else ""
    ])


def sqr_patch_user_activity(user_id):
    data = {"opened_courses": 0, "quiz_attempts": 0, "ats_results": 0}
    try:
        if table_exists("course_enrollments"):
            row = query_db("SELECT COUNT(*) AS total FROM course_enrollments WHERE user_id=%s", (user_id,), fetchone=True)
            data["opened_courses"] = int(row.get("total") or 0) if row else 0
        elif table_exists("progress"):
            row = query_db("SELECT COUNT(*) AS total FROM progress WHERE user_id=%s", (user_id,), fetchone=True)
            data["opened_courses"] = int(row.get("total") or 0) if row else 0
    except Exception:
        pass
    try:
        if table_exists("quiz_attempts"):
            row = query_db("SELECT COUNT(*) AS total FROM quiz_attempts WHERE user_id=%s", (user_id,), fetchone=True)
            data["quiz_attempts"] = int(row.get("total") or 0) if row else 0
    except Exception:
        pass
    try:
        if table_exists("ats_results"):
            row = query_db("SELECT COUNT(*) AS total FROM ats_results WHERE user_id=%s", (user_id,), fetchone=True)
            data["ats_results"] = int(row.get("total") or 0) if row else 0
    except Exception:
        pass
    return data


@app.route("/api/public/bootstrap", methods=["GET"])
def sqr_patch_public_bootstrap():
    return jsonify({
        "message": "SQR dynamic bootstrap loaded",
        "stats": sqr_patch_public_stats(),
        "specializations": sqr_patch_normalize_rows("specializations", sqr_patch_select_rows("specializations", 6)),
        "courses": sqr_patch_normalize_rows("courses", sqr_patch_select_rows("courses", 6)),
        "jobs": sqr_patch_normalize_rows("jobs", sqr_patch_select_rows("jobs", 6)),
        "theme": SQR_COLOR_THEME_TOKENS,
        "pages": SQR_PAGE_BLUEPRINTS
    })


@app.route("/api/home/dashboard", methods=["GET"])
def sqr_patch_home_dashboard():
    return jsonify({
        "stats": sqr_patch_public_stats(),
        "latest_specializations": sqr_patch_normalize_rows("specializations", sqr_patch_select_rows("specializations", 9)),
        "latest_courses": sqr_patch_normalize_rows("courses", sqr_patch_select_rows("courses", 9)),
        "latest_jobs": sqr_patch_normalize_rows("jobs", sqr_patch_select_rows("jobs", 9))
    })


@app.route("/api/recommendation/questions", methods=["GET"])
def sqr_patch_recommendation_questions():
    specialization_key = safe_text(request.args.get("specialization_key")).lower()
    dimension = safe_text(request.args.get("dimension")).lower()
    questions = []
    for question in sqr_recommendation_questions():
        if specialization_key and safe_text(question.get("specialization_key")).lower() != specialization_key:
            continue
        if dimension and safe_text(question.get("dimension")).lower() != dimension:
            continue
        questions.append(question)
    return jsonify({"questions": questions, "count": len(questions), "mode": "quiz_only"})


@app.route("/api/catalog/search", methods=["GET"])
def sqr_patch_catalog_search():
    term = safe_text(request.args.get("q")).lower()
    limit = max(1, min(int(request.args.get("limit", 12) or 12), 40))
    results = {"specializations": [], "courses": [], "jobs": [], "quizzes": []}
    if not term:
        return jsonify(results)
    for table_name, key in [("specializations", "specializations"), ("courses", "courses"), ("jobs", "jobs"), ("quizzes", "quizzes")]:
        rows = sqr_patch_normalize_rows(table_name, sqr_patch_select_rows(table_name, 50))
        filtered = []
        for row in rows:
            body = " ".join(safe_text(v) for v in row.values()).lower()
            if term in body:
                filtered.append(row)
        results[key] = filtered[:limit]
    return jsonify(results)


@app.route("/api/profile/dashboard/advanced", methods=["GET"])
@login_required
def sqr_patch_profile_dashboard_advanced():
    user = clean_user(request.current_user)
    user_id = user.get("id") or user.get("user_id")
    recommendation = sqr_patch_recommend_from_text(sqr_patch_profile_text(user), 5)
    activity = sqr_patch_user_activity(user_id)
    progress_payload = []
    try:
        if "profile_progress" in globals():
            pass
    except Exception:
        pass
    return jsonify({
        "user": user,
        "activity": activity,
        "recommendation_preview": recommendation,
        "profile_completeness": sqr_patch_profile_completeness(user),
        "stats": sqr_patch_public_stats(),
        "progress_hint": "Use /api/profile/progress for real progress bars shown only on profile.html."
    })


def sqr_patch_profile_completeness(user):
    fields = ["name", "email", "skills", "interests", "goal"]
    if not user:
        return 0
    filled = sum(1 for field in fields if safe_text(user.get(field)))
    return round((filled / len(fields)) * 100)


@app.route("/api/admin/dashboard/advanced", methods=["GET"])
@admin_required
def sqr_patch_admin_dashboard_advanced():
    return jsonify({
        "stats": sqr_patch_public_stats(),
        "tables": {
            name: sqr_patch_table_columns(name)
            for name in ["users", "specializations", "courses", "quizzes", "quiz_questions", "jobs", "certificates", "ats_results", "course_enrollments", "quiz_attempts"]
        },
        "recent": {
            "specializations": sqr_patch_normalize_rows("specializations", sqr_patch_select_rows("specializations", 5)),
            "courses": sqr_patch_normalize_rows("courses", sqr_patch_select_rows("courses", 5)),
            "jobs": sqr_patch_normalize_rows("jobs", sqr_patch_select_rows("jobs", 5)),
            "quizzes": sqr_patch_normalize_rows("quizzes", sqr_patch_select_rows("quizzes", 5))
        }
    })


@app.route("/api/schema/check", methods=["GET"])
def sqr_patch_schema_check():
    tables = ["users", "admins", "specializations", "courses", "quizzes", "quiz_questions", "jobs", "certificates", "course_enrollments", "quiz_attempts", "ats_results", "assessments"]
    return jsonify({
        "database": DB_CONFIG.get("database"),
        "connected": bool(pool),
        "tables": [
            {"name": table, "exists": table_exists(table), "columns": sqr_patch_table_columns(table)}
            for table in tables
        ]
    })


@app.route("/api/static/page-blueprint/<page_name>", methods=["GET"])
def sqr_patch_page_blueprint(page_name):
    key = safe_text(page_name).replace(".html", "").lower()
    return jsonify({"page": key, "dynamic_targets": SQR_PAGE_BLUEPRINTS.get(key, []), "theme": SQR_COLOR_THEME_TOKENS})


@app.route("/api/recommendations/preview", methods=["POST"])
@login_required
def sqr_patch_recommendations_preview():
    data = get_json()
    text = " ".join([
        safe_text(data.get("interests")),
        safe_text(data.get("skills")),
        safe_text(data.get("work_style")),
        safe_text(data.get("goal")),
        sqr_patch_profile_text(request.current_user)
    ])
    return jsonify(sqr_patch_recommend_from_text(text, 8))


# Extra backend view-model helpers used by the colorful templates.
# These routes are intentionally unique so they do not replace existing project features.
@app.route("/api/view-model/home", methods=["GET"])
def sqr_patch_view_model_home():
    payload = sqr_patch_public_stats()
    return jsonify({
        "page": "home",
        "title": "Skill Quest Road",
        "counts": payload,
        "sections": SQR_PAGE_BLUEPRINTS.get("home", []),
        "colors": SQR_COLOR_THEME_TOKENS
    })


@app.route("/api/view-model/profile", methods=["GET"])
@login_required
def sqr_patch_view_model_profile():
    user = clean_user(request.current_user)
    return jsonify({
        "page": "profile",
        "user": user,
        "activity": sqr_patch_user_activity(user.get("id") or user.get("user_id")),
        "completeness": sqr_patch_profile_completeness(user),
        "sections": SQR_PAGE_BLUEPRINTS.get("profile", [])
    })

SQR_DYNAMIC_CONTAINER_REGISTRY = [
    {
        "page": "home",
        "target": "homeSpecializations",
        "priority": 1,
        "purpose": "Dynamic container homeSpecializations on home page"
    },
    {
        "page": "home",
        "target": "homeCourses",
        "priority": 2,
        "purpose": "Dynamic container homeCourses on home page"
    },
    {
        "page": "home",
        "target": "homeJobs",
        "priority": 3,
        "purpose": "Dynamic container homeJobs on home page"
    },
    {
        "page": "profile",
        "target": "profileSummary",
        "priority": 1,
        "purpose": "Dynamic container profileSummary on profile page"
    },
    {
        "page": "profile",
        "target": "profileProgressBars",
        "priority": 2,
        "purpose": "Dynamic container profileProgressBars on profile page"
    },
    {
        "page": "profile",
        "target": "profileQuizHistory",
        "priority": 3,
        "purpose": "Dynamic container profileQuizHistory on profile page"
    },
    {
        "page": "profile",
        "target": "profileAtsHistory",
        "priority": 4,
        "purpose": "Dynamic container profileAtsHistory on profile page"
    },
    {
        "page": "specializations",
        "target": "specializationDetails",
        "priority": 1,
        "purpose": "Dynamic container specializationDetails on specializations page"
    },
    {
        "page": "specializations",
        "target": "specializationsBox",
        "priority": 2,
        "purpose": "Dynamic container specializationsBox on specializations page"
    },
    {
        "page": "courses",
        "target": "courseDetails",
        "priority": 1,
        "purpose": "Dynamic container courseDetails on courses page"
    },
    {
        "page": "courses",
        "target": "coursesBox",
        "priority": 2,
        "purpose": "Dynamic container coursesBox on courses page"
    },
    {
        "page": "quiz",
        "target": "quizDetails",
        "priority": 1,
        "purpose": "Dynamic container quizDetails on quiz page"
    },
    {
        "page": "quiz",
        "target": "quizResult",
        "priority": 2,
        "purpose": "Dynamic container quizResult on quiz page"
    },
    {
        "page": "quiz",
        "target": "quizzesBox",
        "priority": 3,
        "purpose": "Dynamic container quizzesBox on quiz page"
    },
    {
        "page": "ats",
        "target": "atsCheckForm",
        "priority": 1,
        "purpose": "Dynamic container atsCheckForm on ats page"
    },
    {
        "page": "ats",
        "target": "atsGenerateForm",
        "priority": 2,
        "purpose": "Dynamic container atsGenerateForm on ats page"
    },
    {
        "page": "ats",
        "target": "atsResult",
        "priority": 3,
        "purpose": "Dynamic container atsResult on ats page"
    },
    {
        "page": "ats",
        "target": "generatedResume",
        "priority": 4,
        "purpose": "Dynamic container generatedResume on ats page"
    },
    {
        "page": "jobs",
        "target": "jobsBox",
        "priority": 1,
        "purpose": "Dynamic container jobsBox on jobs page"
    },
    {
        "page": "jobs",
        "target": "jobDetails",
        "priority": 2,
        "purpose": "Dynamic container jobDetails on jobs page"
    },
    {
        "page": "recommendation",
        "target": "recommendationForm",
        "priority": 1,
        "purpose": "Dynamic container recommendationForm on recommendation page"
    },
    {
        "page": "recommendation",
        "target": "recommendationResult",
        "priority": 2,
        "purpose": "Dynamic container recommendationResult on recommendation page"
    },
    {
        "page": "recommendation",
        "target": "recommendationQuestionBank",
        "priority": 3,
        "purpose": "Dynamic container recommendationQuestionBank on recommendation page"
    },
    {
        "page": "admin",
        "target": "adminStatsBox",
        "priority": 1,
        "purpose": "Dynamic container adminStatsBox on admin page"
    },
    {
        "page": "admin",
        "target": "adminSpecializationsList",
        "priority": 2,
        "purpose": "Dynamic container adminSpecializationsList on admin page"
    },
    {
        "page": "admin",
        "target": "adminCoursesList",
        "priority": 3,
        "purpose": "Dynamic container adminCoursesList on admin page"
    },
    {
        "page": "admin",
        "target": "adminJobsList",
        "priority": 4,
        "purpose": "Dynamic container adminJobsList on admin page"
    },
    {
        "page": "admin",
        "target": "adminQuizzesList",
        "priority": 5,
        "purpose": "Dynamic container adminQuizzesList on admin page"
    },
    {
        "page": "admin",
        "target": "adminCertificatesList",
        "priority": 6,
        "purpose": "Dynamic container adminCertificatesList on admin page"
    },
    {
        "page": "admin",
        "target": "adminUsersList",
        "priority": 7,
        "purpose": "Dynamic container adminUsersList on admin page"
    }
]


@app.route("/api/static/dynamic-containers", methods=["GET"])
def sqr_patch_dynamic_containers():
    page = safe_text(request.args.get("page")).replace(".html", "").lower()
    rows = [row for row in SQR_DYNAMIC_CONTAINER_REGISTRY if not page or row.get("page") == page]
    return jsonify({"containers": rows, "count": len(rows)})


def sqr_patch_runtime_report():
    return {
        "python_file": "SQR.py",
        "db_host_set": bool(DB_CONFIG.get("host")),
        "db_name_set": bool(DB_CONFIG.get("database")),
        "gemini_enabled": bool(gemini_client),
        "upload_folder": app.config.get("UPLOAD_FOLDER"),
        "public_stats": sqr_patch_public_stats(),
        "page_targets": SQR_PAGE_BLUEPRINTS
    }


@app.route("/api/runtime/report", methods=["GET"])
def sqr_patch_runtime_report_route():
    return jsonify(sqr_patch_runtime_report())


@app.route("/api/recommendation", methods=["POST"])
@student_required
def recommendation_alias():
    return recommendations()


@app.route("/api/recommendation/analyze", methods=["POST"])
@student_required
def recommendation_analyze_alias():
    return recommendations()



@app.route("/api/recommendations/jobs", methods=["POST"])
@student_required
def recommendation_jobs_alias():
    data = get_json()
    user = request.current_user
    user_id = user.get("id") or user.get("user_id")

    latest_result = {}
    try:
        if table_exists("recommendation_results"):
            latest = query_db(
                "SELECT recommendation_json FROM recommendation_results WHERE user_id=%s ORDER BY recommendation_id DESC LIMIT 1",
                (user_id,),
                fetchone=True
            ) or {}
            latest_result = json.loads(latest.get("recommendation_json") or "{}") if latest.get("recommendation_json") else {}
    except Exception as exc:
        print("LATEST RECOMMENDATION LOAD ERROR:", exc)
        latest_result = {}

    quiz_scores = latest_result.get("quiz_scores") or sqr_score_recommendation_quiz(data)[0]
    recommended_specs = latest_result.get("recommended_specializations") or []
    profile_text = sqr_recommendation_profile_text({"answers": latest_result.get("quiz_answers") or data.get("answers") or []}, user)
    profile_text += " " + " ".join(safe_text(s.get("name")) for s in recommended_specs if isinstance(s, dict))

    jobs = [normalize_job(row) for row in (query_db(
        """
        SELECT j.*, s.name AS specialization_name
        FROM jobs j
        LEFT JOIN specializations s ON s.specialization_id=j.specialization_id
        ORDER BY j.job_id DESC
        """,
        fetchall=True
    ) or [])]

    spec_name_scores = {safe_text(s.get("name")).lower(): pct_value(s.get("match_percentage") or s.get("score")) for s in recommended_specs if isinstance(s, dict)}
    recommended_jobs = []
    for job in jobs:
        target = f"{job.get('title','')} {job.get('description','')} {job.get('skills','')} {job.get('required_skills','')} {job.get('specialization','')} {job.get('specialization_name','')}"
        skill_score, matches = calculate_match_percentage(profile_text, target)
        spec_label = safe_text(job.get("specialization_name") or job.get("specialization")).lower()
        spec_boost = spec_name_scores.get(spec_label, 0)
        if spec_boost <= 0:
            spec_slug = sqr_slug(spec_label)
            spec_boost = pct_value(quiz_scores.get(spec_slug))
        final_score = round((skill_score * 0.58) + (spec_boost * 0.42)) if (skill_score or spec_boost) else 0
        final_score = max(0, min(100, final_score))
        recommended_jobs.append({
            "id": job.get("id"),
            "job_id": job.get("job_id") or job.get("id"),
            "specialization_id": job.get("specialization_id"),
            "title": job.get("title"),
            "description": job.get("description") or "",
            "match_percentage": final_score,
            "score": final_score,
            "matched_skills": matches,
            "salary": job.get("salary") or job.get("average_salary"),
            "specialization_name": job.get("specialization_name") or job.get("specialization"),
            "link": job.get("link") or job.get("job_link"),
            "reason": "Matched on the Jobs page using your latest specialization quiz result and job skill keywords.",
        })
    recommended_jobs.sort(key=lambda item: item.get("match_percentage", 0), reverse=True)

    return jsonify({
        "recommended_jobs": recommended_jobs[:8],
        "summary": "Job recommendations are separate from specialization recommendations and are shown only on the Jobs page.",
        "recommendation_source": "latest_quiz" if latest_result else "profile_or_current_answers",
        "recommended_specializations_used": recommended_specs[:3],
    })


@app.errorhandler(404)
def not_found(error):
    if request.path.startswith("/api/"):
        return jsonify({"error": "Endpoint not found"}), 404
    return render_template("gp.html"), 404


@app.errorhandler(413)
def too_large(error):
    return jsonify({"error": "File is too large"}), 413


@app.errorhandler(500)
def server_error(error):
    return jsonify({"error": "Server error", "details": str(error)}), 500


try:
    init_db()
    print("SQR database checked")
except Exception as exc:
    print("init_db skipped or failed:", exc)


if __name__ == "__main__":
    app.run(host=os.getenv("FLASK_HOST", "0.0.0.0"), port=int(os.getenv("PORT", os.getenv("FLASK_PORT", 5000))), debug=os.getenv("FLASK_DEBUG", "0") == "1")
