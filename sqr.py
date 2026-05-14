import os
import re
import json
import datetime
from io import BytesIO
from functools import wraps

import jwt
import mysql.connector
from mysql.connector import pooling
from flask import Flask, request, jsonify, send_from_directory, send_file, render_template
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
except Exception:
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    HRFlowable = None
    A4 = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    TA_CENTER = None
    TA_LEFT = None
    colors = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")
UPLOAD_DIR = os.path.join(BASE_DIR, os.getenv("UPLOAD_FOLDER", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = Flask(__name__, template_folder=TEMPLATES_DIR, static_folder=STATIC_DIR)
CORS(app, resources={r"/api/*": {"origins": os.getenv("CORS_ORIGINS", "*").split(",")}})

app.config["SECRET_KEY"] = os.getenv("SQR_SECRET_KEY") or os.getenv("SECRET_KEY") or "CHANGE_THIS_SECRET_KEY_BEFORE_DEPLOYMENT"
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_CONTENT_LENGTH", str(50 * 1024 * 1024)))
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR

DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost").strip(),
    "port": int(os.getenv("DB_PORT", "3306").strip() or "3306"),
    "user": os.getenv("DB_USER", "root").strip(),
    "password": os.getenv("DB_PASSWORD", "").strip(),
    "database": os.getenv("DB_NAME", "railway").strip(),
    "connection_timeout": int(os.getenv("DB_TIMEOUT", "10")),
    "autocommit": True,
}

pool = None
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY")) if OpenAI and os.getenv("OPENAI_API_KEY") else None
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

TECH_SKILLS = [
    "python", "java", "javascript", "typescript", "html", "css", "sql", "mysql", "postgresql",
    "react", "node", "flask", "django", "api", "rest", "git", "github", "docker", "aws", "azure",
    "linux", "security", "cybersecurity", "networking", "forensics", "wireshark", "burp suite", "database",
    "machine learning", "ai", "data analysis", "data engineering", "etl", "cloud", "mongodb", "communication",
    "teamwork", "problem solving", "devops", "kubernetes", "php", "c++", "go", "rust", "swift",
    "figma", "ui", "ux", "testing", "automation", "incident response", "siem"
]

COURSE_LEVEL_META = {
    "beginner": {"label": "Beginner", "class": "level-beginner", "hex": "#22c55e"},
    "intermediate": {"label": "Intermediate", "class": "level-intermediate", "hex": "#eab308"},
    "advanced": {"label": "Advanced", "class": "level-advanced", "hex": "#ef4444"},
}

SPECIALIZATION_HINTS = {
    "cybersecurity": ["security", "cyber", "network", "linux", "forensics", "burp", "wireshark", "soc", "vulnerability"],
    "digital forensics": ["forensics", "evidence", "investigation", "incident", "malware", "security"],
    "software engineering": ["software", "java", "python", "problem", "api", "backend", "testing"],
    "web development": ["html", "css", "javascript", "react", "frontend", "backend", "node", "flask"],
    "data science": ["data", "python", "sql", "analysis", "machine learning", "statistics", "visualization"],
    "artificial intelligence": ["ai", "machine learning", "automation", "model", "python", "nlp", "vision"],
    "cloud computing": ["cloud", "aws", "azure", "deployment", "server", "docker", "devops"],
    "database administration": ["database", "sql", "mysql", "postgresql", "queries", "schema", "admin"],
    "computer networks": ["network", "tcp", "ip", "routing", "switching", "security", "linux"],
    "ui/ux engineering": ["ui", "ux", "design", "interface", "figma", "frontend", "user"],
}


def get_db():
    global pool
    if pool is None:
        pool = pooling.MySQLConnectionPool(
            pool_name="sqr_pool",
            pool_size=int(os.getenv("DB_POOL_SIZE", "5")),
            **DB_CONFIG
        )
    return pool.get_connection()


def query_db(sql, params=None, fetchone=False, fetchall=False, commit=False):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    try:
        cursor.execute(sql, params or ())
        result = None
        if fetchone:
            result = cursor.fetchone()
        elif fetchall:
            result = cursor.fetchall()
        if commit:
            db.commit()
            result = cursor.lastrowid
        return result
    except Exception:
        if commit:
            db.rollback()
        raise
    finally:
        cursor.close()
        db.close()


def exec_db(sql, params=None):
    return query_db(sql, params=params, commit=True)


def get_json():
    return request.get_json(silent=True) or {}


def safe_text(value):
    return str(value or "").strip()


def safe_int(value, default=0):
    try:
        if value is None or value == "":
            return default
        return int(float(value))
    except Exception:
        return default


def table_exists(table_name):
    try:
        row = query_db(
            """
            SELECT COUNT(*) AS total
            FROM INFORMATION_SCHEMA.TABLES
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
            """,
            (DB_CONFIG["database"], table_name),
            fetchone=True
        )
        return bool(row and row.get("total"))
    except Exception:
        return False


def column_exists(table_name, column_name):
    try:
        row = query_db(
            """
            SELECT COUNT(*) AS total
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s AND COLUMN_NAME=%s
            """,
            (DB_CONFIG["database"], table_name, column_name),
            fetchone=True
        )
        return bool(row and row.get("total"))
    except Exception:
        return False


def first_existing_column(table_name, names):
    for name in names:
        if column_exists(table_name, name):
            return name
    return names[0]


def row_value(row, *names):
    for name in names:
        if isinstance(row, dict) and row.get(name) not in [None, ""]:
            return row.get(name)
    return None


def upload_url(filename):
    value = safe_text(filename)
    if not value:
        return ""
    if value.startswith("http://") or value.startswith("https://") or value.startswith("/uploads/"):
        return value
    return f"/uploads/{value}"


def normalize_level(level):
    value = safe_text(level).lower() or "beginner"
    aliases = {
        "begginer": "beginner",
        "beginner": "beginner",
        "medium": "intermediate",
        "intermidiete": "intermediate",
        "intermediate": "intermediate",
        "advance": "advanced",
        "advanced": "advanced",
    }
    return aliases.get(value, "beginner")


def add_level_meta(course):
    level = normalize_level(course.get("level"))
    course["level"] = level
    course["level_badge"] = COURSE_LEVEL_META[level]
    return course


def normalize_specialization(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "specialization_id")
    row["specialization_id"] = row["id"]
    row["image_url"] = upload_url(row_value(row, "image", "image_url"))
    row["skills"] = row_value(row, "skills", "required_skills") or ""
    row["description"] = row.get("description") or ""
    return row


def normalize_course(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "course_id")
    row["course_id"] = row["id"]
    row["specialization_id"] = row_value(row, "specialization_id", "spec_id")
    row["spec_id"] = row["specialization_id"]
    row["link"] = row_value(row, "link", "course_link") or ""
    row["course_link"] = row["link"]
    row["image_url"] = upload_url(row_value(row, "image", "image_url", "thumbnail"))
    row["video_url"] = upload_url(row_value(row, "video", "video_url", "media_url"))
    row["specialization_name"] = row_value(row, "specialization_name", "specialization") or ""
    return add_level_meta(row)


def normalize_quiz(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "quiz_id")
    row["quiz_id"] = row["id"]
    row["course_id"] = row_value(row, "course_id", "course")
    return row


def normalize_question(row, include_answer=False):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "question_id")
    row["question_id"] = row["id"]
    row["question"] = row_value(row, "question", "question_text") or ""
    row["options"] = [
        row_value(row, "option1", "option_a") or "",
        row_value(row, "option2", "option_b") or "",
        row_value(row, "option3", "option_c") or "",
        row_value(row, "option4", "option_d") or "",
    ]
    if not include_answer:
        row.pop("answer", None)
        row.pop("correct_answer", None)
    return row


def normalize_job(row):
    if not row:
        return row
    row = dict(row)
    row["id"] = row_value(row, "id", "job_id")
    row["job_id"] = row["id"]
    row["skills"] = row_value(row, "skills", "required_skills") or ""
    row["required_skills"] = row["skills"]
    row["salary"] = row_value(row, "salary", "average_salary") or ""
    row["average_salary"] = row["salary"]
    row["link"] = row_value(row, "link", "job_link") or ""
    row["job_link"] = row["link"]
    row["specialization"] = row_value(row, "specialization_name", "specialization") or ""
    return row


def clean_user(user):
    if not user:
        return None
    user = dict(user)
    user.pop("password", None)
    user["id"] = row_value(user, "id", "user_id")
    user["user_id"] = user["id"]
    user["current_mode"] = user.get("current_mode") or user.get("role") or "student"
    user["banned"] = safe_int(row_value(user, "banned", "is_banned"), 0)
    return user


def strong_password(password):
    password = str(password or "")
    return (
        len(password) >= 8
        and re.search(r"[A-Z]", password)
        and re.search(r"[a-z]", password)
        and re.search(r"[0-9]", password)
        and re.search(r"[^A-Za-z0-9]", password)
        and not re.search(r"\s", password)
    )


def generate_username(name, email):
    base = (email.split("@")[0] if email and "@" in email else name).strip().lower()
    base = re.sub(r"[^a-z0-9_]+", "_", base).strip("_") or "student"
    candidate = base
    counter = 1
    while query_db("SELECT id FROM users WHERE username=%s", (candidate,), fetchone=True):
        counter += 1
        candidate = f"{base}{counter}"
    return candidate


def generate_token(user):
    uid = row_value(user, "id", "user_id")
    payload = {
        "id": uid,
        "user_id": uid,
        "name": user.get("name"),
        "email": user.get("email"),
        "role": user.get("role", "student"),
        "current_mode": user.get("current_mode", user.get("role", "student")),
        "exp": datetime.datetime.utcnow() + datetime.timedelta(hours=int(os.getenv("JWT_HOURS", "24"))),
    }
    return jwt.encode(payload, app.config["SECRET_KEY"], algorithm="HS256")


def get_current_user():
    token = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
    if not token:
        return None
    try:
        data = jwt.decode(token, app.config["SECRET_KEY"], algorithms=["HS256"])
        uid = data.get("user_id") or data.get("id")
        user = query_db("SELECT * FROM users WHERE id=%s", (uid,), fetchone=True)
        if not user:
            return None
        if safe_int(user.get("banned"), 0) == 1:
            return None
        return user
    except Exception:
        return None


def login_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def admin_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        if safe_text(user.get("role")).lower() != "admin":
            return jsonify({"error": "Admin only"}), 403
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def student_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({"error": "Unauthorized"}), 401
        if safe_text(user.get("role")).lower() == "admin" and safe_text(user.get("current_mode") or "admin").lower() != "student":
            return jsonify({"error": "Admins can only access the admin page unless switched to student mode."}), 403
        request.current_user = user
        return func(*args, **kwargs)
    return wrapper


def allowed_file(filename):
    allowed = {"png", "jpg", "jpeg", "gif", "webp", "mp4", "mov", "webm", "ogg", "pdf", "docx", "txt"}
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed


def save_file(field_name):
    file = request.files.get(field_name)
    if not file or not file.filename:
        return ""
    if not allowed_file(file.filename):
        return ""
    original = secure_filename(file.filename)
    stamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
    filename = f"{stamp}_{original}"
    file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
    return filename


def request_data():
    if request.content_type and "multipart/form-data" in request.content_type:
        return dict(request.form)
    return get_json()


def calculate_match_percentage(profile_text, target_text):
    profile = safe_text(profile_text).lower()
    target = safe_text(target_text).lower()
    words = []
    for skill in TECH_SKILLS:
        if skill in profile or skill in target:
            words.append(skill)
    matched = [skill for skill in words if skill in profile and skill in target]
    unique_target = sorted(set([skill for skill in words if skill in target]))
    if not unique_target:
        tokens = set(re.findall(r"[a-zA-Z][a-zA-Z+#.]{2,}", target))
        user_tokens = set(re.findall(r"[a-zA-Z][a-zA-Z+#.]{2,}", profile))
        if not tokens:
            return 0, []
        score = round((len(tokens & user_tokens) / max(len(tokens), 1)) * 100)
        return max(0, min(100, score)), sorted(tokens & user_tokens)[:12]
    score = round((len(set(matched)) / max(len(unique_target), 1)) * 100)
    return max(0, min(100, score)), sorted(set(matched))[:12]


def extract_resume_text(file):
    if not file or not file.filename:
        return ""
    name = file.filename.lower()
    raw = file.read()
    file.seek(0)
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf") and PdfReader:
        reader = PdfReader(BytesIO(raw))
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    if name.endswith(".docx") and Document:
        doc = Document(BytesIO(raw))
        return "\n".join([p.text for p in doc.paragraphs])
    return ""


def ai_json(prompt, fallback):
    if not client:
        return fallback
    try:
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Return valid JSON only. Do not invent user experience."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )
        text = response.choices[0].message.content or "{}"
        match = re.search(r"\{.*\}", text, re.DOTALL)
        return json.loads(match.group(0) if match else text)
    except Exception:
        return fallback


def init_db():
    statements = [
        """
        CREATE TABLE IF NOT EXISTS users (
            id INT AUTO_INCREMENT PRIMARY KEY,
            username VARCHAR(120) UNIQUE,
            name VARCHAR(160) NOT NULL,
            email VARCHAR(180) NOT NULL UNIQUE,
            password VARCHAR(255) NOT NULL,
            role ENUM('student','admin') DEFAULT 'student',
            current_mode ENUM('student','admin') DEFAULT 'student',
            banned TINYINT DEFAULT 0,
            skills TEXT,
            interests TEXT,
            goal TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS admins (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL UNIQUE,
            admin_level VARCHAR(80) DEFAULT 'manager',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS specializations (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(180) NOT NULL,
            description TEXT,
            skills TEXT,
            image VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS courses (
            id INT AUTO_INCREMENT PRIMARY KEY,
            spec_id INT,
            title VARCHAR(220) NOT NULL,
            description TEXT,
            level VARCHAR(40) DEFAULT 'beginner',
            link TEXT,
            image VARCHAR(255),
            video VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (spec_id) REFERENCES specializations(id) ON DELETE SET NULL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS certificates (
            id INT AUTO_INCREMENT PRIMARY KEY,
            specialization_id INT,
            name VARCHAR(220) NOT NULL,
            description TEXT,
            link TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (specialization_id) REFERENCES specializations(id) ON DELETE SET NULL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quizzes (
            id INT AUTO_INCREMENT PRIMARY KEY,
            course_id INT,
            title VARCHAR(220) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quiz_questions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            quiz_id INT NOT NULL,
            question TEXT NOT NULL,
            option1 TEXT,
            option2 TEXT,
            option3 TEXT,
            option4 TEXT,
            answer VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS jobs (
            id INT AUTO_INCREMENT PRIMARY KEY,
            specialization_id INT,
            title VARCHAR(220) NOT NULL,
            description TEXT,
            skills TEXT,
            salary VARCHAR(120),
            link TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (specialization_id) REFERENCES specializations(id) ON DELETE SET NULL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS course_progress (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            course_id INT NOT NULL,
            opened TINYINT DEFAULT 1,
            video_started TINYINT DEFAULT 0,
            completed TINYINT DEFAULT 0,
            opened_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_course_user (user_id, course_id),
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY (course_id) REFERENCES courses(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS specialization_progress (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            specialization_id INT NOT NULL,
            progress INT DEFAULT 0,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_spec_user (user_id, specialization_id),
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY (specialization_id) REFERENCES specializations(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS quiz_attempts (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            quiz_id INT NOT NULL,
            course_id INT,
            score INT DEFAULT 0,
            total INT DEFAULT 0,
            score_percentage INT DEFAULT 0,
            answers_json TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE CASCADE,
            FOREIGN KEY (course_id) REFERENCES courses(id) ON DELETE SET NULL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS ats_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            target_job VARCHAR(220),
            score INT DEFAULT 0,
            summary TEXT,
            matched_keywords TEXT,
            missing_keywords TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS recommendation_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT NOT NULL,
            recommendation_json LONGTEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
        """,
    ]
    for statement in statements:
        exec_db(statement)


def render_page(template_name):
    return render_template(template_name)


@app.route("/")
def home():
    return render_page("gp.html")


@app.route("/home")
def page_home():
    return render_page("gp.html")


@app.route("/specializations")
def page_specializations():
    return render_page("Specialization.html")


@app.route("/courses")
def page_courses():
    return render_page("Courses.html")


@app.route("/quizzes")
def page_quizzes():
    return render_page("Quiz.html")


@app.route("/ats")
def page_ats():
    return render_page("ATS.html")


@app.route("/jobs")
def page_jobs():
    return render_page("jobs.html")


@app.route("/recommendation")
def page_recommendation():
    return render_page("recommendation.html")


@app.route("/profile")
def page_profile():
    return render_page("profile.html")


@app.route("/admin")
def page_admin():
    return render_page("admin.html")


@app.route("/signin")
def page_signin():
    return render_page("signin.html")


@app.route("/signup")
def page_signup():
    return render_page("signup.html")


@app.route("/<path:page>.html")
def legacy_html_pages(page):
    aliases = {
        "gp": "gp.html",
        "Specialization": "Specialization.html",
        "Sepecialization": "Specialization.html",
        "Courses": "Courses.html",
        "courses": "Courses.html",
        "Quiz": "Quiz.html",
        "ATS": "ATS.html",
        "ats": "ATS.html",
        "jobs": "jobs.html",
        "JobDetails": "JobDetails.html",
        "recommendation": "recommendation.html",
        "profile": "profile.html",
        "admin": "admin.html",
        "signin": "signin.html",
        "signup": "signup.html",
    }
    return render_page(aliases.get(page, f"{page}.html"))


@app.route("/uploads/<path:filename>")
def uploads(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


@app.route("/api/health")
def health():
    return jsonify({
        "message": "SQR Backend is running",
        "features": ["auth", "admin", "specializations", "courses", "quizzes", "jobs", "profile", "progress", "ATS", "recommendation"],
        "database": DB_CONFIG.get("database")
    })


@app.route("/api/signup", methods=["POST"])
def signup():
    data = get_json()
    name = safe_text(data.get("name"))
    email = safe_text(data.get("email")).lower()
    password = safe_text(data.get("password"))
    if not name or not email or not password:
        return jsonify({"error": "Name, email, and password are required"}), 400
    if not strong_password(password):
        return jsonify({"error": "Password must be at least 8 characters and include uppercase, lowercase, number, and symbol"}), 400
    if query_db("SELECT id FROM users WHERE email=%s", (email,), fetchone=True):
        return jsonify({"error": "Email already exists"}), 409
    username = generate_username(name, email)
    hashed_password = generate_password_hash(password, method="pbkdf2:sha256", salt_length=16)
    user_id = query_db(
        """
        INSERT INTO users (username, name, email, password, role, current_mode, banned)
        VALUES (%s,%s,%s,%s,'student','student',0)
        """,
        (username, name, email, hashed_password),
        commit=True
    )
    user = query_db("SELECT * FROM users WHERE id=%s", (user_id,), fetchone=True)
    return jsonify({"message": "Account created", "token": generate_token(user), "user": clean_user(user)}), 201


@app.route("/api/signin", methods=["POST"])
def signin():
    data = get_json()
    email = safe_text(data.get("email")).lower()
    password = safe_text(data.get("password"))
    user = query_db("SELECT * FROM users WHERE email=%s", (email,), fetchone=True)
    if not user or not check_password_hash(user.get("password") or "", password):
        return jsonify({"error": "Invalid email or password"}), 401
    if safe_int(user.get("banned"), 0) == 1:
        return jsonify({"error": "Your account is banned"}), 403
    return jsonify({"message": "Login successful", "token": generate_token(user), "user": clean_user(user)})


@app.route("/api/me")
@login_required
def me():
    return jsonify(clean_user(request.current_user))


@app.route("/api/profile", methods=["GET"])
@login_required
def get_profile():
    user = clean_user(request.current_user)
    user_id = user["id"]
    quiz_history = query_db(
        """
        SELECT qa.id, qa.score, qa.total, qa.score_percentage, qa.created_at, q.title AS quiz_title, c.title AS course_title
        FROM quiz_attempts qa
        LEFT JOIN quizzes q ON q.id=qa.quiz_id
        LEFT JOIN courses c ON c.id=qa.course_id
        WHERE qa.user_id=%s
        ORDER BY qa.created_at DESC
        LIMIT 30
        """,
        (user_id,),
        fetchall=True
    ) or []
    ats_history = query_db(
        """
        SELECT id,target_job,score,summary,matched_keywords,missing_keywords,created_at
        FROM ats_results
        WHERE user_id=%s
        ORDER BY created_at DESC
        LIMIT 20
        """,
        (user_id,),
        fetchall=True
    ) or []
    return jsonify({"user": user, "quiz_history": quiz_history, "ats_history": ats_history})


@app.route("/api/profile", methods=["PUT"])
@login_required
def update_profile():
    data = get_json()
    name = safe_text(data.get("name")) or request.current_user.get("name")
    skills = safe_text(data.get("skills"))
    interests = safe_text(data.get("interests"))
    goal = safe_text(data.get("goal"))
    exec_db(
        "UPDATE users SET name=%s, skills=%s, interests=%s, goal=%s WHERE id=%s",
        (name, skills, interests, goal, request.current_user["id"])
    )
    user = query_db("SELECT * FROM users WHERE id=%s", (request.current_user["id"],), fetchone=True)
    return jsonify({"message": "Profile updated", "user": clean_user(user)})


def compute_user_progress(user_id):
    specs = query_db("SELECT * FROM specializations ORDER BY name", fetchall=True) or []
    progress_rows = []
    for spec in specs:
        spec_id = spec["id"]
        total_courses_row = query_db("SELECT COUNT(*) AS total FROM courses WHERE spec_id=%s", (spec_id,), fetchone=True) or {"total": 0}
        total_courses = safe_int(total_courses_row.get("total"), 0)
        opened_row = query_db(
            """
            SELECT COUNT(DISTINCT cp.course_id) AS total
            FROM course_progress cp
            JOIN courses c ON c.id=cp.course_id
            WHERE cp.user_id=%s AND c.spec_id=%s AND cp.opened=1
            """,
            (user_id, spec_id),
            fetchone=True
        ) or {"total": 0}
        quiz_row = query_db(
            """
            SELECT COUNT(DISTINCT qa.course_id) AS completed_quizzes, COALESCE(ROUND(AVG(qa.score_percentage),0),0) AS average_score
            FROM quiz_attempts qa
            JOIN courses c ON c.id=qa.course_id
            WHERE qa.user_id=%s AND c.spec_id=%s AND qa.score_percentage >= 60
            """,
            (user_id, spec_id),
            fetchone=True
        ) or {"completed_quizzes": 0, "average_score": 0}
        opened_courses = safe_int(opened_row.get("total"), 0)
        completed_quizzes = safe_int(quiz_row.get("completed_quizzes"), 0)
        average_score = safe_int(quiz_row.get("average_score"), 0)
        if total_courses <= 0:
            percent_value = 0
        else:
            opened_part = (opened_courses / total_courses) * 50
            quiz_part = (completed_quizzes / total_courses) * 50
            percent_value = max(0, min(100, round(opened_part + quiz_part)))
        exec_db(
            """
            INSERT INTO specialization_progress (user_id, specialization_id, progress)
            VALUES (%s,%s,%s)
            ON DUPLICATE KEY UPDATE progress=%s
            """,
            (user_id, spec_id, percent_value, percent_value)
        )
        progress_rows.append({
            "specialization_id": spec_id,
            "specialization_name": spec.get("name"),
            "total_courses": total_courses,
            "opened_courses": opened_courses,
            "completed_quizzes": completed_quizzes,
            "average_quiz_score": average_score,
            "progress": percent_value,
        })
    return progress_rows


@app.route("/api/profile/progress")
@login_required
def profile_progress():
    return jsonify({"progress": compute_user_progress(request.current_user["id"])})


@app.route("/api/specializations", methods=["GET"])
def get_specializations():
    rows = query_db("SELECT * FROM specializations ORDER BY id DESC", fetchall=True) or []
    return jsonify({"specializations": [normalize_specialization(row) for row in rows]})


@app.route("/api/specializations/<int:spec_id>", methods=["GET"])
def get_specialization(spec_id):
    spec = query_db("SELECT * FROM specializations WHERE id=%s", (spec_id,), fetchone=True)
    if not spec:
        return jsonify({"error": "Specialization not found"}), 404
    courses = query_db("SELECT c.*, s.name AS specialization_name FROM courses c LEFT JOIN specializations s ON s.id=c.spec_id WHERE c.spec_id=%s ORDER BY c.id DESC", (spec_id,), fetchall=True) or []
    jobs = query_db("SELECT j.*, s.name AS specialization_name FROM jobs j LEFT JOIN specializations s ON s.id=j.specialization_id WHERE j.specialization_id=%s ORDER BY j.id DESC", (spec_id,), fetchall=True) or []
    certificates = query_db("SELECT * FROM certificates WHERE specialization_id=%s ORDER BY id DESC", (spec_id,), fetchall=True) or []
    return jsonify({
        "specialization": normalize_specialization(spec),
        "courses": [normalize_course(row) for row in courses],
        "jobs": [normalize_job(row) for row in jobs],
        "certificates": certificates,
    })


@app.route("/api/specializations", methods=["POST"])
@admin_required
def add_specialization():
    data = request_data()
    image = save_file("image") or safe_text(data.get("image"))
    name = safe_text(data.get("name"))
    if not name:
        return jsonify({"error": "Specialization name is required"}), 400
    spec_id = query_db(
        "INSERT INTO specializations (name, description, skills, image) VALUES (%s,%s,%s,%s)",
        (name, safe_text(data.get("description")), safe_text(data.get("skills")), image),
        commit=True
    )
    return jsonify({"message": "Specialization added", "id": spec_id})


@app.route("/api/specializations/<int:spec_id>", methods=["PUT"])
@admin_required
def update_specialization(spec_id):
    data = request_data()
    old = query_db("SELECT * FROM specializations WHERE id=%s", (spec_id,), fetchone=True)
    if not old:
        return jsonify({"error": "Specialization not found"}), 404
    image = save_file("image") or safe_text(data.get("image")) or old.get("image")
    exec_db(
        "UPDATE specializations SET name=%s, description=%s, skills=%s, image=%s WHERE id=%s",
        (safe_text(data.get("name")) or old.get("name"), safe_text(data.get("description")), safe_text(data.get("skills")), image, spec_id)
    )
    return jsonify({"message": "Specialization updated"})


@app.route("/api/specializations/<int:spec_id>", methods=["DELETE"])
@admin_required
def delete_specialization(spec_id):
    exec_db("DELETE FROM specializations WHERE id=%s", (spec_id,))
    return jsonify({"message": "Specialization deleted"})


@app.route("/api/courses", methods=["GET"])
def get_courses():
    spec_id = request.args.get("spec_id") or request.args.get("specialization_id")
    search = safe_text(request.args.get("search"))
    sql = "SELECT c.*, s.name AS specialization_name FROM courses c LEFT JOIN specializations s ON s.id=c.spec_id WHERE 1=1"
    params = []
    if spec_id:
        sql += " AND c.spec_id=%s"
        params.append(spec_id)
    if search:
        sql += " AND (c.title LIKE %s OR c.description LIKE %s OR c.level LIKE %s)"
        params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
    sql += " ORDER BY c.id DESC"
    rows = query_db(sql, tuple(params), fetchall=True) or []
    return jsonify({"courses": [normalize_course(row) for row in rows]})


@app.route("/api/courses/<int:course_id>", methods=["GET"])
def get_course(course_id):
    course = query_db("SELECT c.*, s.name AS specialization_name FROM courses c LEFT JOIN specializations s ON s.id=c.spec_id WHERE c.id=%s", (course_id,), fetchone=True)
    if not course:
        return jsonify({"error": "Course not found"}), 404
    quizzes = query_db("SELECT * FROM quizzes WHERE course_id=%s ORDER BY id DESC", (course_id,), fetchall=True) or []
    return jsonify({"course": normalize_course(course), "quizzes": [normalize_quiz(row) for row in quizzes]})


@app.route("/api/courses", methods=["POST"])
@admin_required
def add_course():
    data = request_data()
    image = save_file("image") or safe_text(data.get("image"))
    video = save_file("video") or safe_text(data.get("video"))
    title = safe_text(data.get("title"))
    if not title:
        return jsonify({"error": "Course title is required"}), 400
    course_id = query_db(
        """
        INSERT INTO courses (spec_id,title,description,level,link,image,video)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (
            safe_int(data.get("spec_id") or data.get("specialization_id"), None),
            title,
            safe_text(data.get("description")),
            normalize_level(data.get("level")),
            safe_text(data.get("link") or data.get("course_link")),
            image,
            video,
        ),
        commit=True
    )
    return jsonify({"message": "Course added", "id": course_id})


@app.route("/api/courses/<int:course_id>", methods=["PUT"])
@admin_required
def update_course(course_id):
    data = request_data()
    old = query_db("SELECT * FROM courses WHERE id=%s", (course_id,), fetchone=True)
    if not old:
        return jsonify({"error": "Course not found"}), 404
    image = save_file("image") or safe_text(data.get("image")) or old.get("image")
    video = save_file("video") or safe_text(data.get("video")) or old.get("video")
    exec_db(
        """
        UPDATE courses SET spec_id=%s,title=%s,description=%s,level=%s,link=%s,image=%s,video=%s WHERE id=%s
        """,
        (
            safe_int(data.get("spec_id") or data.get("specialization_id"), old.get("spec_id")),
            safe_text(data.get("title")) or old.get("title"),
            safe_text(data.get("description")),
            normalize_level(data.get("level") or old.get("level")),
            safe_text(data.get("link") or data.get("course_link") or old.get("link")),
            image,
            video,
            course_id,
        )
    )
    return jsonify({"message": "Course updated"})


@app.route("/api/courses/<int:course_id>", methods=["DELETE"])
@admin_required
def delete_course(course_id):
    exec_db("DELETE FROM courses WHERE id=%s", (course_id,))
    return jsonify({"message": "Course deleted"})


@app.route("/api/courses/<int:course_id>/open", methods=["POST"])
@student_required
def open_course(course_id):
    course = query_db("SELECT id FROM courses WHERE id=%s", (course_id,), fetchone=True)
    if not course:
        return jsonify({"error": "Course not found"}), 404
    data = get_json()
    video_started = 1 if data.get("video_started") else 0
    completed = 1 if data.get("completed") else 0
    exec_db(
        """
        INSERT INTO course_progress (user_id,course_id,opened,video_started,completed)
        VALUES (%s,%s,1,%s,%s)
        ON DUPLICATE KEY UPDATE opened=1, video_started=GREATEST(video_started,VALUES(video_started)), completed=GREATEST(completed,VALUES(completed))
        """,
        (request.current_user["id"], course_id, video_started, completed)
    )
    compute_user_progress(request.current_user["id"])
    return jsonify({"message": "Course progress tracked"})


@app.route("/api/quizzes", methods=["GET"])
def get_quizzes():
    course_id = request.args.get("course_id")
    sql = "SELECT q.*, c.title AS course_title FROM quizzes q LEFT JOIN courses c ON c.id=q.course_id WHERE 1=1"
    params = []
    if course_id:
        sql += " AND q.course_id=%s"
        params.append(course_id)
    sql += " ORDER BY q.id DESC"
    rows = query_db(sql, tuple(params), fetchall=True) or []
    return jsonify({"quizzes": [normalize_quiz(row) for row in rows]})


@app.route("/api/quizzes/<int:quiz_id>", methods=["GET"])
def get_quiz(quiz_id):
    quiz = query_db("SELECT q.*, c.title AS course_title FROM quizzes q LEFT JOIN courses c ON c.id=q.course_id WHERE q.id=%s", (quiz_id,), fetchone=True)
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404
    questions = query_db("SELECT * FROM quiz_questions WHERE quiz_id=%s ORDER BY id", (quiz_id,), fetchall=True) or []
    return jsonify({"quiz": normalize_quiz(quiz), "questions": [normalize_question(row) for row in questions]})


@app.route("/api/quizzes", methods=["POST"])
@admin_required
def add_quiz():
    data = get_json()
    title = safe_text(data.get("title") or data.get("name"))
    course_id = safe_int(data.get("course_id"), None)
    if not title or not course_id:
        return jsonify({"error": "Quiz title and course are required"}), 400
    quiz_id = query_db("INSERT INTO quizzes (course_id,title) VALUES (%s,%s)", (course_id, title), commit=True)
    questions = data.get("questions") or []
    if not questions and data.get("questions_json"):
        try:
            questions = json.loads(data.get("questions_json"))
        except Exception:
            questions = []
    for q in questions:
        add_question_to_quiz(quiz_id, q)
    return jsonify({"message": "Quiz added", "id": quiz_id})


def add_question_to_quiz(quiz_id, data):
    return query_db(
        """
        INSERT INTO quiz_questions (quiz_id,question,option1,option2,option3,option4,answer)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (
            quiz_id,
            safe_text(data.get("question") or data.get("question_text")),
            safe_text(data.get("option1") or data.get("option_a")),
            safe_text(data.get("option2") or data.get("option_b")),
            safe_text(data.get("option3") or data.get("option_c")),
            safe_text(data.get("option4") or data.get("option_d")),
            safe_text(data.get("answer") or data.get("correct_answer")),
        ),
        commit=True
    )


@app.route("/api/quizzes/<int:quiz_id>/questions", methods=["POST"])
@admin_required
def add_quiz_question(quiz_id):
    question_id = add_question_to_quiz(quiz_id, get_json())
    return jsonify({"message": "Question added", "id": question_id})


@app.route("/api/quizzes/<int:quiz_id>", methods=["DELETE"])
@admin_required
def delete_quiz(quiz_id):
    exec_db("DELETE FROM quizzes WHERE id=%s", (quiz_id,))
    return jsonify({"message": "Quiz deleted"})


@app.route("/api/quizzes/<int:quiz_id>/submit", methods=["POST"])
@student_required
def submit_quiz(quiz_id):
    data = get_json()
    answers = data.get("answers") or {}
    questions = query_db("SELECT * FROM quiz_questions WHERE quiz_id=%s ORDER BY id", (quiz_id,), fetchall=True) or []
    quiz = query_db("SELECT * FROM quizzes WHERE id=%s", (quiz_id,), fetchone=True)
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404
    score = 0
    details = []
    for q in questions:
        qid = str(q["id"])
        given = safe_text(answers.get(qid) or answers.get(q["id"])).lower()
        correct = safe_text(q.get("answer")).lower()
        ok = given == correct or given in ["1", "a"] and correct in ["a", "1", safe_text(q.get("option1")).lower()] or given in ["2", "b"] and correct in ["b", "2", safe_text(q.get("option2")).lower()] or given in ["3", "c"] and correct in ["c", "3", safe_text(q.get("option3")).lower()] or given in ["4", "d"] and correct in ["d", "4", safe_text(q.get("option4")).lower()]
        if ok:
            score += 1
        details.append({"question_id": q["id"], "given": given, "correct": correct, "correct_boolean": bool(ok)})
    total = len(questions)
    percentage = round((score / total) * 100) if total else 0
    attempt_id = query_db(
        """
        INSERT INTO quiz_attempts (user_id,quiz_id,course_id,score,total,score_percentage,answers_json)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (request.current_user["id"], quiz_id, quiz.get("course_id"), score, total, percentage, json.dumps(details)),
        commit=True
    )
    if quiz.get("course_id"):
        exec_db(
            """
            INSERT INTO course_progress (user_id,course_id,opened,completed)
            VALUES (%s,%s,1,%s)
            ON DUPLICATE KEY UPDATE opened=1, completed=GREATEST(completed,VALUES(completed))
            """,
            (request.current_user["id"], quiz.get("course_id"), 1 if percentage >= 60 else 0)
        )
    compute_user_progress(request.current_user["id"])
    return jsonify({"message": "Quiz submitted", "attempt_id": attempt_id, "score": score, "total": total, "score_percentage": percentage, "details": details})


@app.route("/api/jobs", methods=["GET"])
def get_jobs():
    search = safe_text(request.args.get("search"))
    spec_id = request.args.get("specialization_id") or request.args.get("spec_id")
    sql = "SELECT j.*, s.name AS specialization_name FROM jobs j LEFT JOIN specializations s ON s.id=j.specialization_id WHERE 1=1"
    params = []
    if search:
        sql += " AND (j.title LIKE %s OR j.description LIKE %s OR j.skills LIKE %s)"
        params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
    if spec_id:
        sql += " AND j.specialization_id=%s"
        params.append(spec_id)
    sql += " ORDER BY j.id DESC"
    rows = query_db(sql, tuple(params), fetchall=True) or []
    return jsonify({"jobs": [normalize_job(row) for row in rows]})


@app.route("/api/jobs/<int:job_id>", methods=["GET"])
def get_job(job_id):
    row = query_db("SELECT j.*, s.name AS specialization_name FROM jobs j LEFT JOIN specializations s ON s.id=j.specialization_id WHERE j.id=%s", (job_id,), fetchone=True)
    if not row:
        return jsonify({"error": "Job not found"}), 404
    return jsonify({"job": normalize_job(row)})


@app.route("/api/jobs", methods=["POST"])
@admin_required
def add_job():
    data = get_json()
    title = safe_text(data.get("title"))
    if not title:
        return jsonify({"error": "Job title is required"}), 400
    job_id = query_db(
        "INSERT INTO jobs (specialization_id,title,description,skills,salary,link) VALUES (%s,%s,%s,%s,%s,%s)",
        (safe_int(data.get("specialization_id") or data.get("spec_id"), None), title, safe_text(data.get("description")), safe_text(data.get("skills") or data.get("required_skills")), safe_text(data.get("salary")), safe_text(data.get("link") or data.get("job_link"))),
        commit=True
    )
    return jsonify({"message": "Job added", "id": job_id})


@app.route("/api/jobs/<int:job_id>", methods=["PUT"])
@admin_required
def update_job(job_id):
    data = get_json()
    old = query_db("SELECT * FROM jobs WHERE id=%s", (job_id,), fetchone=True)
    if not old:
        return jsonify({"error": "Job not found"}), 404
    exec_db(
        "UPDATE jobs SET specialization_id=%s,title=%s,description=%s,skills=%s,salary=%s,link=%s WHERE id=%s",
        (safe_int(data.get("specialization_id") or data.get("spec_id"), old.get("specialization_id")), safe_text(data.get("title")) or old.get("title"), safe_text(data.get("description")), safe_text(data.get("skills") or data.get("required_skills")), safe_text(data.get("salary")), safe_text(data.get("link") or data.get("job_link")), job_id)
    )
    return jsonify({"message": "Job updated"})


@app.route("/api/jobs/<int:job_id>", methods=["DELETE"])
@admin_required
def delete_job(job_id):
    exec_db("DELETE FROM jobs WHERE id=%s", (job_id,))
    return jsonify({"message": "Job deleted"})


@app.route("/api/certificates", methods=["GET"])
def get_certificates():
    rows = query_db("SELECT c.*, s.name AS specialization_name FROM certificates c LEFT JOIN specializations s ON s.id=c.specialization_id ORDER BY c.id DESC", fetchall=True) or []
    return jsonify({"certificates": rows})


@app.route("/api/certificates", methods=["POST"])
@admin_required
def add_certificate():
    data = get_json()
    name = safe_text(data.get("name"))
    if not name:
        return jsonify({"error": "Certificate name is required"}), 400
    cert_id = query_db(
        "INSERT INTO certificates (specialization_id,name,description,link) VALUES (%s,%s,%s,%s)",
        (safe_int(data.get("specialization_id") or data.get("spec_id"), None), name, safe_text(data.get("description")), safe_text(data.get("link"))),
        commit=True
    )
    return jsonify({"message": "Certificate added", "id": cert_id})


@app.route("/api/certificates/<int:cert_id>", methods=["DELETE"])
@admin_required
def delete_certificate(cert_id):
    exec_db("DELETE FROM certificates WHERE id=%s", (cert_id,))
    return jsonify({"message": "Certificate deleted"})


@app.route("/api/recommendations", methods=["POST"])
@student_required
def recommendations():
    data = get_json()
    interests = safe_text(data.get("interests"))
    skills = safe_text(data.get("skills"))
    preferred_work = safe_text(data.get("preferred_work") or data.get("goal"))
    profile_text = f"{interests} {skills} {preferred_work}".lower()
    specs = [normalize_specialization(row) for row in (query_db("SELECT * FROM specializations", fetchall=True) or [])]
    jobs = [normalize_job(row) for row in (query_db("SELECT j.*, s.name AS specialization_name FROM jobs j LEFT JOIN specializations s ON s.id=j.specialization_id", fetchall=True) or [])]
    recommended_specs = []
    for spec in specs:
        target = f"{spec.get('name','')} {spec.get('description','')} {spec.get('skills','')}"
        score, matches = calculate_match_percentage(profile_text, target)
        lower_name = safe_text(spec.get("name")).lower()
        for key, hints in SPECIALIZATION_HINTS.items():
            if key in lower_name:
                score = max(score, min(100, len([h for h in hints if h in profile_text]) * 18))
        recommended_specs.append({"id": spec.get("id"), "name": spec.get("name"), "match_percentage": score, "matched_skills": matches, "reason": "Matched your interests and skills with specialization content."})
    recommended_jobs = []
    for job in jobs:
        target = f"{job.get('title','')} {job.get('description','')} {job.get('skills','')} {job.get('specialization','')}"
        score, matches = calculate_match_percentage(profile_text, target)
        recommended_jobs.append({"id": job.get("id"), "title": job.get("title"), "match_percentage": score, "matched_skills": matches, "salary": job.get("salary"), "reason": "Matched your profile with job skills and description."})
    recommended_specs.sort(key=lambda item: item["match_percentage"], reverse=True)
    recommended_jobs.sort(key=lambda item: item["match_percentage"], reverse=True)
    fallback = {
        "recommended_specializations": recommended_specs[:5],
        "recommended_jobs": recommended_jobs[:8],
        "roadmap": [
            "Choose the highest matching specialization.",
            "Start beginner courses, then move to intermediate and advanced content.",
            "Complete course quizzes to update your profile progress.",
            "Use ATS tools to improve your resume for the highest matching jobs."
        ],
        "summary": "Recommendations are based on your answers and current database content."
    }
    prompt = f"Return JSON with recommended_specializations, recommended_jobs, roadmap, summary. User interests: {interests}. Skills: {skills}. Preferred work: {preferred_work}. Specs: {json.dumps(specs[:20])}. Jobs: {json.dumps(jobs[:30])}."
    result = ai_json(prompt, fallback)
    query_db("INSERT INTO recommendation_results (user_id,recommendation_json) VALUES (%s,%s)", (request.current_user["id"], json.dumps(result)), commit=True)
    return jsonify(result)


@app.route("/api/recommendations/analyze", methods=["POST"])
@student_required
def recommendations_analyze():
    return recommendations()


@app.route("/api/ats/check", methods=["POST"])
@student_required
def ats_check():
    target_job = safe_text(request.form.get("target_job") or request.form.get("job_description"))
    resume_file = request.files.get("resume_file") or request.files.get("resume")
    if not target_job:
        return jsonify({"error": "Target job or job description is required"}), 400
    if not resume_file:
        return jsonify({"error": "Please upload a PDF, DOCX, or TXT resume"}), 400
    resume_text = extract_resume_text(resume_file)
    if not resume_text.strip():
        return jsonify({"error": "Could not read resume text from this file"}), 400
    score, matched = calculate_match_percentage(resume_text, target_job)
    target_skills = [skill for skill in TECH_SKILLS if skill in target_job.lower()]
    missing = [skill for skill in target_skills if skill not in resume_text.lower()]
    section_scores = {
        "keywords": score,
        "structure": 85 if re.search(r"education|experience|projects|skills", resume_text, re.I) else 55,
        "clarity": 80 if len(resume_text.split()) > 120 else 60,
    }
    ats_score = round((section_scores["keywords"] * 0.5) + (section_scores["structure"] * 0.3) + (section_scores["clarity"] * 0.2))
    fallback = {
        "ats_score": ats_score,
        "score": ats_score,
        "summary": "ATS analysis completed using resume text, target job keywords, and structure checks.",
        "matched_keywords": matched,
        "missing_keywords": missing[:20],
        "strengths": ["Resume file was readable", "Relevant keywords were detected"] if matched else ["Resume file was readable"],
        "weaknesses": ["Add more target-job keywords"] if missing else [],
        "improvements": ["Add measurable achievements", "Add missing role-specific keywords", "Keep sections clear: Summary, Skills, Projects, Education"],
        "section_scores": section_scores,
    }
    prompt = f"Return JSON ATS result with ats_score, summary, matched_keywords, missing_keywords, strengths, weaknesses, improvements, section_scores. Target job: {target_job}. Resume: {resume_text[:5000]}"
    result = ai_json(prompt, fallback)
    ats_score = safe_int(result.get("ats_score") or result.get("score"), ats_score)
    query_db(
        "INSERT INTO ats_results (user_id,target_job,score,summary,matched_keywords,missing_keywords) VALUES (%s,%s,%s,%s,%s,%s)",
        (request.current_user["id"], target_job[:220], ats_score, safe_text(result.get("summary")), json.dumps(result.get("matched_keywords", [])), json.dumps(result.get("missing_keywords", []))),
        commit=True
    )
    return jsonify(result)


@app.route("/api/ats/generate", methods=["POST"])
@student_required
def ats_generate():
    data = get_json()
    name = safe_text(data.get("name") or request.current_user.get("name"))
    target_role = safe_text(data.get("target_role") or data.get("target_job"))
    summary = safe_text(data.get("summary"))
    skills = safe_text(data.get("skills"))
    projects = safe_text(data.get("projects"))
    education = safe_text(data.get("education"))
    if not name or not target_role or not skills or not summary:
        return jsonify({"error": "Name, target role, summary, and skills are required"}), 400
    resume = f"""{name}\n{target_role}\n\nSUMMARY\n{summary}\n\nSKILLS\n{skills}\n\nPROJECTS\n{projects or 'Add relevant projects with tools and measurable results.'}\n\nEDUCATION\n{education or 'Add education details.'}\n"""
    fallback = {
        "resume": resume,
        "enhanced_summary": summary,
        "ats_score": 82,
        "matched_keywords": [skill for skill in TECH_SKILLS if skill in skills.lower()][:12],
    }
    prompt = f"Return JSON with resume, enhanced_summary, ats_score, matched_keywords. Create ATS-friendly resume text, no fake claims. Name: {name}. Role: {target_role}. Summary: {summary}. Skills: {skills}. Projects: {projects}. Education: {education}."
    result = ai_json(prompt, fallback)
    return jsonify(result)


def build_resume_pdf(text):
    if not SimpleDocTemplate:
        return None
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=38, bottomMargin=38)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=14, textColor=colors.HexColor("#111827"))
    heading = ParagraphStyle("Heading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#111827"), spaceBefore=10)
    story = []
    lines = str(text or "").splitlines()
    for line in lines:
        value = line.strip()
        if not value:
            story.append(Spacer(1, 8))
        elif value.isupper() and len(value) < 35:
            story.append(Paragraph(value, heading))
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#d1d5db")))
        else:
            story.append(Paragraph(value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body))
    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route("/api/ats/export/pdf", methods=["POST"])
@student_required
def export_pdf():
    data = get_json()
    text = safe_text(data.get("resume") or data.get("text"))
    if not text:
        return jsonify({"error": "Resume text is required"}), 400
    pdf = build_resume_pdf(text)
    if not pdf:
        return jsonify({"error": "PDF export library is not installed"}), 500
    return send_file(pdf, mimetype="application/pdf", as_attachment=True, download_name="sqr_resume.pdf")


@app.route("/api/ats/export/docx", methods=["POST"])
@student_required
def export_docx():
    if not Document:
        return jsonify({"error": "DOCX export library is not installed"}), 500
    data = get_json()
    text = safe_text(data.get("resume") or data.get("text"))
    if not text:
        return jsonify({"error": "Resume text is required"}), 400
    doc = Document()
    for line in text.splitlines():
        value = line.strip()
        if not value:
            doc.add_paragraph("")
        elif value.isupper() and len(value) < 35:
            doc.add_heading(value, level=2)
        else:
            doc.add_paragraph(value)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="sqr_resume.docx")


@app.route("/api/admin/stats")
@admin_required
def admin_stats():
    def count(table):
        if not table_exists(table):
            return 0
        row = query_db(f"SELECT COUNT(*) AS total FROM `{table}`", fetchone=True) or {"total": 0}
        return safe_int(row.get("total"), 0)
    return jsonify({
        "users": count("users"),
        "specializations": count("specializations"),
        "courses": count("courses"),
        "quizzes": count("quizzes"),
        "jobs": count("jobs"),
        "certificates": count("certificates"),
        "ats_results": count("ats_results"),
    })


@app.route("/api/admin/users")
@admin_required
def admin_users():
    rows = query_db("SELECT id,username,name,email,role,current_mode,banned,created_at FROM users ORDER BY id DESC", fetchall=True) or []
    return jsonify({"users": [clean_user(row) for row in rows]})


@app.route("/api/admin/users/<int:user_id>/ban", methods=["PUT", "POST"])
@admin_required
def ban_user(user_id):
    if user_id == request.current_user["id"]:
        return jsonify({"error": "You cannot ban yourself"}), 400
    exec_db("UPDATE users SET banned=1 WHERE id=%s", (user_id,))
    return jsonify({"message": "User banned"})


@app.route("/api/admin/users/<int:user_id>/unban", methods=["PUT", "POST"])
@admin_required
def unban_user(user_id):
    exec_db("UPDATE users SET banned=0 WHERE id=%s", (user_id,))
    return jsonify({"message": "User unbanned"})


@app.route("/api/admin/users/<int:user_id>/role", methods=["PUT", "POST"])
@admin_required
def change_user_role(user_id):
    data = get_json()
    role = safe_text(data.get("role")).lower()
    if role not in ["student", "admin"]:
        return jsonify({"error": "Role must be student or admin"}), 400
    current_mode = "admin" if role == "admin" else "student"
    exec_db("UPDATE users SET role=%s,current_mode=%s WHERE id=%s", (role, current_mode, user_id))
    if role == "admin":
        exec_db("INSERT IGNORE INTO admins (user_id,admin_level) VALUES (%s,'manager')", (user_id,))
    else:
        exec_db("DELETE FROM admins WHERE user_id=%s", (user_id,))
    return jsonify({"message": "User role updated"})


@app.route("/api/mode", methods=["PUT", "POST"])
@login_required
def switch_mode():
    data = get_json()
    mode = safe_text(data.get("mode")).lower()
    if request.current_user.get("role") != "admin":
        return jsonify({"error": "Only admins can switch mode"}), 403
    if mode not in ["student", "admin"]:
        return jsonify({"error": "Mode must be student or admin"}), 400
    exec_db("UPDATE users SET current_mode=%s WHERE id=%s", (mode, request.current_user["id"]))
    user = query_db("SELECT * FROM users WHERE id=%s", (request.current_user["id"],), fetchone=True)
    return jsonify({"message": "Mode updated", "token": generate_token(user), "user": clean_user(user)})


@app.errorhandler(404)
def not_found(error):
    if request.path.startswith("/api/"):
        return jsonify({"error": "Endpoint not found"}), 404
    return render_template("gp.html"), 404


@app.errorhandler(413)
def too_large(error):
    return jsonify({"error": "File is too large"}), 413


@app.errorhandler(500)
def server_error(error):
    return jsonify({"error": "Server error", "details": str(error)}), 500


try:
    init_db()
    print("SQR database checked")
except Exception as exc:
    print("init_db skipped or failed:", exc)


if __name__ == "__main__":
    app.run(host=os.getenv("FLASK_HOST", "0.0.0.0"), port=int(os.getenv("PORT", os.getenv("FLASK_PORT", 5000))), debug=os.getenv("FLASK_DEBUG", "0") == "1")
